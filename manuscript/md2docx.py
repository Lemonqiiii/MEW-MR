#!/usr/bin/env python3
"""Convert the JITC submission markdown to a formatted Word document."""

import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

INPUT = r"E:\medical-review\manuscript\jitc_submission.md"
OUTPUT = r"E:\medical-review\manuscript\jitc_submission.docx"


def add_run(paragraph, text, bold=False, italic=False, superscript=False, size=None, font_name=None, color=None):
    """Add a run with formatting to a paragraph."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if superscript:
        run.font.superscript = True
    if size:
        run.font.size = size
    if font_name:
        run.font.name = font_name
    if color:
        run.font.color.rgb = color
    return run


def parse_inline_formatting(paragraph, text):
    """Parse inline markdown formatting and add runs to paragraph.
    Handles: **bold**, *italic*, ***bold+italic***, [ref] superscript references
    """
    # First, split on bold markers
    # Pattern: **text** or __text__ for bold
    # Pattern: *text* or _text_ for italic
    # We need to handle nested formatting

    i = 0
    buffer = ""
    while i < len(text):
        # Check for bold+italic ***...***
        if text[i:i+3] == "***" and i + 3 < len(text):
            end = text.find("***", i + 3)
            if end != -1:
                if buffer:
                    add_run(paragraph, buffer)
                    buffer = ""
                add_run(paragraph, text[i+3:end], bold=True, italic=True)
                i = end + 3
                continue
        # Check for bold **...**
        if text[i:i+2] == "**" and i + 2 < len(text):
            end = text.find("**", i + 2)
            if end != -1:
                if buffer:
                    add_run(paragraph, buffer)
                    buffer = ""
                # Check for italic inside bold
                inner = text[i+2:end]
                # Handle *italic* inside bold
                add_run(paragraph, inner, bold=True)
                i = end + 2
                continue
        # Check for italic *...* (but not **)
        if text[i] == "*" and text[i:i+2] != "**" and (i == 0 or text[i-1] != "*"):
            end = text.find("*", i + 1)
            if end != -1 and end > i + 1:
                if buffer:
                    add_run(paragraph, buffer)
                    buffer = ""
                add_run(paragraph, text[i+1:end], italic=True)
                i = end + 1
                continue
        buffer += text[i]
        i += 1
    if buffer:
        add_run(paragraph, buffer)


def process_superscript_refs(paragraph):
    """Process citation references like [1], [1,2], [5-7], [1,2,5-7] to superscript."""
    # We need to process runs and convert reference patterns to superscript
    full_text = paragraph.text
    # Find patterns like [1], [1,2], [5-7], [1,2,5-7], [8–10], [5,6], [8,11], etc.
    ref_pattern = re.compile(r'\[(\d+(?:[,–-]\d+)*(?:[,;]\s*\d+(?:[,–-]\d+)*)*)\]')

    # Rebuild the paragraph by processing each run
    runs = paragraph.runs
    if not runs:
        return

    # Collect all text and references
    new_runs_data = []
    for run in runs:
        text = run.text
        bold = run.bold
        italic = run.italic

        last_end = 0
        for m in ref_pattern.finditer(text):
            # Add text before the reference
            if m.start() > last_end:
                new_runs_data.append({
                    'text': text[last_end:m.start()],
                    'bold': bold,
                    'italic': italic,
                    'superscript': False,
                })
            # Add the reference as superscript
            new_runs_data.append({
                'text': m.group(),
                'bold': bold,
                'italic': italic,
                'superscript': True,
            })
            last_end = m.end()

        # Add remaining text
        if last_end < len(text):
            new_runs_data.append({
                'text': text[last_end:],
                'bold': bold,
                'italic': italic,
                'superscript': False,
            })

    # Clear existing runs
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)

    # Add new runs
    for data in new_runs_data:
        add_run(paragraph, data['text'], bold=data['bold'], italic=data['italic'],
                superscript=data['superscript'])


def convert_md_to_docx(input_path, output_path):
    """Convert a markdown file to a Word document."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

    # Configure heading styles
    for i, (style_name, size, bold) in enumerate([
        ('Heading 1', Pt(16), True),
        ('Heading 2', Pt(14), True),
        ('Heading 3', Pt(13), True),
        ('Heading 4', Pt(12), True),
    ], 1):
        try:
            h_style = doc.styles[style_name]
            h_style.font.name = 'Times New Roman'
            h_style.font.size = size
            h_style.font.bold = bold
            h_style.font.color.rgb = RGBColor(0, 0, 0)
            h_style.paragraph_format.space_before = Pt(12)
            h_style.paragraph_format.space_after = Pt(6)
        except KeyError:
            pass

    # Read the markdown file
    with open(input_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_references = False
    in_key_messages = False

    for line in lines:
        line = line.rstrip()

        # Skip empty lines
        if not line:
            continue

        # Handle horizontal rules
        if line.strip() == '---':
            doc.add_page_break()
            continue

        # Handle headings
        if line.startswith('# '):
            # Level 1 heading - use as title
            text = line[2:]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, text, bold=True, size=Pt(18))
            continue

        if line.startswith('## '):
            in_references = line.startswith('## References')
            text = line[3:]
            p = doc.add_heading(text, level=1)
            continue

        if line.startswith('### '):
            text = line[4:]
            doc.add_heading(text, level=2)
            continue

        if line.startswith('#### '):
            text = line[5:]
            doc.add_heading(text, level=3)
            continue

        if line.startswith('##### '):
            text = line[6:]
            doc.add_heading(text, level=4)
            continue

        # Handle bullet lists
        if line.startswith('- ') or line.startswith('  - '):
            text = re.sub(r'^\s*-\s+', '', line)
            p = doc.add_paragraph(style='List Bullet')
            # Clear default text
            p.clear()
            parse_inline_formatting(p, text)
            process_superscript_refs(p)
            continue

        # Handle numbered lists (1. ...)
        if re.match(r'^\d+\.\s+', line):
            text = re.sub(r'^\d+\.\s+', '', line)
            p = doc.add_paragraph(style='List Number')
            p.clear()
            parse_inline_formatting(p, text)
            process_superscript_refs(p)
            continue

        # Handle Key Messages special lines
        if line.startswith('**What ') or line.startswith('**How '):
            p = doc.add_paragraph()
            parse_inline_formatting(p, line)
            process_superscript_refs(p)
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        parse_inline_formatting(p, line)
        process_superscript_refs(p)

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Save
    doc.save(output_path)
    print(f"Word document saved to: {output_path}")


if __name__ == "__main__":
    convert_md_to_docx(INPUT, OUTPUT)
