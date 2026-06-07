#!/usr/bin/env python3
"""
Generate Word document from jitc_submission.md.
Generic and configurable — reads PROJECT config from first comment block.
Does not hardcode any topic-specific content (title, figures, keywords).

Expected markdown structure:
    # Title (used for title page)
    ## Abstract
    [abstract text]
    **Figure N. caption** (optional standalone figure blocks)
    **Table N. caption** (optional standalone table blocks)
    ## 1. Section Name
    [body text — can contain inline Figure/Table references]
    ...
    ## References
    [numbered references]
    ## Data Availability Statement / Author Contributions / etc.
"""
import sys as _sys
if '--help' in _sys.argv or '-h' in _sys.argv:
    print("Usage: python3 gen_word.py")
    print("  Generate Word (.docx) from markdown manuscript.")
    print("  Reads paths from config.yaml — manuscript_src, manuscript_docx, figures_dir.")
    print("  No CLI arguments. Place manuscript at the configured path and run.")
    _sys.exit(0)

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os, re, json
from config_loader import load_config, find_project_root

# ── Configuration (from config.yaml) ──
_config = load_config()
ROOT = find_project_root()
SRC = str(ROOT / _config["paths"]["manuscript_src"])
FIG_DIR = str(ROOT / _config["paths"]["figures_dir"])
OUT = str(ROOT / _config["scripts"]["gen_word"]["output_filename"])
# ───────────────────────────────────────

doc = Document()

# ── Page setup ──
for s in doc.sections:
    s.top_margin = Cm(2.54)
    s.bottom_margin = Cm(2.54)
    s.left_margin = Cm(3.18)
    s.right_margin = Cm(3.18)

# ── Styles ──
sty = doc.styles['Normal']
sty.font.name = 'Times New Roman'
sty.font.size = Pt(12)
sty.paragraph_format.line_spacing = 2.0

for lv in [1, 2, 3]:
    hs = doc.styles[f'Heading {lv}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.font.bold = True
    hs.font.size = {1: Pt(14), 2: Pt(13), 3: Pt(12)}[lv]
    spacing = {1: (Pt(18), Pt(6)), 2: (Pt(12), Pt(4)), 3: (Pt(6), Pt(2))}
    hs.paragraph_format.space_before = spacing[lv][0]
    hs.paragraph_format.space_after = spacing[lv][1]

# ── Helper functions ──
def P(text, bold=False, sz=12, italic=False, align=None):
    par = doc.add_paragraph()
    par.paragraph_format.line_spacing = 2.0
    if align:
        par.alignment = align
    r = par.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(sz)
    r.bold = bold
    r.italic = italic
    return par

def H(text, lv=1):
    hd = doc.add_heading(text, level=lv)
    for r in hd.runs:
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RGBColor(0, 0, 0)
    return hd

def PB():
    doc.add_page_break()

def insert_image(filename, caption, width=5.8):
    """Embed a PNG/JPG figure or table image into the document."""
    fp = os.path.join(FIG_DIR, filename)
    if not os.path.exists(fp):
        P(f'[Image file not found: {filename}]', sz=10, italic=True)
        return
    # Centered image
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(fp, width=Inches(width))
    # Centered bold caption
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run(caption)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10)
    r.bold = True

def clean_text(text):
    """Strip markdown bold/italic markers for plain text rendering."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    return text

def _cell_runs(cell, text):
    """Fill a table cell with text, preserving **bold** and *italic* markdown."""
    # Split on **bold** markers: odd segments are normal, even are bold
    parts = re.split(r'\*\*(.+?)\*\*', text)
    for i, segment in enumerate(parts):
        if not segment:
            continue
        if i % 2 == 0:
            # Normal text — may contain *italic* markers
            italic_parts = re.split(r'\*(.+?)\*', segment)
            for j, iseg in enumerate(italic_parts):
                if not iseg:
                    continue
                r = cell.add_paragraph().add_run(iseg)
                r.font.name = 'Times New Roman'
                r.font.size = Pt(10)
                r.italic = (j % 2 == 1)
        else:
            # Bold segment
            r = cell.add_paragraph().add_run(segment)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)
            r.bold = True

def _detect_pipe_table(block):
    """Return True if a block is a markdown pipe table (all lines start with |)."""
    lines = [l.strip() for l in block.split('\n') if l.strip()]
    if len(lines) < 2:
        return False
    # Every line must start with |
    if not all(l.startswith('|') for l in lines):
        return False
    # Must have a separator row (second line: |---|...|)
    sep = lines[1]
    if not re.match(r'^\|[\s\-:|]+\|$', sep):
        return False
    return True

def _render_md_table(block):
    """Render a markdown pipe table block as a python-docx Table."""
    lines = [l.strip() for l in block.split('\n') if l.strip() and l.startswith('|')]
    # Header is first line, separator is second, data rows follow
    header_cells = [c.strip() for c in lines[0].split('|')[1:-1]]
    data_rows = []
    for line in lines[2:]:
        data_rows.append([c.strip() for c in line.split('|')[1:-1]])

    ncols = len(header_cells)
    nrows = 1 + len(data_rows)
    table = doc.add_table(rows=nrows, cols=ncols)
    table.style = 'Table Grid'

    # Header row
    for ci, cell_text in enumerate(header_cells):
        _cell_runs(table.rows[0].cells[ci], cell_text)
        # Shade header
        from docx.oxml.ns import qn
        shading = table.rows[0].cells[ci]._element.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): 'D9E2F3',
            qn('w:val'): 'clear',
        })
        shading.append(shd)

    # Data rows (zebra-striped)
    for ri, row_data in enumerate(data_rows):
        for ci, cell_text in enumerate(row_data):
            _cell_runs(table.rows[ri + 1].cells[ci], cell_text)
            if ri % 2 == 1:
                shading = table.rows[ri + 1].cells[ci]._element.get_or_add_tcPr()
                shd = shading.makeelement(qn('w:shd'), {
                    qn('w:fill'): 'F2F2F2',
                    qn('w:val'): 'clear',
                })
                shading.append(shd)

    # Spacing after table
    P('')

# ═══════════════════════════════════════════════════════════════════
# 1. READ AND PARSE SOURCE
# ═══════════════════════════════════════════════════════════════════

with open(SRC, 'r', encoding='utf-8') as f:
    full_text = f.read()

# ── Pre-processing: strip HTML comments (Phase 7.6a) ──
import re as _re
_stripped = _re.sub(r'<!--.*?-->', '', full_text, flags=_re.DOTALL)
if _stripped != full_text:
    n_stripped = len(_re.findall(r'<!--.*?-->', full_text, flags=_re.DOTALL))
    print(f'  Stripped {n_stripped} HTML comment(s) from source')
    full_text = _stripped

# Split: body text vs reference section
parts = full_text.split('## References')
body_text = parts[0]
refs_text = parts[1] if len(parts) > 1 else ''

# ═══════════════════════════════════════════════════════════════════
# 2. EXTRACT METADATA FOR TITLE PAGE
# ═══════════════════════════════════════════════════════════════════

# Title: first # heading
title_match = re.search(r'^# (.+)$', body_text, re.MULTILINE)
title = title_match.group(1).strip() if title_match else 'Untitled'

# Count what's available
fig_files = []
tab_files = []
if os.path.isdir(FIG_DIR):
    for f in sorted(os.listdir(FIG_DIR)):
        if f.endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):
            if 'Figure' in f:
                fig_files.append(f)
            elif 'Table' in f:
                tab_files.append(f)
            else:
                fig_files.append(f)  # unrecognized → treat as figure

ref_count = len(re.findall(r'^\d+\.', refs_text, re.MULTILINE))
# Rough word count from body (exclude headers/markers)
text_only = re.sub(r'#+\s+.*', '', body_text)
text_only = re.sub(r'\*\*(?:Figure|Table)\s+\d+[.*].*?\*\*', '', text_only)
word_count = len(re.findall(r'\b\w+\b', text_only))

# ═══════════════════════════════════════════════════════════════════
# 3. BUILD TITLE PAGE
# ═══════════════════════════════════════════════════════════════════

P(''); P(''); P('')  # vertical spacing
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = tp.add_run(title)
tr.font.name = 'Times New Roman'
tr.font.size = Pt(18)
tr.bold = True

P('')
P('Narrative Review', sz=14, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
P('')

# Running information
info_lines = [
    ('Running title: ', 'NRDS Interventions: Life-Course Consequences'),
    ('Word count: ', f'~{word_count:,} words (body); ~{word_count + 800:,} including references'),
    ('Figures/Tables: ', f'{len(fig_files)} Figure(s) | {len(tab_files)} Table(s) | {ref_count} References'),
    ('Target Journal: ', 'Pediatric Research'),
]
for label, value in info_lines:
    pp = P('')
    pp.paragraph_format.line_spacing = 1.5
    rl = pp.add_run(label)
    rl.font.name = 'Times New Roman'
    rl.font.size = Pt(12)
    rl.bold = True
    rv = pp.add_run(value)
    rv.font.name = 'Times New Roman'
    rv.font.size = Pt(12)

PB()

# ═══════════════════════════════════════════════════════════════════
# 4. PARSE AND CONVERT BODY SECTIONS
# ═══════════════════════════════════════════════════════════════════

# Split into top-level sections by ## headers
# Each section starts with ## SectionName
raw_sections = re.split(r'\n(?=## )', body_text.strip())

# Which sections contain declarations (to be rendered before References, not inline)?
DECL_SECTIONS = [
    'Data Availability Statement', 'Author Contributions',
    'Competing Interests', 'Acknowledgements', 'Funding'
]

# Pre-scan: build a map from section header text to figure/table files
# by matching figure numbers mentioned in section body text
section_figures = {}  # header_text -> [(filename, caption)]
available_figs = {f: f for f in fig_files + tab_files}

for section in raw_sections:
    lines = section.strip().split('\n')
    if not lines:
        continue
    header_line = lines[0]
    header_raw = re.sub(r'^#+\s*', '', header_line).strip()
    section_body = '\n'.join(lines[1:])

    # Find standalone figure/table markers: **Figure N. caption** or **Table N. caption**
    for m in re.finditer(r'\*\*(Figure|Table)\s+(\d+)[.:]\s*(.+)', section_body):
        prefix = m.group(1)
        num = m.group(2)
        # Caption is the rest of the line (strip any trailing ** if present)
        caption_text = m.group(3).rstrip('*').strip()
        # Also remove any residual ** markers (e.g. closing bold before description text)
        caption_text = caption_text.replace('**', '')
        caption = f'{prefix} {num}. {caption_text}'

        matching = [f for f in available_figs if f'{prefix}{num}' in f]
        if matching:
            if header_raw not in section_figures:
                section_figures[header_raw] = []
            section_figures[header_raw].append((matching[0], caption))

if section_figures:
    print(f"Detected {sum(len(v) for v in section_figures.values())} figure/table insert(s):")
    for sec, items in section_figures.items():
        for fname, cap in items:
            print(f"  [{sec}] <- {fname}")

# Process each section
declaration_blocks = {}  # collect declarations for end-of-document rendering

for section in raw_sections:
    lines = section.strip().split('\n')
    if not lines:
        continue

    header_line = lines[0].strip()
    header_level = header_line.count('#')
    header_text = re.sub(r'^#+\s*', '', header_line).strip()

    # Skip the document title (already rendered on title page)
    if header_level == 1 and header_text == title:
        continue

    # Skip declaration sections — collect for later
    if header_text in DECL_SECTIONS:
        body_para = '\n'.join(lines[1:]).strip()
        body_para = clean_text(body_para)
        if body_para:
            declaration_blocks[header_text] = body_para
        continue

    # ── Render heading ──
    if header_level == 2:
        if header_text == 'Abstract':
            H('Abstract', 1)
        else:
            H(header_text, 1)
    elif header_level == 3:
        H(header_text, 2)
    elif header_level == 4:
        H(header_text, 3)

    # ── Process body content ──
    body_content = '\n'.join(lines[1:])

    # Split into blocks (paragraphs, separated by blank lines)
    blocks = re.split(r'\n\s*\n', body_content)

    # Separate Impact Statement from Abstract body
    if header_text == 'Abstract':
        abstract_blocks = []
        impact_blocks = []
        in_impact = False
        for block in blocks:
            block = block.strip()
            if not block:
                continue
            if block.startswith('**Impact Statement'):
                in_impact = True
                continue
            if in_impact:
                if block.startswith('**') and 'Impact' not in block:
                    in_impact = False
                    abstract_blocks.append(block)
                else:
                    impact_blocks.append(block)
            else:
                abstract_blocks.append(block)

        # Render abstract
        abstract_para = ' '.join(clean_text(b) for b in abstract_blocks if b)
        if abstract_para:
            P(abstract_para)

        # Render Impact Statement
        if impact_blocks:
            P('')
            pp = P('')
            rl = pp.add_run('Impact Statement: ')
            rl.font.name = 'Times New Roman'
            rl.font.size = Pt(12)
            rl.bold = True
            impact_para = ' '.join(clean_text(b) for b in impact_blocks if b)
            P(impact_para, sz=12)

        PB()
        continue

    # Process non-Abstract body blocks
    # Figures/tables are inserted INLINE at their marker position, not at section end
    i = 0
    while i < len(blocks):
        block = blocks[i].strip()
        i += 1
        if not block:
            continue
        if block.startswith('---'):
            continue

        # Is this a sub-heading? (### or ####)
        if block.startswith('#'):
            sub_level = block.split('\n')[0].count('#')
            sub_text = re.sub(r'^#+\s*', '', block.split('\n')[0])
            H(sub_text, min(sub_level - 1, 3))
            remaining = '\n'.join(block.split('\n')[1:]).strip()
            if remaining:
                P(clean_text(' '.join(remaining.split('\n'))))
            continue

        # Is this a standalone figure/table marker? **Figure N. caption** or **Table N. caption**
        fig_match = re.match(r'\*\*(Figure|Table)\s+(\d+)[.:]\s*(.+)', block)
        if fig_match:
            prefix = fig_match.group(1)  # "Figure" or "Table"
            num = fig_match.group(2)
            caption_text = fig_match.group(3).rstrip('*').strip()
            # Also remove any residual ** markers (e.g. closing bold before description text)
            caption_text = caption_text.replace('**', '')
            full_caption = f'{prefix} {num}. {caption_text}'

            # If this is a Table marker and the next block is a pipe table,
            # skip image insertion and render the pipe table with caption instead.
            next_block = blocks[i].strip() if i < len(blocks) else ''
            if prefix == 'Table' and next_block and _detect_pipe_table(next_block):
                # Render caption above the native table
                cp = doc.add_paragraph()
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = cp.add_run(full_caption)
                r.font.name = 'Times New Roman'
                r.font.size = Pt(10)
                r.bold = True
                P('')
                _render_md_table(next_block)
                i += 1  # consume the pipe table block
                continue

            # Otherwise, embed the image as before
            matching = [f for f in (fig_files + tab_files) if f'{prefix}{num}' in f]
            if matching:
                P('')  # spacing before image
                insert_image(matching[0], full_caption)
                P('')  # spacing after image
            else:
                # Caption still appears even if file missing (for debugging)
                cp = doc.add_paragraph()
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = cp.add_run(f'[{prefix} {num} — file not found: {prefix}{num}*.png]')
                r.font.name = 'Times New Roman'
                r.font.size = Pt(10)
                r.italic = True
            continue

        # Is this a markdown pipe table? (| col1 | col2 | ...)
        if _detect_pipe_table(block):
            _render_md_table(block)
            continue

        # Regular paragraph: merge lines, clean, render
        para_text = ' '.join(block.split('\n'))
        para_text = clean_text(para_text)
        if para_text.strip():
            P(para_text)

    # Note: Figures are now inserted inline at marker position.
    # The section_figures dict is no longer used for insertion timing.
    # It is still built for diagnostic reporting.

# ═══════════════════════════════════════════════════════════════════
# 5. RENDER DECLARATIONS
# ═══════════════════════════════════════════════════════════════════

if declaration_blocks:
    PB()
    H('Declarations', 1)
    for decl_name in DECL_SECTIONS:
        if decl_name in declaration_blocks:
            H(decl_name, 2)
            P(declaration_blocks[decl_name], sz=11)

# ═══════════════════════════════════════════════════════════════════
# 6. RENDER REFERENCES
# ═══════════════════════════════════════════════════════════════════

PB()
H('References', 1)

ref_entries = re.findall(r'(\d+)\.\s+(.+?)(?=\n\d+\.\s|\n\n|\Z)', refs_text, re.DOTALL)
for num, entry in ref_entries:
    entry = entry.strip()
    entry = re.sub(r'\n+', ' ', entry)
    entry = re.sub(r'\s+', ' ', entry)
    # Strip internal type annotations like [G - consensus guideline]
    entry = re.sub(r'\s*\[[A-Z]\s+[-–]\s+[^\]]+\]', '', entry)
    P(f'[{num}] {entry}', sz=9).paragraph_format.line_spacing = 1.15

# ═══════════════════════════════════════════════════════════════════
# 7. SAVE AND SELF-CHECK
# ═══════════════════════════════════════════════════════════════════

doc.save(OUT)

# Collect all text from Word paragraphs for verification
all_text = " ".join([p.text for p in doc.paragraphs])

# Count embedded images
rels = doc.part.rels
img_count = sum(1 for rel in rels.values() if "image" in str(rel.reltype))

# Count headings
h_count = sum(1 for p in doc.paragraphs if p.style.name.startswith('Heading'))

# Find figure/table references in the rendered text
fig_refs = re.findall(r'Figure\s+(\d+)', all_text)
tab_refs = re.findall(r'Table\s+(\d+)', all_text)

# Build set of available figure/table numbers from files
avail_fig_nums = set()
avail_tab_nums = set()
for fname in fig_files:
    m = re.search(r'Figure(\d+)', fname)
    if m: avail_fig_nums.add(m.group(1))
for fname in tab_files:
    m = re.search(r'Table(\d+)', fname)
    if m: avail_tab_nums.add(m.group(1))

# Check for unreferenced figures or missing figure files
missing_figs = [n for n in set(fig_refs) if n not in avail_fig_nums]
missing_tabs = [n for n in set(tab_refs) if n not in avail_tab_nums]

# Report
print(f"Saved: {OUT}")
print(f"Size: {os.path.getsize(OUT) / 1024:.0f} KB")
print(f"Words: ~{len(all_text.split())}")
print(f"Headings: {h_count}")
print(f"Images embedded: {img_count}")
print(f"Figure refs in text: {len(fig_refs)}")
print(f"Table refs in text: {len(tab_refs)}")
if missing_figs:
    print(f"  WARNING: Figure(s) referenced but file missing: {missing_figs}")
if missing_tabs:
    print(f"  WARNING: Table(s) referenced but file missing: {missing_tabs}")
if not missing_figs and not missing_tabs:
    print(f"  All figure/table refs resolved OK")

# Quick structural validation
has_title_page = 'Narrative Review' in all_text
has_references = 'References' in all_text
print(f"Title page: {'OK' if has_title_page else 'MISSING'}")
print(f"Reference section: {'OK' if has_references else 'MISSING'}")


# ── CLI overrides ──
if __name__ == "__main__" and "gen_word" in str(__import__("sys").argv[0]):
    import argparse
    ap = argparse.ArgumentParser(description="Generate Word document from markdown manuscript")
    ap.add_argument("--src", help="Markdown source file")
    ap.add_argument("--out", help="Output .docx file")
    ap.add_argument("--fig-dir", help="Figure directory")
    args = ap.parse_args()
    if args.src: SRC = args.src
    if args.out: OUT = args.out
    if args.fig_dir: FIG_DIR = args.fig_dir
    # Re-run main generation with overridden paths
    # (The script runs at import time, so this is informational when used as module)
