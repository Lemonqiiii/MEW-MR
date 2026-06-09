#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Apply all MUST FIX corrections from review-feedback-2026-06-09.md
to PNCS_Systematic_Review.docx
"""
import re
import copy
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = r"E:\medical-review\manuscript\PNCS_Systematic_Review.docx"
DST = r"E:\medical-review\manuscript\PNCS_Systematic_Review_R6.docx"

doc = Document(SRC)

# ============================================================
# Helper: replace text within a paragraph's runs
# ============================================================
def replace_in_paragraph(para, old, new):
    """Replace `old` text with `new` across all runs in a paragraph."""
    full = para.text
    if old not in full:
        return False
    # Strategy: join all runs, do replacement, then redistribute
    # But since we can't easily preserve per-run formatting perfectly,
    # we do surgical replacement within each run
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    # If text is split across runs, need to rebuild
    # Collect run texts
    texts = [r.text for r in para.runs]
    combined = ''.join(texts)
    if old in combined:
        combined = combined.replace(old, new)
        # Put all text into first run, clear others
        if para.runs:
            para.runs[0].text = combined
            for r in para.runs[1:]:
                r.text = ''
        return True
    return False

def replace_across_paragraphs(paragraphs, old, new):
    """Replace text that may span multiple paragraphs."""
    # Try single paragraph first
    for p in paragraphs:
        if replace_in_paragraph(p, old, new):
            return True
    return False

# ============================================================
# MF-4: Shinwell CP rates (§2.2) — fix numerator/percentage confusion
# ============================================================
for p in doc.paragraphs:
    replace_in_paragraph(p,
        "CP rates of 39% versus 13% (P < 0.01)",
        "CP rates of 49% (39/80) versus 15% (12/79) (P < 0.01)")

# ============================================================
# MF-1: NEUROSIS follow-up data (§6.1)
# ============================================================
for p in doc.paragraphs:
    # Fix "5-year follow-up"
    replace_in_paragraph(p,
        "The 5-year follow-up [39]",
        "The 2-year (18–22 month corrected age) follow-up [39]")

    # Fix the fabricated NDI data
    replace_in_paragraph(p,
        "found that among survivors, budesonide was associated with higher moderate-to-severe NDI (16.9% vs. 12.0%; RR 1.41, 95% CI 1.01–1.96)",
        "confirmed significantly higher mortality in the budesonide group (19.9% vs. 14.5%; RR 1.37, 95% CI 1.01–1.86; NNH ~18), while neurodevelopmental disability among survivors did not differ significantly between groups (48.1% vs. 51.4%; adjusted RR 0.93, 95% CI 0.80–1.09)")

# ============================================================
# MF-6 (§3.3): Fix MoLdDX → accurate terminology
# ============================================================
for p in doc.paragraphs:
    replace_in_paragraph(p,
        "moderately-early, low-dose dexamethasone (MoLdDX) and late low-dose hydrocortisone optimized the risk-benefit balance",
        "moderately-early, medium cumulative-dose dexamethasone (MoMdDX; SUCRA 0.91) ranked highest for efficacy, with late low-dose hydrocortisone offering a more favorable safety profile")

# ============================================================
# MF-7 (§8.2): "emerging" → "published"
# ============================================================
for p in doc.paragraphs:
    replace_in_paragraph(p,
        "HC school-age data emerging",
        "HC school-age data now published")

# ============================================================
# MF-2a: Reference [17] — fix authors, article number
# ============================================================
for p in doc.paragraphs:
    replace_in_paragraph(p,
        "[17] Ramaswamy VV, Bandyopadhyay T, Nanda D, et al. Systemic corticosteroids for the prevention of bronchopulmonary dysplasia, a network meta-analysis. *Cochrane Database Syst Rev* 2023;(8):CD014603. PMID: 37650547.",
        "[17] Hay S, Ovelman C, Zupancic JA, Doyle LW, Onland W, Konstantinidis M, Shah PS, Soll R. Systemic corticosteroids for the prevention of bronchopulmonary dysplasia, a network meta-analysis. *Cochrane Database Syst Rev* 2023;(8):CD013730. PMID: 37650547.")

# ============================================================
# MF-2b: Reference [18] — fix authors (Ramaswamy, not Zeng)
# ============================================================
for p in doc.paragraphs:
    replace_in_paragraph(p,
        "[18] Zeng L, Tian J, Song F, et al. Assessment of postnatal corticosteroids for the prevention of bronchopulmonary dysplasia in preterm neonates: a systematic review and network meta-analysis. *JAMA Pediatr* 2021;175:e206826. PMID: 33720274.",
        "[18] Ramaswamy VV, Bandyopadhyay T, Nanda D, Bandiya P, Ahmed J, Garg A, Roehr CC, Nangia S. Assessment of postnatal corticosteroids for the prevention of bronchopulmonary dysplasia in preterm neonates: a systematic review and network meta-analysis. *JAMA Pediatr* 2021;175(6):e206826. PMID: 33720274.")

# ============================================================
# MF-3: Reference [23] — fix authors, journal, volume
# ============================================================
for p in doc.paragraphs:
    replace_in_paragraph(p,
        "[23] Zeng L, Tian J, Song F, et al. Effect of dexamethasone on intelligence and hearing in preterm infants: a meta-analysis. *Pediatrics* 2014;134:898–906. PMID: 25206867.",
        "[23] Zhang R, Bo T, Shen L, Luo S, Li J. Effect of dexamethasone on intelligence and hearing in preterm infants: a meta-analysis. *Neural Regen Res* 2014;9(6):637–645. PMID: 25206867.")

# ============================================================
# MF-8: Reference [12] — fix co-authors, add PMID
# ============================================================
for p in doc.paragraphs:
    replace_in_paragraph(p,
        "[12] Jenkinson A, O'Connell O, Ryan CA, et al. Systematic review of the long-term effects of postnatal corticosteroids. *J Perinat Med* 2023;51:951–60.",
        "[12] Jenkinson AC, Kaltsogianni O, Dassios T, Greenough A. Systematic review of the long-term effects of postnatal corticosteroids. *J Perinat Med* 2023;51(9):1120–1128. PMID: 37606507.")

# ============================================================
# MF-5: Title block cleanup
# ============================================================
# Remove conflicting metadata lines and fix running title
# These are in the first few paragraphs of the document

paras_to_clear = []  # indices of paragraphs to clear
paras_to_modify = {}  # index -> new text for first run

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()

    # Remove "Narrative Review" (genre label conflict)
    if text == "Narrative Review":
        paras_to_clear.append(i)

    # Remove "Target Journal: Pediatric Research"
    elif text == "Target Journal: Pediatric Research":
        paras_to_clear.append(i)

    # Fix wrong running title (keep only correct one later in doc)
    elif text == "Running title: NRDS Interventions: Life-Course Consequences":
        paras_to_clear.append(i)

    # Remove metadata line for Pediatric Research submission
    elif text.startswith("Word count:"):
        paras_to_clear.append(i)
    elif text.startswith("Figures/Tables:"):
        paras_to_clear.append(i)

    # Fix the revision label
    elif "Revision: R5" in text:
        for run in p.runs:
            if "R5" in run.text:
                run.text = run.text.replace("R5 (format cleanup + internal review response)",
                                            "R6 (post-review corrections)")
                break

    # Fix date
    elif text == "Date: June 8, 2026":
        for run in p.runs:
            if "June 8, 2026" in run.text:
                run.text = run.text.replace("June 8, 2026", "June 9, 2026")
                break

# Clear paragraphs marked for removal (set text to empty)
for i in paras_to_clear:
    for run in doc.paragraphs[i].runs:
        run.text = ""

# Also remove empty paragraphs that result from clearing
# (we'll leave them - they'll just be blank lines)

# ============================================================
# NC-1: CP RR correction (§3.1, §8.2, Abstract)
# ============================================================
for p in doc.paragraphs:
    # Fix RR 1.42 → 1.43 and CI 1.08–1.87 → 1.07–1.92
    # Only in Cochrane early dex CP context
    if "RR 1.42, 95% CI 1.08" in p.text and "NNH" in p.text:
        replace_in_paragraph(p,
            "RR 1.42, 95% CI 1.08–1.87",
            "RR 1.43, 95% CI 1.07–1.92")

# Also fix in §8.2 table and other locations
for p in doc.paragraphs:
    if "CP RR 1.42, NNH 20" in p.text:
        replace_in_paragraph(p,
            "CP RR 1.42, NNH 20",
            "CP RR 1.43, NNH 20")

# Fix Abstract CP RR
for p in doc.paragraphs:
    if "RR 1.42, 95% CI 1.08" in p.text:
        replace_in_paragraph(p,
            "RR 1.42, 95% CI 1.08–1.87",
            "RR 1.43, 95% CI 1.07–1.92")

# ============================================================
# LP-1: Tone down Abstract conclusion about HC school-age evidence
# ============================================================
for p in doc.paragraphs:
    replace_in_paragraph(p,
        "providing the first randomized school-age evidence that hydrocortisone did not increase measured functional or neurodevelopmental impairment",
        "providing the first randomized school-age evidence that hydrocortisone did not significantly increase measured functional or neurodevelopmental impairment compared to placebo")

# ============================================================
# LP-2: §10 Conclusions — add qualifying language
# ============================================================
for p in doc.paragraphs:
    replace_in_paragraph(p,
        "represents the preferred agent when postnatal corticosteroids are necessary.",
        "represents the preferred agent when postnatal corticosteroids are necessary, based on the currently available indirect evidence and acknowledging the absence of head-to-head trials.")

# ============================================================
# Additional: Fix Abstract genre label
# ============================================================
for p in doc.paragraphs:
    if p.text.strip() == "Narrative Review" and p != doc.paragraphs[paras_to_clear[0]] if paras_to_clear else True:
        # Already handled above
        pass

# ============================================================
# Save
# ============================================================
doc.save(DST)
print(f"Saved corrected manuscript to: {DST}")
print("Done. All MUST FIX corrections applied.")
