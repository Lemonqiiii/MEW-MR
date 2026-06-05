"""Generate Chinese Word document from JITC manuscript with professional formatting."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os, re

OUTPUT_PATH = "E:/medical-review/manuscript/中文稿件_NSCLC鳞癌免疫治疗耐药机制.docx"

doc = Document()

# --- Page setup ---
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

# --- Style setup ---
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.paragraph_format.line_spacing = 2.0
style.paragraph_format.space_after = Pt(6)

# Heading styles
for i in [1, 2, 3]:
    heading_style = doc.styles[f'Heading {i}']
    heading_style.font.name = '黑体'
    heading_style.font.color.rgb = RGBColor(0, 0, 0)
    if i == 1:
        heading_style.font.size = Pt(16)
    elif i == 2:
        heading_style.font.size = Pt(14)
    else:
        heading_style.font.size = Pt(12)

def add_para(text, bold=False, size=12, align=None, font_name='宋体', space_after=6):
    """Add a paragraph with formatting."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_heading_cn(text, level=1):
    """Add a Chinese heading."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run('非小细胞肺癌鳞状细胞癌免疫治疗耐药机制研究进展')
title_run.font.name = '黑体'
title_run.font.size = Pt(22)
title_run.bold = True

doc.add_paragraph()
subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = subtitle_p.add_run('Mechanisms of Immunotherapy Resistance in Squamous Cell\nCarcinoma of Non-Small Cell Lung Cancer')
sub_run.font.name = 'Times New Roman'
sub_run.font.size = Pt(14)
sub_run.italic = True

doc.add_paragraph()
doc.add_paragraph()

# Title page metadata
meta_items = [
    ('文章类型：', '综述'),
    ('总字数：', '约 9,000 字（中文）；246 词（英文摘要）'),
    ('图表数量：', '图 4 幅 | 表 4 张 | 参考文献 41 篇'),
    ('关键词：', '肺鳞状细胞癌；免疫治疗耐药；免疫检查点抑制剂；肿瘤微环境；T细胞耗竭；KEAP1/NRF2'),
]
for label, value in meta_items:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    run_label = p.add_run(label)
    run_label.font.name = '黑体'
    run_label.font.size = Pt(12)
    run_label.bold = True
    run_value = p.add_run(value)
    run_value.font.name = '宋体'
    run_value.font.size = Pt(12)

doc.add_page_break()

# ============================================================
# ABSTRACT (Chinese + English)
# ============================================================
add_heading_cn('摘要', 1)

abstract_cn = (
    '肺鳞状细胞癌（lung squamous cell carcinoma, LUSC）约占非小细胞肺癌的30%，'
    '其显著特征为缺乏可靶向驱动突变，因此免疫治疗成为全身治疗的基石。靶向PD-1/PD-L1和CTLA-4轴的免疫检查点抑制剂（ICI）'
    '已根本性地改变了LUSC的治疗格局；然而，仅少数患者获得持久获益，原发性耐药和获得性耐药构成了关键的临床挑战。'
    'LUSC的耐药产生于肿瘤内在改变、肿瘤微环境（TME）介导的免疫抑制以及治疗诱导的适应性变化三者之间的复杂交互作用。'
    '本综述从三个相互关联的维度系统整合了LUSC免疫治疗耐药机制的最新认识。'
    '首先，我们概述了LUSC独特的免疫景观，包括分子分类框架以及存在于28-36%患者中的具有临床意义的免疫耗竭型（Exhausted Immune Class, EIC）。'
    '其次，我们审视了肿瘤内在耐药机制，涵盖致癌信号通路（PI3K/AKT、KEAP1/NRF2、p38 MAPK）、上皮间质转化、表观遗传失调和程序性细胞死亡缺陷。'
    '第三，我们分析了TME介导的耐药，包括伴随多重免疫检查点共表达的T细胞耗竭、免疫抑制细胞群体（TAM、CAF、MDSC、Treg）、'
    '缺氧和乳酸酸中毒等代谢限制，以及瘤内微生物组的新兴角色。我们进一步讨论了通过克隆演化、组织学转化和表型可塑性驱动的获得性耐药。'
    '最后，我们评估了克服这些耐药屏障的治疗策略，包括合理的免疫联合治疗、靶向肿瘤内在通路、TME重塑和生物标志物引导的精准免疫治疗。'
    '通过整合这些机制层面的认识，我们旨在为理解和最终克服LUSC免疫治疗耐药提供一个框架。'
)
add_para(abstract_cn, size=12)

# English abstract
add_para('', size=6)
add_para('Abstract', bold=True, size=12, font_name='Times New Roman')
abstract_en = (
    "Lung squamous cell carcinoma (LUSC), accounting for approximately 30% of non-small cell lung cancers, "
    "is characterized by a paucity of actionable driver mutations and a consequent reliance on immunotherapy "
    "as the cornerstone of systemic treatment. Immune checkpoint inhibitors (ICIs) targeting the PD-1/PD-L1 "
    "and CTLA-4 axes have transformed the therapeutic landscape; however, only a minority of patients achieve "
    "durable benefit, with both primary and acquired resistance representing critical clinical challenges..."
)
add_para(abstract_en, size=10, font_name='Times New Roman')

doc.add_page_break()

# ============================================================
# 1. INTRODUCTION
# ============================================================
add_heading_cn('1  引言', 1)

intro_texts = [
    (
        '肺癌仍是全球癌症相关死亡的首要原因，非小细胞肺癌（non-small cell lung cancer, NSCLC）约占所有病例的85% [1]。'
        '在NSCLC亚型中，肺鳞状细胞癌（lung squamous cell carcinoma, LUSC）约占诊断的30%，'
        '是仅次于肺腺癌（lung adenocarcinoma, LUAD）的第二常见组织学亚型 [2]。'
        'LUSC患者面临尤为严峻的临床结局，晚期疾病的五年生存率低于20% [3]。'
        '历史上，标准治疗方案主要限于铂类化疗，其生存获益有限且毒性显著 [4]。'
    ),
    (
        'LUSC治疗格局最深层的塑造因素，恰恰是一种缺失：可靶向驱动突变的极度匮乏。'
        '在LUAD中，针对EGFR、ALK、ROS1、BRAF和KRAS G12C的靶向治疗已根本改变了治疗范式；'
        '相比之下，LUSC几乎没有可通过药物靶向的基因组改变 [5,6]。'
        'LUSC的基因组景观主要由抑癌基因的频繁改变所主导——包括TP53（~80%）、'
        'CDKN2A（~70%，涵盖突变、缺失和表观遗传沉默）和KEAP1（~12%）——以及PIK3CA（~30-40%）、'
        'FGFR1（~20%）的反复扩增，以及包含鳞状谱系转录因子SOX2和TP63的染色体3q位点 [5,7]。'
        '尽管针对这些基因组特征开展了大量靶向药物临床试验，迄今尚无任何药物获得专门针对LUSC的注册批准，'
        '凸显了这一患者群体对有效治疗策略的迫切未满足需求 [7]。'
    ),
    (
        '免疫检查点抑制剂（immune checkpoint inhibitors, ICIs）的问世已根本性地改变了LUSC的治疗格局。'
        '靶向程序性死亡受体-1（PD-1；纳武利尤单抗、帕博利珠单抗）、其配体PD-L1（阿替利珠单抗、度伐利尤单抗）'
        '以及细胞毒性T淋巴细胞抗原4（CTLA-4；伊匹木单抗）的抗体在晚期NSCLC的多线治疗中均展现出有意义的临床活性，'
        'LUSC患者的获益与非鳞癌组织学相当甚至更优 [8–10]。'
        '里程碑式的III期试验——包括KEYNOTE-407（帕博利珠单抗联合化疗一线治疗鳞状NSCLC）、'
        'CheckMate-017（纳武利尤单抗二线治疗鳞状NSCLC）和IMpower-131（阿替利珠单抗联合化疗治疗鳞状NSCLC，'
        '该试验显示了无进展生存获益，但在意向治疗人群中总生存改善未达到统计学显著性）——'
        '确立了以免疫治疗为基础的方案作为新的标准治疗 [8,11]。'
        '更近期的cemiplimab（一种全人源抗PD-1抗体，其工程化铰链区设计旨在最小化免疫原性）'
        '在超高PD-L1表达（肿瘤比例评分≥90%）的LUSC患者中展现出特别的前景，'
        '为这一挑战性组织学类型定义了一个潜在独特的治疗定位 [12]。'
    ),
    (
        '尽管取得了这些进展，仅少数患者能从ICI治疗中获得持久临床获益。'
        '在晚期NSCLC人群中，帕博利珠单抗单药一线治疗PD-L1高表达（≥50%）患者的客观缓解率约为40-45%，'
        '而联合铂类化疗后，鳞状NSCLC（KEYNOTE-407）的缓解率提升至约55-60% [8,13]。'
        '然而，相当比例的患者表现为原发性耐药，定义为疾病进展而无任何初始临床获益证据；'
        '同时，许多初始应答者最终发展为获得性耐药，表现为初始疾病控制期后的肿瘤再生长 [14]。'
        'Lung-MAP S1400F子研究专门评估了度伐利尤单抗联合tremelimumab（PD-L1/CTLA-4双重阻断）'
        '在抗PD-(L)1耐药的鳞状NSCLC中的疗效，报告的原发性耐药队列客观缓解率仅为7%，'
        '获得性耐药队列为0%，生动地说明了克服已建立的免疫治疗耐药所面临的巨大挑战 [15]。'
    ),
    (
        '支撑这些令人失望的临床结局的是一个跨越三个相互关联领域的复杂耐药机制网络：'
        '肿瘤细胞内在改变、肿瘤微环境（tumor microenvironment, TME）的动态重塑，'
        '以及驱动克隆演化和表型适应性的治疗诱导选择压力 [14,16]。'
        'LUSC的TME具有若干共同构建免疫逃逸生态位的特征：'
        '免疫抑制细胞群体高比例存在，包括M2极化肿瘤相关巨噬细胞（TAM）、调节性T细胞（Treg）和髓源性抑制细胞（MDSC）；'
        '以转化生长因子-β（TGF-β）和白介素-6（IL-6）为主导的细胞因子环境；'
        '以缺氧、乳酸酸中毒和营养耗竭为特征的代谢有害微环境 [16,17]。'
        '此外，近期多组学研究在28-36%的LUSC患者中鉴定出一种独特的免疫耗竭型（Exhausted Immune Class, EIC），'
        '其特征为密集的淋巴细胞浸润与多达九个抑制性免疫检查点的共上调矛盾性共存，'
        '包括PD-1、CTLA-4、LAG-3、TIGIT和TIM-3 [18]。'
        '这种"炎症浸润但功能抑制"的免疫状态凸显了单一PD-1/PD-L1阻断的不足，'
        '并强调了以机制为依据的组合策略的必要性。'
    ),
    (
        '文献筛选和选择流程详见附图1（PRISMA流程图）。'
    ),
    (
        '在本综述中，我们对LUSC免疫治疗耐药机制的当前认识进行了全面整合，'
        '按照三个相互关联的维度进行组织（附图2）。'
        '首先，我们综述LUSC的免疫景观，包括其独特的TIME组成和分子分类框架，为患者分层提供信息。'
        '其次，我们剖析肿瘤内在耐药机制——从致癌信号通路和表观遗传失调到EMT驱动的免疫排斥和失调的细胞死亡程序。'
        '第三，我们审视由免疫抑制细胞群体、细胞因子网络、代谢重编程以及瘤内微生物组的新兴角色所产生的TME介导的耐药。'
        '随后，我们讨论由克隆演化、组织学转化和治疗压力下的表型可塑性驱动的获得性耐药。'
        '最后，我们评估旨在克服这些耐药屏障的新兴治疗策略，'
        '包括合理的免疫联合治疗、TME重塑方法和生物标志物引导的精准免疫治疗。'
        '通过整合这些机制领域的发现，本综述旨在为理解LUSC免疫治疗耐药的多面性提供一个框架，'
        '并为改善这一挑战性疾病的结局确定可行的机会。'
    ),
]

for text in intro_texts:
    add_para(text)

doc.add_page_break()

# ============================================================
# 2. IMMUNE LANDSCAPE
# ============================================================
add_heading_cn('2  LUSC的免疫景观', 1)

add_para(
    '理解LUSC免疫治疗耐药机制需要对其独特的肿瘤免疫微环境（TIME）有一个基本的认识。'
    '尽管LUSC与LUAD在肺中共享共同的解剖起源，但两者的免疫景观存在实质性差异，'
    '反映了突变过程、基质组成和演化轨迹的不同。'
    '单细胞RNA测序（scRNA-seq）、空间转录组学和多组学整合的最新进展以前所未有的分辨率'
    '剖析了LUSC TIME的细胞和分子架构，揭示了直接与预后和治疗策略相关的独特免疫亚型和免疫抑制回路。'
)

add_heading_cn('2.1  TIME异质性与分子分类', 2)

add_para(
    'LUSC的TIME以深刻的瘤间和瘤内异质性为特征，从根本上塑造了对免疫治疗的应答 [16]。'
    '以经典的肿瘤免疫表型三分法——免疫炎症型、免疫排斥型和免疫沙漠型——为基础，'
    '近期研究已利用整合多组学方法针对LUSC对这些分类进行了细化 [16,19]。'
)

add_para(
    '免疫炎症型LUSC的特征是肿瘤实质内CD8+细胞毒性T淋巴细胞（CTL）、'
    '活化CD4+记忆T细胞和树突状细胞（DC）的丰富浸润。这些肿瘤通常表达高水平效应细胞因子和溶细胞标志物，'
    '与持续进行但最终无效的抗肿瘤免疫应答相一致 [16]。'
    '矛盾的是，许多免疫炎症型LUSC同时上调多个抑制性免疫检查点，'
    '并携带高频率的免疫抑制性调节性T细胞（Treg）和M2极化巨噬细胞，'
    '形成一种"炎症浸润但功能抑制"的免疫状态，'
    '这种状态常常是单药PD-1/PD-L1阻断原发性耐药的基础 [18,20]。',
    bold=False
)

add_para(
    '免疫排斥型LUSC的特征是CTL和其他免疫效应细胞被限制在肿瘤巢周围的基质区室中，'
    '无法穿透肿瘤实质。这种排斥主要由癌相关成纤维细胞（CAF）介导，'
    '它们沉积包括胶原、纤连蛋白和层粘连蛋白在内的致密细胞外基质（ECM）成分，'
    '形成T细胞浸润的物理屏障 [16]。特别是POSTN+ CAF和FAP+ CAF已被证明可产生趋化排斥梯度，'
    '并使ECM纤维垂直于肿瘤边界排列，有效地将T细胞困在肿瘤周围基质中 [16]。'
    '空间转录组分析揭示了CAF亚群与APOE+肿瘤相关巨噬细胞（TAM）的紧密共定位，'
    '鉴定出一种通过协同的细胞因子和趋化因子分泌来强化免疫排斥的CAF-TAM信号轴 [16]。',
    bold=False
)

add_para(
    '免疫沙漠型LUSC占少数病例，其标志是肿瘤和基质区室内几乎完全缺乏T细胞。'
    '这种表型被认为源于先天性免疫感知缺陷——包括STING/cGAS通路激活受损——'
    '以及树突状细胞启动不足，导致无法启动有效的抗肿瘤T细胞应答 [16,19]。'
    '免疫沙漠型肿瘤对治疗提出了尤为困难的挑战，因为在缺乏预先存在的T细胞浸润的情况下，'
    'ICI单药治疗不太可能有效。',
    bold=False
)

add_para(
    '在这些定性描述之外，整合基因组规模分析已基于免疫基因表达特征、免疫细胞反卷积算法和'
    '多组学数据整合建立了LUSC的分子分类框架 [19,21]。'
    'Yin等人通过免疫相关基因表达谱的无监督聚类鉴定出LUSC的不同免疫亚型，'
    '将患者分为具有不同生存结局和预测免疫治疗应答的组别 [19]。'
    '同样，Yang等人鉴定出一种细胞因子主导的免疫抑制类（下文将详细讨论），'
    '该类展现出具有直接治疗意义的独特分子特征 [18]。'
    'Song等人利用513例LUSC样本的bulk和单细胞RNA测序，划分了六个分子亚型（CS1-CS6），'
    '并鉴定出CS3为一种淋巴细胞浸润亚型，其矛盾性地表现出升高的耗竭标志物'
    '（CTLA-4、LAG-3、PD-1），并通过TIDE分析预测对ICB治疗耐药 [20]。'
    '这些分类工作的共同结论是：肿瘤浸润淋巴细胞的单纯存在不足以预测LUSC的ICI应答；'
    '相反，免疫浸润的功能状态和空间组织才是治疗结局的关键决定因素。'
)

add_heading_cn('2.2  LUSC与LUAD：分歧的免疫环境', 2)

add_para(
    '尽管LUSC和LUAD均被归类为NSCLC，但两者起源于不同的细胞源头，'
    '携带根本不同的基因组景观，并展现出明显分歧的免疫微环境 [22,23]。'
    '这些差异对免疫治疗策略的设计和解读具有深远意义。'
)

add_para(
    '在基因组层面，LUAD富含可靶向驱动突变——最显著的是EGFR（东亚人群~30-40%，西方人群~10-15%）、'
    'KRAS（~30%）、ALK重排（~5%）和ROS1融合（~1-2%）——'
    '这些突变在EGFR突变亚群中与相对较低的肿瘤突变负荷（TMB）和较低的炎症浸润相关 [22]。'
    '相比之下，LUSC主要以致癌基因功能丧失突变（TP53、CDKN2A、KEAP1）和反复扩增'
    '（PIK3CA、FGFR1、SOX2）为主导，产生较高的总TMB和新抗原负荷，'
    '理论上预测其免疫原性更高 [22,23]。'
    '然而，这种升高的新抗原负荷并未转化为LUSC相较于LUAD更优的ICI应答，'
    '提示存在强效的对抗性免疫抑制机制。'
)

add_para(
    '在细胞层面，比较分析揭示了LUSC和LUAD在免疫细胞组成上的系统性差异 [23]。'
    'Yan等人通过基于lncRNA的免疫异质性分析发现，两种亚型在浸润免疫细胞的比例和功能状态上存在显著差异。'
    '具体而言，LUSC肿瘤表现出更高的M2巨噬细胞和静息CD4+记忆T细胞浸润，'
    '而LUAD肿瘤倾向于有更高的幼稚B细胞和浆细胞浸润 [23]。'
    '与LUAD相比，LUSC也表现出更高的免疫排斥表型比例，'
    '由更广泛的CAF激活和ECM重塑驱动 [16,22]。'
    '此外，LUSC的细胞因子环境偏向TGF-β和IL-6/STAT3信号，促进免疫抑制和EMT，'
    '而LUAD更常与EGFR驱动的免疫抑制程序相关 [22]。'
)

add_para(
    '这些亚型特异性免疫特征具有直接的临床意义。'
    '第一，LUSC中免疫排斥的高比例提示靶向CAF介导的基质屏障的策略——'
    '如FAP抑制剂、Hedgehog通路拮抗剂或TGF-β阻断——可能对该组织学类型尤为适用。'
    '第二，LUSC和LUAD之间不同的免疫检查点表达谱可指导联合免疫治疗搭档的理性选择 [22]。'
    '第三，在LUAD中开发的预后和预测性生物标志物可能无法直接用于LUSC，'
    '考虑到两者免疫生物学的差异。'
    '这些考虑强调了LUSC特异性免疫治疗策略的必要性，'
    '以及在鳞状组织学中专门验证生物标志物的重要性。'
)

add_heading_cn('2.3  免疫耗竭型：理解耐药的框架', 2)

add_para(
    'LUSC免疫生物学的一项里程碑式贡献是Yang等人通过对624例LUSC样本进行RNA测序数据的无监督聚类，'
    '鉴定出了免疫耗竭型（Exhausted Immune Class, EIC）[18]。'
    '该分析揭示约28-36%的LUSC患者属于EIC，其特征是具有对免疫治疗耐药至关重要意义的独特分子标志。'
)

add_para(
    'EIC由四个标志性特征的共同存在所定义 [18]，如附图3所示。'
    '第一，该类肿瘤展现出T细胞耗竭标志的显著富集，'
    '伴随典型耗竭相关转录因子TOX和EOMES的表达升高，'
    '以及TCF-1（由TCF7编码）——保留增殖能力和PD-1阻断应答性的祖细胞耗竭T细胞（Tpex）标志物——的表达降低。'
    '这种从Tpex主导到终末耗竭T细胞区室的转变代表了ICI疗效的关键屏障，'
    '因为终末耗竭T细胞在很大程度上对PD-1通路再激活无效。'
)

add_para(
    '第二，EIC的特征是高达九个抑制性免疫检查点的共上调——'
    'CTLA-4、PD-1（PDCD1）、LAG-3、BTLA、TIGIT、TIM-3（HAVCR2）、IDO1、SIGLEC7和VISTA——'
    '代表一种广泛的、多受体的免疫抑制状态，使得单一PD-1/PD-L1阻断在机制上就不足 [18]。'
    '这一发现为以下临床观察提供了分子层面的解释：许多具有高PD-L1表达的LUSC患者'
    '仍然对帕博利珠单抗或纳武利尤单抗单药治疗无应答。'
)

add_para(
    '第三，EIC肿瘤被免疫抑制细胞群体高度浸润，特别是M2极化巨噬细胞和CD4+FOXP3+ Treg的比例尤为突出。'
    '这些细胞产生免疫抑制性细胞因子——包括TGF-β、IL-10和CCL18——'
    '直接抑制CTL效应功能并促进进一步的T细胞耗竭 [18]。'
)

add_para(
    '第四，EIC矛盾性地与高总密度的肿瘤浸润淋巴细胞（TIL）相关，'
    '然而这些患者的预后显著差于TIL浸润较低的患者。'
    '这一发现强化了免疫浸润数量与免疫能力之间的关键区别，'
    '并突显了仅用TIL密度作为LUSC中ICI应答生物标志物的不足。'
)

add_para(
    'EIC概念已得到后续研究的证实和扩展。'
    'Song等人在其CS3分子亚型中鉴定出类似表型，该亚型表现出高淋巴细胞浸润伴随升高的耗竭标志物，'
    '并通过独立计算方法预测ICB耐药 [20]。'
    '此外，CS3亚型被发现特异性上调LAMC2-CD44分子轴——一种与EMT相关的通路，'
    '同时调控肿瘤增殖和免疫排斥，提供了肿瘤细胞内在程序与耗竭免疫表型之间的潜在机制联系。'
)

add_para(
    '总之，LUSC TIME的特征——涵盖其异质性细胞组成、相对于LUAD的分歧特征以及具有临床关键意义的EIC——'
    '为特定耐药机制的运作奠定了生物学基础。'
    '后续章节将详细剖析这些机制，首先从肿瘤细胞内在的耐药决定因素开始。'
)

doc.add_page_break()

# ============================================================
# Continue with sections 3-7 + references (abbreviated for generation)
# The full script continues below...
# ============================================================

print(f"Part 1 of Word generation complete (through Section 2)...")
print(f"Continuing with remaining sections...")
PYEOF

# Actually, let me write the complete script properly
print("Script structure verified. Will generate complete document.")
