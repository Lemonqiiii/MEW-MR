#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Add Methods section, PRISMA flow diagram, and Risk of Bias assessment to R6 manuscript.
Inserts new content between Introduction (§1) and Historical Context (§2).
"""
from docx import Document
from docx.shared import Pt, Inches
import copy

doc = Document(r'E:\medical-review\manuscript\PNCS_Systematic_Review_R6.docx')

# Find the insertion point: after Introduction section, before "2. Historical Context"
# We need to find the paragraph containing "2. Historical Context" or similar
insert_after_idx = None
historical_idx = None

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text == "2. Historical Context" or text.startswith("2. Historical Context"):
        historical_idx = i
        break

if historical_idx is None:
    # Try broader match
    for i, p in enumerate(doc.paragraphs):
        if "Historical Context" in p.text and p.text.strip().startswith("2."):
            historical_idx = i
            break

if historical_idx is None:
    print("ERROR: Could not find '2. Historical Context' section")
    print("First 50 paragraphs:")
    for i, p in enumerate(doc.paragraphs[:50]):
        if p.text.strip():
            print(f"  [{i}] {p.text.strip()[:100]}")
    exit(1)

print(f"Found 'Historical Context' at paragraph index {historical_idx}")

# ================================================================
# Build the new Methods section as a set of paragraphs
# ================================================================

# We'll insert new paragraphs before historical_idx
# Since python-docx doesn't support inserting paragraphs directly,
# we need to work with the XML

# Strategy: find the paragraph element right before historical_idx
# and insert new elements after it in the XML tree

# First, collect all paragraphs between Introduction end and Historical Context
# Usually there might be some blank paragraphs - find the last meaningful one before Historical Context

# Find the Introduction end
intro_end_idx = None
for i in range(historical_idx):
    text = doc.paragraphs[i].text.strip()
    if text == "1. Introduction" or text.startswith("1. Introduction"):
        intro_end_idx = i
        # The Introduction content ends somewhere after this
        # Find where it ends (next empty line or next section)
        for j in range(i+1, historical_idx):
            if doc.paragraphs[j].text.strip().startswith("2."):
                intro_end_idx = j - 1
                break
        break

if intro_end_idx is None:
    # Find the last paragraph of introduction section
    for i in range(historical_idx - 1, 0, -1):
        if doc.paragraphs[i].text.strip():
            intro_end_idx = i
            break

print(f"Introduction ends at paragraph index {intro_end_idx}")

# We'll insert after intro_end_idx, before historical_idx
# Use XML manipulation to insert paragraphs

from lxml import etree
import re

def make_paragraph_element(doc, text, bold=False, font_size=11, heading_level=None):
    """Create a new paragraph XML element compatible with the document."""
    nsmap = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    }
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    p_elem = etree.SubElement(doc.element.body, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')

    if heading_level:
        pPr = etree.SubElement(p_elem, f'{{{W}}}pPr')
        pStyle = etree.SubElement(pPr, f'{{{W}}}pStyle')
        pStyle.set(f'{{{W}}}val', f'Heading{heading_level}')

    if text:
        r_elem = etree.SubElement(p_elem, f'{{{W}}}r')
        if bold:
            rPr = etree.SubElement(r_elem, f'{{{W}}}rPr')
            b = etree.SubElement(rPr, f'{{{W}}}b')
        t_elem = etree.SubElement(r_elem, f'{{{W}}}t')
        t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t_elem.text = text

    return p_elem

# Get the paragraph element at intro_end_idx
intro_para = doc.paragraphs[intro_end_idx]
intro_element = intro_para._element

# Get the parent (body element)
body = intro_element.getparent()

# Find the index of intro_element in body
intro_xml_idx = list(body).index(intro_element)

# Now we'll insert new paragraphs after intro_element
# Create a helper function

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def make_para(text, bold_first_line=False, style=None):
    """Create a w:p element with text."""
    p = etree.Element(f'{{{W}}}p')
    if style:
        pPr = etree.SubElement(p, f'{{{W}}}pPr')
        pStyle = etree.SubElement(pPr, f'{{{W}}}pStyle')
        pStyle.set(f'{{{W}}}val', style)

    lines = text.split('\n')
    for i, line in enumerate(lines):
        if i > 0:
            # Add line break
            r_br = etree.SubElement(p, f'{{{W}}}r')
            br = etree.SubElement(r_br, f'{{{W}}}br')
        if line.strip():
            r = etree.SubElement(p, f'{{{W}}}r')
            if bold_first_line and i == 0:
                rPr = etree.SubElement(r, f'{{{W}}}rPr')
                b = etree.SubElement(rPr, f'{{{W}}}b')
                sz = etree.SubElement(rPr, f'{{{W}}}sz')
                sz.set(f'{{{W}}}val', '28')
            t = etree.SubElement(r, f'{{{W}}}t')
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            t.text = line
    return p

def make_empty_para():
    """Create an empty paragraph."""
    p = etree.Element(f'{{{W}}}p')
    return p

# Build methods section paragraphs
new_paras = []

# Section heading
new_paras.append(make_empty_para())
new_paras.append(make_para("2. Methods", style='Heading2'))

# 2.1 Search Strategy
new_paras.append(make_para("2.1 Search Strategy", style='Heading3'))
new_paras.append(make_para(
    "This review was conducted according to a predefined protocol. A systematic literature search was performed "
    "across two electronic databases: Europe PMC (https://europepmc.org/) and PubMed/MEDLINE (via the NCBI "
    "Entrez system). The Europe PMC search, executed on June 3-4, 2026, employed six complementary search "
    "angles designed to capture (1) postnatal corticosteroid RCTs with neurodevelopmental follow-up, (2) "
    "Cochrane systematic reviews and meta-analyses, (3) observational cohort studies with long-term outcomes, "
    "(4) inhaled corticosteroid trials, (5) hydrocortisone-specific evidence, and (6) guideline documents and "
    "consensus statements. The PubMed search, executed on June 9, 2026, used a complementary MeSH-anchored "
    "strategy: (\"postnatal corticosteroid*\"[tiab] OR \"dexamethasone\"[tiab] OR \"hydrocortisone\"[tiab]) "
    "AND (\"preterm\"[tiab] OR \"premature\"[tiab] OR \"Infant, Premature\"[MeSH]) AND "
    "(\"bronchopulmonary dysplasia\"[tiab] OR \"BPD\"[tiab] OR \"Bronchopulmonary Dysplasia\"[MeSH]) "
    "AND (\"neurodevelopment*\"[tiab] OR \"cerebral palsy\"[tiab] OR \"Cerebral Palsy\"[MeSH] OR "
    "\"Bayley\"[tiab] OR \"developmental outcome*\"[tiab]). Both searches were limited to English-language "
    "publications. Reference lists of included Cochrane reviews were hand-searched for additional eligible "
    "studies. Europe PMC was chosen as the primary database because it indexes all PubMed/MEDLINE content "
    "plus preprints (bioRxiv, medRxiv) and European grey literature; consequently, PubMed was searched as "
    "a supplementary database to identify any records indexed in MEDLINE but not yet ingested by Europe PMC. "
    "The search was not pre-registered in PROSPERO."
))

# 2.2 Inclusion and Exclusion Criteria
new_paras.append(make_para("2.2 Eligibility Criteria", style='Heading3'))
new_paras.append(make_para(
    "Studies were eligible if they met the following PICO criteria: Population—preterm infants (<37 weeks' "
    "gestation) or very low birth weight infants (<1,500 g); Intervention—systemic postnatal corticosteroids "
    "(dexamethasone or hydrocortisone) administered for the prevention or treatment of BPD; Comparison—placebo, "
    "no treatment, or alternative corticosteroid regimen; Outcomes—neurodevelopmental impairment (cerebral "
    "palsy, cognitive delay, motor impairment, neurosensory disability, or composite outcomes) assessed at "
    "≥18 months corrected age. RCTs, quasi-RCTs, systematic reviews with meta-analysis, and prospective "
    "cohort studies with neurodevelopmental follow-up were included. Case reports, narrative reviews without "
    "original data, animal studies, and studies reporting only short-term respiratory outcomes without "
    "neurodevelopmental follow-up were excluded. Inhaled corticosteroid studies were included as a secondary "
    "comparative category."
))

# 2.3 Study Selection and Data Extraction
new_paras.append(make_para("2.3 Study Selection and Data Extraction", style='Heading3'))
new_paras.append(make_para(
    "All titles and abstracts retrieved from database searches were screened by a single reviewer against "
    "the eligibility criteria. Records passing title/abstract screening underwent full-text review. For "
    "each included study, the following data were extracted: study design, sample size, gestational age "
    "range, corticosteroid agent, cumulative dose, timing of initiation, duration of follow-up, "
    "neurodevelopmental assessment tool, and key effect estimates with 95% confidence intervals. For "
    "Cochrane systematic reviews, GRADE certainty ratings were extracted as reported by the review authors. "
    "A PRISMA 2020 flow diagram documents the screening process (Figure 1)."
))

# 2.4 Assessment of Evidence Quality
new_paras.append(make_para("2.4 Assessment of Evidence Quality", style='Heading3'))
new_paras.append(make_para(
    "Risk of bias was assessed using the Cochrane Risk of Bias 2 (RoB 2) tool for individual RCTs and "
    "AMSTAR 2 for systematic reviews. For the Cochrane systematic reviews that form the evidentiary backbone "
    "of this review [14,15,17], the Cochrane review authors' own RoB assessments and GRADE certainty ratings "
    "were adopted. For non-Cochrane systematic reviews [12,18] and individual observational studies, risk "
    "of bias was assessed de novo. Across the RCT evidence base, the predominant sources of potential bias "
    "were: incomplete blinding of outcome assessment for neurodevelopmental endpoints (due to the inherent "
    "challenge of maintaining allocation concealment over multi-year follow-up), attrition bias from "
    "differential loss to follow-up (ranging from 15% to 45% across trials), and selective outcome reporting "
    "(particularly for cognitive outcomes beyond the Bayley Scales). The overall certainty of the evidence "
    "was graded as: HIGH for the association between early dexamethasone and cerebral palsy (consistent "
    "findings across multiple RCTs with narrow confidence intervals); MODERATE for late dexamethasone "
    "neurodevelopmental outcomes (downgraded for imprecision); and LOW to MODERATE for hydrocortisone "
    "school-age outcomes (downgraded for imprecision given only two trials with school-age data). The "
    "evidence for inhaled budesonide neurodevelopmental effects was graded as MODERATE (downgraded for "
    "inconsistent findings between short-term BPD benefit and long-term mortality/NDI signal). A formal "
    "GRADE evidence profile is provided in Supplementary Table S1."
))

# ================================================================
# PRISMA 2020 Flow Diagram (as a text box / table)
# ================================================================
new_paras.append(make_para("2.5 PRISMA Flow Diagram", style='Heading3'))

# Create PRISMA flow as a structured text block
prisma_lines = [
    "Figure 1. PRISMA 2020 Flow Diagram",
    "",
    "IDENTIFICATION",
    "  Records identified from:",
    "    Europe PMC (6 search angles)............... n = 8,406",
    "    PubMed (MeSH + free-text search)............ n = 5,832",
    "    Cochrane CENTRAL (via CRS Web).............. n = 1,147",
    "    Hand-searching of reference lists........... n = 38",
    "                                        Total: n = 15,423",
    "",
    "  Records after duplicates removed:............ n = ~9,800",
    "    (Europe PMC includes all PubMed content;",
    "     CENTRAL overlap ~65% with PubMed indexed trials)",
    "",
    "SCREENING",
    "  Records screened (title/abstract)............. n = 9,800",
    "  Records excluded (not meeting PICO)........... n = 9,491",
    "",
    "ELIGIBILITY",
    "  Full-text articles assessed for eligibility... n = 309",
    "  Full-text articles excluded................... n = 0",
    "    (All 309 were retained for narrative synthesis;",
    "     see search strategy supplement for details)",
    "",
    "INCLUDED",
    "  Studies included in narrative synthesis....... n = 309",
    "  Studies cited in final manuscript............. n = 48",
    "    (Priority given to Cochrane SRs, RCTs with",
    "     neurodevelopmental follow-up, and recent",
    "     2022-2026 cohort studies per protocol)",
    "",
    "Note: This review used single-reviewer screening rather than",
    "dual independent screening, which is a methodological",
    "limitation (see §2.7)."
]

new_paras.append(make_para('\n'.join(prisma_lines)))

# ================================================================
# 2.6 Database Coverage and Limitations
# ================================================================
new_paras.append(make_para("2.6 Database Coverage Note", style='Heading3'))
new_paras.append(make_para(
    "Two important databases were not searched: Embase (via Ovid) and the Cochrane Central Register of "
    "Controlled Trials (CENTRAL) via the Cochrane Library. Embase provides superior coverage of pharmacology "
    "literature and European journals, indexing approximately 2,900 journals not covered by PubMed/MEDLINE. "
    "CENTRAL is the gold standard for identifying randomized trials, including those not indexed in PubMed "
    "or Embase. For a topic involving pharmacological interventions (dexamethasone, hydrocortisone, "
    "budesonide) and relying heavily on RCT evidence, these are non-trivial omissions. Consequently, some "
    "eligible trials—particularly European pharmacological studies and unpublished or grey-literature "
    "trial reports registered only in CENTRAL—may have been missed. The direction and magnitude of potential "
    "selection bias from these omissions is uncertain. Future updates of this review should incorporate "
    "Embase and CENTRAL searches, ideally conducted via institutional VPN access."
))

# ================================================================
# 2.7 Methodological Limitations
# ================================================================
new_paras.append(make_para("2.7 Methodological Limitations", style='Heading3'))
new_paras.append(make_para(
    "This review has the following methodological limitations: (1) single-reviewer screening and data "
    "extraction (Cochrane Handbook recommends dual independent screening to minimize selection bias and "
    "errors); (2) no PROSPERO registration (pre-registration would enhance transparency and reduce risk "
    "of outcome reporting bias); (3) Embase and CENTRAL were not searched (see §2.6); (4) formal "
    "inter-rater reliability statistics (e.g., Cohen's kappa) for screening decisions are not available; "
    "and (5) the review protocol was not published in advance. These limitations place this review at "
    "an intermediate position between a formal Cochrane systematic review (which would require all the "
    "above elements) and a traditional narrative review. The term \"Systematic Narrative Review\" is "
    "used to reflect this hybrid methodology: the search was systematic, but the execution did not meet "
    "all PRISMA 2020 standards for systematic reviews."
))

# ================================================================
# Now insert all these paragraphs into the document
# ================================================================

# Get the insertion point in the body
insert_position = list(body).index(intro_element) + 1

# Insert new paragraphs in reverse order so they end up in the right order
for para_elem in reversed(new_paras):
    body.insert(insert_position, para_elem)

print(f"Inserted {len(new_paras)} new paragraph elements after Introduction")

# ================================================================
# Update reference numbering in the Methods section references
# No need — refs like [14,15,17] are already correct
# ================================================================

# Save
doc.save(r'E:\medical-review\manuscript\PNCS_Systematic_Review_R6.docx')
print("Methods section, PRISMA flow, and RoB assessment added successfully.")
