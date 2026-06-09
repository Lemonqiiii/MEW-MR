#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Apply additional reference fixes discovered during extended verification."""
from docx import Document

doc = Document(r'E:\medical-review\manuscript\PNCS_Systematic_Review_R6.docx')

fixes_applied = []

# ============ FIX 1: Ref [24] — completely wrong author, journal, and interpretation ============
for p in doc.paragraphs:
    for run in p.runs:
        if "O'Brien CE, et al. Postnatal dexamethasone treatment for preterm infants" in run.text:
            run.text = run.text.replace(
                "O'Brien CE, et al. Postnatal dexamethasone treatment for preterm infants at high risk for bronchopulmonary dysplasia is associated with altered regional brain volumes at term-equivalent age. *Pediatr Res* 2025. PMID: 40360237.",
                "Chandwani R, Kline J, Altaye M, Parikh N. Postnatal dexamethasone treatment for preterm infants at high risk for bronchopulmonary dysplasia is associated with improved regional brain volumes: a prospective cohort study. *Arch Dis Child Fetal Neonatal Ed* 2025;111:F74-F81. PMID: 40360237."
            )
            fixes_applied.append("Ref [24] author/journal/volume")

# Fix paragraph text that misinterprets Chandwani paper
for p in doc.paragraphs:
    full = p.text
    if "A 2025 cohort study showed that low-dose postnatal dexamethasone was associated with altered" in full:
        texts = [r.text for r in p.runs]
        combined = "".join(texts)
        combined = combined.replace(
            "A 2025 cohort study showed that low-dose postnatal dexamethasone was associated with altered regional brain volumes on term-equivalent MRI [24]",
            "A 2025 prospective cohort study found that low-dose postnatal dexamethasone (0.89 mg/kg cumulative, initiated at median day 36) was associated with larger cerebellar and subcortical grey matter volumes on term-equivalent MRI, and higher motor scores at 2 years [24]"
        )
        if p.runs:
            p.runs[0].text = combined
            for r in p.runs[1:]:
                r.text = ""
            fixes_applied.append("Section 4.2: Chandwani interpretation corrected")

    if "suggesting that even" in full and "low-dose" in full:
        texts = [r.text for r in p.runs]
        combined = "".join(texts)
        old_text = 'suggesting that even "low-dose" dexamethasone may have measurable structural brain effects, though the functional significance of these volume differences is unclear'
        new_text = "suggesting that low-dose dexamethasone initiated after the first postnatal week does not have adverse macrostructural brain effects and may have a protective effect on motor development"
        if old_text in combined:
            combined = combined.replace(old_text, new_text)
            if p.runs:
                p.runs[0].text = combined
                for r in p.runs[1:]:
                    r.text = ""
                fixes_applied.append("Section 4.2: conclusion sentence corrected")

# ============ FIX 2: Ref [8] — Shinwell year/volume/pages ============
for p in doc.paragraphs:
    for run in p.runs:
        if "Shinwell ES, Karplus M, Reich D" in run.text and "1997;76:F283" in run.text:
            run.text = run.text.replace(
                "Shinwell ES, Karplus M, Reich D, et al. Early postnatal dexamethasone therapy and increased incidence of cerebral palsy. *Arch Dis Child Fetal Neonatal Ed* 1997;76:F283-7.",
                "Shinwell ES, Karplus M, Reich D, et al. Early postnatal dexamethasone treatment and increased incidence of cerebral palsy. *Arch Dis Child Fetal Neonatal Ed* 2000;83:F177-F181."
            )
            fixes_applied.append("Ref [8] year/volume/pages 1997->2000")

# ============ FIX 3: Ref [36] — JAMA -> JAMA Pediatr ============
for p in doc.paragraphs:
    for run in p.runs:
        if "DeMauro SB, Kirpalani H, Hintz S" in run.text and "JAMA* 2026" in run.text:
            run.text = run.text.replace(
                "DeMauro SB, Kirpalani H, Hintz S, et al. Hydrocortisone in preterm infants and school-age functional outcomes: follow-up of a randomized clinical trial. *JAMA* 2026. PMID: 41359352.",
                "DeMauro SB, Kirpalani H, Hintz S, Watterberg KL, et al. Hydrocortisone in preterm infants and school-age functional outcomes: follow-up of a randomized clinical trial. *JAMA Pediatr* 2026;180(2):134-143. PMID: 41359352."
            )
            fixes_applied.append("Ref [36] JAMA->JAMA Pediatr + volume + Watterberg")

# ============ FIX 4: Ref [38] NEUROSIS — add PMID ============
for p in doc.paragraphs:
    for run in p.runs:
        if "Bassler D, Plavka R, Shinwell ES" in run.text and "1497-506" in run.text and "PMID" not in run.text:
            run.text = run.text.replace(
                "Bassler D, Plavka R, Shinwell ES, et al. Early inhaled budesonide for the prevention of bronchopulmonary dysplasia. *N Engl J Med* 2015;373:1497-506.",
                "Bassler D, Plavka R, Shinwell ES, et al. Early inhaled budesonide for the prevention of bronchopulmonary dysplasia. *N Engl J Med* 2015;373:1497-506. PMID: 26469126."
            )
            fixes_applied.append("Ref [38] added PMID 26469126")

# ============ FIX 5: update Section 4.2 heading context - Beyond Cerebral Palsy ============
# The section narrative needs to flow correctly with the corrected Chandwani interpretation
for p in doc.paragraphs:
    full = p.text
    if "A 2014 meta-analysis of dexamethasone RCTs" in full and "2025 cohort study" in full:
        texts = [r.text for r in p.runs]
        combined = "".join(texts)
        old_transition = "A 2014 meta-analysis of dexamethasone RCTs [23] found adverse effects on intelligence and hearing."
        new_transition = "A 2014 meta-analysis of dexamethasone RCTs [23] found adverse effects on intelligence and hearing with early treatment, and increased hearing loss with late treatment."
        if old_transition in combined:
            combined = combined.replace(old_transition, new_transition)
            if p.runs:
                p.runs[0].text = combined
                for r in p.runs[1:]:
                    r.text = ""
                fixes_applied.append("Section 4.2: 2014 meta-analysis context refined")

doc.save(r'E:\medical-review\manuscript\PNCS_Systematic_Review_R6.docx')

for f in fixes_applied:
    print(f"  ✅ {f}")
print(f"\n{len(fixes_applied)} additional fixes applied and saved.")
