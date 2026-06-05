#!/usr/bin/env python3
"""Generate complete Chinese Word document from JITC manuscript."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUT = "E:/medical-review/manuscript/中文稿件_NSCLC鳞癌免疫治疗耐药机制.docx"
doc = Document()

# Page setup
s = doc.sections[0]
s.top_margin = Cm(2.54); s.bottom_margin = Cm(2.54)
s.left_margin = Cm(3.0); s.right_margin = Cm(3.0)

# Styles
sty = doc.styles['Normal']; sty.font.name = '宋体'; sty.font.size = Pt(12)
sty.paragraph_format.line_spacing = 1.5
for lv in [1,2,3]:
    hs = doc.styles[f'Heading {lv}']; hs.font.name = '黑体'
    hs.font.color.rgb = RGBColor(0,0,0)
    if lv==1: hs.font.size=Pt(16)
    elif lv==2: hs.font.size=Pt(14)
    else: hs.font.size=Pt(12)

def p(text, bold=False, sz=12, font='宋体', align=None, sp=1.5):
    par = doc.add_paragraph(); par.paragraph_format.line_spacing = sp
    if align is not None: par.alignment = align
    r = par.add_run(text); r.font.name = font; r.font.size = Pt(sz); r.bold = bold
    return par

def h(text, lv=1):
    hd = doc.add_heading(text, level=lv)
    for r in hd.runs: r.font.name = '黑体'; r.font.color.rgb = RGBColor(0,0,0)
    return hd

def page_break():
    doc.add_page_break()

# ============= TITLE PAGE =============
p(''); p('')
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run('非小细胞肺癌鳞状细胞癌\n免疫治疗耐药机制研究进展'); r.font.name='黑体'; r.font.size=Pt(22); r.bold=True
p('')
sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sp.add_run('Mechanisms of Immunotherapy Resistance in\nSquamous Cell Carcinoma of Non-Small Cell Lung Cancer')
sr.font.name='Times New Roman'; sr.font.size=Pt(14); sr.italic=True
p(''); p('')
for lb, vl in [('文章类型：','综述'),('总字数：','约9,000字（中文正文）'),('图表：','图4幅 | 表4张 | 参考文献41篇'),
               ('关键词：','肺鳞状细胞癌；免疫治疗耐药；免疫检查点抑制剂；肿瘤微环境；T细胞耗竭；KEAP1/NRF2')]:
    pp = p(''); rl = pp.add_run(lb); rl.font.name='黑体'; rl.font.size=Pt(12); rl.bold=True
    rv = pp.add_run(vl); rv.font.name='宋体'; rv.font.size=Pt(12)

page_break()

# ============= ABSTRACT =============
h('摘要', 1)
p('肺鳞状细胞癌（lung squamous cell carcinoma, LUSC）约占非小细胞肺癌的30%，其显著特征为缺乏可靶向驱动突变，因此免疫治疗成为全身治疗的基石。靶向PD-1/PD-L1和CTLA-4轴的免疫检查点抑制剂（ICI）已根本性地改变了LUSC的治疗格局；然而，仅少数患者获得持久获益，原发性耐药和获得性耐药构成了关键的临床挑战。LUSC的耐药产生于肿瘤内在改变、肿瘤微环境（TME）介导的免疫抑制以及治疗诱导的适应性变化三者之间的复杂交互作用。本综述从三个相互关联的维度系统整合了LUSC免疫治疗耐药机制的最新认识：首先概述LUSC独特的免疫景观，包括存在于28-36%患者中的免疫耗竭型（Exhausted Immune Class, EIC）；其次审视肿瘤内在耐药机制（PI3K/AKT、KEAP1/NRF2、p38 MAPK、上皮间质转化、表观遗传失调、程序性细胞死亡缺陷）；第三分析TME介导的耐药（T细胞耗竭伴多重检查点共表达、TAM/CAF/MDSC/Treg等免疫抑制细胞、代谢限制、瘤内微生物组）。进一步讨论获得性耐药（克隆演化、组织学转化、表型可塑性），最后评估克服策略（联合免疫治疗、靶向内在通路、TME重塑、精准免疫治疗）。')
p('')
p('Abstract', bold=True, sz=11, font='Times New Roman')
p("Lung squamous cell carcinoma (LUSC), accounting for approximately 30% of non-small cell lung cancers, is characterized by a paucity of actionable driver mutations and a consequent reliance on immunotherapy. Immune checkpoint inhibitors targeting the PD-1/PD-L1 and CTLA-4 axes have transformed treatment; however, only a minority achieve durable benefit. This review synthesizes resistance mechanisms across three dimensions: tumor-intrinsic alterations, TME-mediated suppression, and acquired resistance, concluding with strategies to overcome these barriers.", sz=10, font='Times New Roman')

page_break()

# ============= 1. INTRODUCTION =============
h('1  引言', 1)

p('肺癌仍是全球癌症相关死亡的首要原因，非小细胞肺癌（NSCLC）约占所有病例的85% [1]。在NSCLC亚型中，肺鳞状细胞癌（LUSC）约占诊断的30%，是仅次于肺腺癌（LUAD）的第二常见组织学亚型 [2]。LUSC患者面临尤为严峻的临床结局，晚期疾病五年生存率低于20% [3]。历史上，标准治疗主要限于铂类化疗，生存获益有限且毒性显著 [4]。')

p('LUSC治疗格局最深层的塑造因素恰恰是一种缺失：可靶向驱动突变的极度匮乏。在LUAD中，针对EGFR、ALK、ROS1、BRAF和KRAS G12C的靶向治疗已根本改变了治疗范式；相比之下，LUSC几乎没有可通过药物靶向的基因组改变 [5,6]。LUSC的基因组景观主要由抑癌基因频繁改变所主导——包括TP53（~80%）、CDKN2A（~70%，涵盖突变、缺失和表观遗传沉默）和KEAP1（~12%）——以及PIK3CA（~30-40%）、FGFR1（~20%）的反复扩增，以及包含鳞状谱系转录因子SOX2和TP63的染色体3q位点 [5,7]。尽管针对这些基因组特征开展了大量靶向药物临床试验，迄今尚无任何药物获得专门针对LUSC的注册批准，凸显了这一患者群体对有效治疗策略的迫切未满足需求 [7]。')

p('免疫检查点抑制剂（ICIs）的问世已根本性地改变了LUSC的治疗格局。靶向程序性死亡受体-1（PD-1；纳武利尤单抗、帕博利珠单抗）、其配体PD-L1（阿替利珠单抗、度伐利尤单抗）以及CTLA-4（伊匹木单抗）的抗体在晚期NSCLC的多线治疗中均展现出有意义的临床活性，LUSC患者的获益与非鳞癌组织学相当甚至更优 [8–10]。里程碑式的III期试验——包括KEYNOTE-407（帕博利珠单抗联合化疗一线治疗鳞状NSCLC）、CheckMate-017（纳武利尤单抗二线治疗鳞状NSCLC）和IMpower-131（阿替利珠单抗联合化疗治疗鳞状NSCLC，显示PFS获益但在ITT人群中OS改善未达统计学显著性）——确立了以免疫治疗为基础的方案作为新标准治疗 [8,11]。更近期的cemiplimab在超高PD-L1表达（TPS≥90%）的LUSC患者中展现出特别的前景 [12]。')

p('尽管取得这些进展，仅少数患者获得持久临床获益。帕博利珠单抗单药一线治疗PD-L1高表达（≥50%）患者的ORR约为40-45%，鳞状NSCLC联合化疗后提升至约55-60%（KEYNOTE-407）[8,13]。然而，相当比例患者表现为原发性耐药（无初始临床获益的疾病进展），许多初始应答者最终发展为获得性耐药（初始疾病控制后的肿瘤再生长）[14]。Lung-MAP S1400F子研究专门评估了度伐利尤单抗联合tremelimumab在抗PD-(L)1耐药的鳞状NSCLC中的疗效，报告原发性耐药队列ORR仅7%，获得性耐药队列为0%，生动说明了克服已建立免疫治疗耐药所面临的巨大挑战 [15]。')

p('支撑这些令人失望临床结局的是一个跨越三个相互关联领域的复杂耐药机制网络：肿瘤细胞内在改变、TME动态重塑，以及驱动克隆演化和表型适应性的治疗诱导选择压力 [14,16]。LUSC的TME具有若干共同构建免疫逃逸生态位的特征：M2极化TAM、Treg和MDSC高比例存在；以TGF-β和IL-6为主导的细胞因子环境；以及以缺氧、乳酸酸中毒和营养耗竭为特征的代谢有害微环境 [16,17]。此外，近期多组学研究在28-36%的LUSC患者中鉴定出独特的免疫耗竭型（EIC），其特征为密集的淋巴细胞浸润与多达九个抑制性免疫检查点的共上调矛盾性共存 [18]。这种"炎症浸润但功能抑制"的免疫状态凸显了单一PD-1/PD-L1阻断的不足，并强调了以机制为依据的组合策略的必要性。')

p('文献筛选和选择流程详见附图1（PRISMA流程图）。', sz=10)

p('在本综述中，我们对LUSC免疫治疗耐药机制的当前认识进行了全面整合，按照三个相互关联的维度进行组织（附图2）。首先综述LUSC的免疫景观及分子分类框架；其次剖析肿瘤内在耐药机制（致癌信号通路、表观遗传失调、EMT驱动的免疫排斥、细胞死亡程序缺陷）；第三审视TME介导的耐药（免疫抑制细胞群体、细胞因子网络、代谢重编程、瘤内微生物组）；随后讨论获得性耐药（克隆演化、组织学转化、表型可塑性）；最后评估克服策略（合理联合治疗、TME重塑、精准免疫治疗）。通过整合这些机制领域的发现，本综述旨在为理解LUSC免疫治疗耐药的多面性提供框架，并为改善这一挑战性疾病的结局确定可行机会。')

page_break()

# ============= 2. IMMUNE LANDSCAPE =============
h('2  LUSC的免疫景观', 1)

p('理解LUSC免疫治疗耐药机制需要对其独特的肿瘤免疫微环境（TIME）有一个基本的认识。尽管LUSC与LUAD在肺中共享共同的解剖起源，但两者的免疫景观存在实质性差异，反映了突变过程、基质组成和演化轨迹的不同。scRNA-seq、空间转录组学和多组学整合的最新进展以前所未有的分辨率剖析了LUSC TIME的细胞和分子架构。')

h('2.1  TIME异质性与分子分类', 2)

p('LUSC的TIME以深刻的瘤间和瘤内异质性为特征，从根本上塑造了对免疫治疗的应答 [16]。以经典的肿瘤免疫表型三分法——免疫炎症型、免疫排斥型和免疫沙漠型——为基础，近期研究已利用整合多组学方法针对LUSC进行了细化 [16,19]。')

p('免疫炎症型LUSC的特征是肿瘤实质内CD8+细胞毒性T淋巴细胞（CTL）、活化CD4+记忆T细胞和树突状细胞（DC）的丰富浸润，表达高水平效应细胞因子和溶细胞标志物，与持续进行但最终无效的抗肿瘤免疫应答相一致 [16]。矛盾的是，许多免疫炎症型LUSC同时上调多个抑制性免疫检查点并携带高频率的免疫抑制性Treg和M2极化巨噬细胞，形成一种"炎症浸润但功能抑制"的免疫状态（frequently underpins primary resistance to single-agent PD-1/PD-L1 blockade）[18,20]。')

p('免疫排斥型LUSC的特征是CTL和其他免疫效应细胞被限制在肿瘤巢周围的基质区室中，无法穿透肿瘤实质。这种排斥主要由癌相关成纤维细胞（CAF）介导，它们沉积包括胶原、纤连蛋白和层粘连蛋白在内的致密ECM成分，形成T细胞浸润的物理屏障 [16]。POSTN+ CAF和FAP+ CAF已被证明可产生趋化排斥梯度并使ECM纤维垂直于肿瘤边界排列，有效地将T细胞困在肿瘤周围基质中。空间转录组分析揭示了CAF亚群与APOE+ TAM的紧密共定位，鉴定出一种通过协同细胞因子和趋化因子分泌来强化免疫排斥的CAF-TAM信号轴 [16]。')

p('免疫沙漠型LUSC占少数病例，其标志是肿瘤和基质区室内几乎完全缺乏T细胞。这种表型被认为源于先天性免疫感知缺陷——包括STING/cGAS通路激活受损——以及DC启动不足，导致无法启动有效的抗肿瘤T细胞应答 [16,19]。')

p('在这些定性描述之外，整合基因组规模分析已建立了LUSC的分子分类框架 [19,21]。Yin等人通过免疫相关基因表达谱的无监督聚类鉴定出LUSC的不同免疫亚型 [19]；Yang等人鉴定出一种细胞因子主导的免疫抑制类（EIC），展现出具有直接治疗意义的独特分子特征 [18]；Song等人利用513例LUSC样本划分了六个分子亚型（CS1-CS6），并鉴定出CS3为一种淋巴细胞浸润亚型，矛盾性地表现出升高的耗竭标志物（CTLA-4、LAG-3、PD-1）并预测对ICB治疗耐药 [20]。这些分类工作的共同结论是：肿瘤浸润淋巴细胞的单纯存在不足以预测LUSC的ICI应答；免疫浸润的功能状态和空间组织才是治疗结局的关键决定因素。')

h('2.2  LUSC与LUAD：分歧的免疫环境', 2)

p('尽管LUSC和LUAD均被归类为NSCLC，但两者起源于不同的细胞源头，携带根本不同的基因组景观，并展现出明显分歧的免疫微环境 [22,23]。')

p('在基因组层面，LUAD富含可靶向驱动突变——最显著的是EGFR（东亚~30-40%，西方~10-15%）、KRAS（~30%）、ALK重排（~5%）和ROS1融合（~1-2%）——这些突变在EGFR突变亚群中与相对较低的TMB和较低的炎症浸润相关 [22]。相比之下，LUSC主要以致癌基因功能丧失突变（TP53、CDKN2A、KEAP1）和反复扩增（PIK3CA、FGFR1、SOX2）为主导，产生较高的总TMB和新抗原负荷，理论上预测其免疫原性更高 [22,23]。然而，这种升高的新抗原负荷并未转化为LUSC相较于LUAD更优的ICI应答，提示存在强效的对抗性免疫抑制机制。')

p('在细胞层面，比较分析揭示了LUSC与LUAD在免疫细胞组成上的系统性差异 [23]。LUSC肿瘤表现出更高的M2巨噬细胞和静息CD4+记忆T细胞浸润，而LUAD肿瘤倾向于有更高的幼稚B细胞和浆细胞浸润。LUSC也表现出更高的免疫排斥型比例，由更广泛的CAF激活和ECM重塑驱动 [16,22]。LUSC的细胞因子环境偏向TGF-β和IL-6/STAT3信号，而LUAD更常与EGFR驱动的免疫抑制程序相关 [22]。')

p('这些亚型特异性免疫特征具有直接的临床意义。第一，LUSC中免疫排斥的高比例提示靶向CAF介导的基质屏障的策略可能对该组织学类型尤为适用。第二，LUSC和LUAD之间不同的免疫检查点表达谱可指导联合免疫治疗搭档的理性选择 [22]。第三，在LUAD中开发的预后和预测性生物标志物可能无法直接用于LUSC。这些考虑强调了LUSC特异性免疫治疗策略的必要性。')

h('2.3  免疫耗竭型（EIC）：理解耐药的框架', 2)

p('LUSC免疫生物学的一项里程碑式贡献是Yang等人通过对624例LUSC样本进行RNA测序数据的无监督聚类，鉴定出了免疫耗竭型（EIC）[18]。约28-36%的LUSC患者属于EIC，其特征是具有对免疫治疗耐药至关重要意义的独特分子标志。')

p('EIC由四个标志性特征的共同存在所定义 [18]（附图3）。第一，该类肿瘤展现出T细胞耗竭标志的显著富集，伴随典型耗竭相关转录因子TOX和EOMES的表达升高，以及TCF-1——保留增殖能力和PD-1阻断应答性的祖细胞耗竭T细胞（Tpex）标志物——的表达降低。这种从Tpex主导到终末耗竭T细胞区室的转变代表了ICI疗效的关键屏障。')

p('第二，EIC的特征是高达九个抑制性免疫检查点的共上调——CTLA-4、PD-1（PDCD1）、LAG-3、BTLA、TIGIT、TIM-3（HAVCR2）、IDO1、SIGLEC7和VISTA——代表一种广泛的、多受体的免疫抑制状态，使得单一PD-1/PD-L1阻断在机制上就不足 [18]。这一发现为以下临床观察提供了分子层面的解释：许多具有高PD-L1表达的LUSC患者仍然对帕博利珠单抗或纳武利尤单抗单药治疗无应答。')

p('第三，EIC肿瘤被免疫抑制细胞群体高度浸润，特别是M2极化巨噬细胞和CD4+FOXP3+ Treg。这些细胞产生免疫抑制性细胞因子——包括TGF-β、IL-10和CCL18——直接抑制CTL效应功能并促进进一步的T细胞耗竭 [18]。')

p('第四，EIC矛盾性地与高总密度的TIL相关，然而这些患者的预后显著差于TIL浸润较低的患者。这一发现强化了免疫浸润数量与免疫能力之间的关键区别，并突显了仅用TIL密度作为LUSC中ICI应答生物标志物的不足。')

p('EIC概念已得到后续研究的证实和扩展。Song等人在其CS3分子亚型中鉴定出类似表型，该亚型表现出高淋巴细胞浸润伴随升高的耗竭标志物并通过独立计算方法预测ICB耐药 [20]。此外，CS3亚型被发现特异性上调LAMC2-CD44分子轴——一种与EMT相关的通路，同时调控肿瘤增殖和免疫排斥，提供了肿瘤细胞内在程序与耗竭免疫表型之间的潜在机制联系。')

p('总之，LUSC TIME的特征——涵盖其异质性细胞组成、相对于LUAD的分歧特征以及具有临床关键意义的EIC——为特定耐药机制的运作奠定了生物学基础。后续章节将详细剖析这些机制，首先从肿瘤细胞内在的耐药决定因素开始。')

page_break()

# ============= 3. TUMOR-INTRINSIC =============
h('3  肿瘤内在耐药机制', 1)

p('肿瘤细胞内在改变代表了LUSC免疫治疗耐药的一个基本层面，通过多种机制运作，汇聚于一个共同终点：免疫效应细胞无法识别和清除恶性细胞。这些机制涵盖从积极塑造免疫抑制微环境的致癌信号通路，到沉默抗原呈递机制的表观遗传程序，再到使肿瘤细胞抵抗CTL介导杀伤的细胞死亡通路缺陷。')

h('3.1  致癌信号通路', 2)

h('3.1.1  PI3K/AKT/mTOR通路', 3)
p('磷脂酰肌醇3-激酶（PI3K）通路是LUSC中最频繁激活的致癌网络之一，PIK3CA扩增和激活突变发生于约30-40%的病例 [5]。PI3K/AKT/mTOR信号的激活在促进肿瘤细胞增殖和存活的同时，通过多种机制驱动免疫逃逸：通过AKT介导的PD-L1 mRNA稳定化上调PD-L1表达；增加免疫抑制性细胞因子（IL-10、TGF-β）的分泌；促进TME内M2巨噬细胞极化 [5,7]。临床前研究已证明mTOR抑制可逆转PI3K驱动的PD-L1表达并恢复T细胞介导的杀伤，为PIK3CA改变的LUSC中PI3K/mTOR抑制剂与ICI联合提供了理论基础 [5]。')

h('3.1.2  KEAP1/NRF2通路', 3)
p('KEAP1突变发生于约12%的LUSC，已成为NSCLC亚型中对ICI原发性耐药的最强有力的预测因子之一 [5,7]。生理条件下，KEAP1靶向转录因子NRF2进行蛋白酶体降解。KEAP1功能丧失突变导致NRF2持续稳定和核转位，驱动广泛的抗氧化和细胞保护基因转录程序 [5]。这一激活的NRF2程序通过多种汇聚机制深刻改变肿瘤细胞的免疫原性：增强的活性氧（ROS）清除减少免疫原性细胞死亡并损害DC激活；外排转运蛋白上调限制CTL释放的细胞毒性分子的胞内积累；抑制促炎细胞因子产生（包括I型干扰素）减弱先天性免疫感知 [5,7]。临床上，KEAP1突变的LUSC与ICI治疗显著缩短的无进展生存期相关，新兴证据提示这些肿瘤可能需要独特的治疗策略——可能包括NRF2抑制剂、谷氨酰胺酶拮抗剂或STING激动剂 [5]。')

h('3.1.3  MAPK/p38信号', 3)
p('在充分表征的RAS/RAF/MEK/ERK级联之外，应激激活的p38 MAPK通路近期已被牵涉到LUSC免疫逃逸中。Lin等人证明转谷氨酰胺酶2（TGM2）——一种催化蛋白质交联的酶——在LUSC中显著过表达，并独立预测较差的总生存（P=0.00018）[24]。机制上，TGM2激活p38 MAPK信号，进而促进肿瘤细胞增殖、迁移、侵袭和抗凋亡。同时，TGM2高表达肿瘤展现出以Th1细胞浸润减少和免疫抑制细胞群体富集为特征的免疫抑制性TME [24]。TGM2的上游调控因子被鉴定为糖皮质激素受体NR3C1，其直接结合TGM2启动子并驱动其转录。这一NR3C1-TGM2-p38 MAPK轴代表了将应激信号与LUSC免疫逃逸联系起来的潜在可靶向的肿瘤内在机制 [24]。')

h('3.1.4  FGFR1和EGFR：RTK过表达的悖论', 3)
p('FGFR1扩增发生于约20%的LUSC，EGFR蛋白在肿瘤细胞表面频繁过表达 [5,7]。然而，针对这些受体的靶向药物在LUSC中的临床经验均令人失望。FGFR抑制剂（erdafitinib、infigratinib）在FGFR1扩增的LUSC中展现出极小的单药活性，可能因为FGFR1扩增在多种共发生基因组改变的背景下并不一致地转化为通路成瘾 [5]。EGFR靶向治疗的失败或许更具启发意义：尽管许多LUSC肿瘤中存在强大的EGFR蛋白表达，EGFR-TKI在临床上无效，而激活EGFR突变——LUAD中TKI敏感性的核心预测因子——在LUSC中极为罕见 [25]。机制研究揭示LUSC中的EGFR蛋白表达伴随下游RAS/MAPK和PI3K/AKT通路的组成性配体非依赖性激活，这些通路由来自FGFR1、PI3KCA和其他RTK的平行输入所维持 [25]。这一现象阐明了LUSC生物学的一个更广泛原理：致癌信号网络的冗余性和互联性限制了单一通路抑制的疗效 [5,7,25]。')

h('3.2  上皮间质转化（EMT）', 2)
p('上皮间质转化是一种发育程序，当在癌细胞中异常激活时，赋予侵袭和转移能力，同时驱动免疫逃逸 [21]。在LUSC中，EMT激活与免疫抑制性TME和ICI耐药强相关，通过多种交叉机制运作。Zhang等人构建了一个六基因EMT预后模型（GAB2、ALDOA、PCDHA3、TMEM92、ERH、IRS4），将LUSC患者分为具有显著不同免疫浸润特征的风险组：低风险肿瘤富集活化CD8+ T细胞、活化CD4+记忆T细胞和幼稚B细胞，而高风险肿瘤表现出高比例的静息CD4+记忆T细胞和M0巨噬细胞——这与适应性免疫参与减弱的模式一致 [21]。')

p('LUSC中EMT与免疫排斥之间的一个关键分子联系由Song等人鉴定：层粘连蛋白γ2（LAMC2）-CD44轴是CS3分子亚型中免疫耐药的关键驱动因子 [20]。LAMC2和CD44均为EMT相关基因，其共表达与免疫排斥密切相关。多重免疫荧光揭示了LAMC2-CD44表达与CD8+ T细胞浸润之间的反向空间关系。在LAMC2-CD44轴之外，EMT通过多种互补机制促进免疫耐药：间质肿瘤细胞上调多种免疫抑制因子包括TGF-β、VEGF和IL-6；E-cadherin下调损害CTL与肿瘤细胞之间的免疫突触形成；EMT相关转录因子ZEB1和SNAIL驱动PD-L1表达增加 [20,21]。')

h('3.3  表观遗传失调', 2)
p('异常DNA甲基化是LUSC的一个标志，通过表观遗传沉默对免疫识别至关重要的基因来促进免疫逃逸 [26]。全基因组甲基化分析揭示了多个抑癌基因以及抗原加工和呈递基因（TAP1、TAP2、LMP2、LMP7）启动子CpG岛的高甲基化。Sasa等人全面回顾了LUSC中DNA甲基化、lncRNA失调和免疫治疗应答之间的交互作用，强调DNA甲基转移酶抑制剂（DNMTi）如地西他滨和阿扎胞苷可恢复表观遗传沉默的免疫基因表达，并在临床前模型中与ICI协同 [26]。')

p('翻译后组蛋白修饰调控染色质可及性和基因表达。在LUSC中，组蛋白去乙酰化酶（HDAC）和PRC2催化亚基EZH2的过表达已被牵涉到免疫逃逸。HDAC介导的去乙酰化沉默干扰素刺激基因（ISG）和趋化因子位点，而EZH2介导的H3K27三甲基化抑制抑癌基因和MHC表达 [26]。临床前证据支持HDAC抑制剂与ICI联合增强肿瘤免疫原性。')

p('非编码RNA方面，哺乳动物转录组编码数千个长链非编码RNA（lncRNA）和环状RNA（circRNA），充当基因表达的主调控因子。在LUSC中，lncRNAs通过多种机制调控免疫相关通路：作为竞争性内源RNA（ceRNA）海绵吸附靶向免疫检查点基因的microRNA；将染色质修饰复合物支架到免疫基因位点；调控细胞因子和趋化因子的mRNA稳定性 [26]。一个显著实例是circHMGB2，它通过miRNA海绵机制驱动免疫抑制和抗PD-1耐药。circHMGB2的敲低在临床前模型中与抗PD-1治疗协同，导致CD8+ T细胞浸润增强和肿瘤消退 [27]。')

h('3.4  细胞死亡通路与免疫耐药', 2)
p('肿瘤细胞死亡的模式和免疫原性是抗肿瘤免疫的关键决定因素。免疫原性细胞死亡（ICD）——以钙网蛋白、ATP和HMGB1等损伤相关分子模式（DAMP）的释放为特征——促进DC成熟和肿瘤特异性CTL的高效交叉启动。反之，细胞死亡通路缺陷或向非免疫原性死亡模式的转变可导致免疫逃逸和ICI耐药 [28,29]。')

p('失巢凋亡耐药——肿瘤细胞在脱离ECM后存活的能力——近期已被联系到LUSC的免疫调控。Lu等人鉴定出三个失巢凋亡相关基因——S100A7、S100A8和SPP1——与LUSC免疫浸润模式显著相关，特别是与免疫抑制性Treg、M2巨噬细胞和DC的丰度 [28]。铁死亡——一种以脂质过氧化为特征的铁依赖性调节性坏死——已被牵涉到肿瘤抑制和免疫治疗应答。SLC7A11的表观遗传激活被证明定义了一个与LUSC DNA甲基化分类相关的铁死亡-免疫轴 [30]。焦亡——由gasdermin家族成员介导的高度免疫原性程序性细胞死亡——也已在LUSC中得到研究 [31]。这些发现共同突显了细胞死亡通路失调与LUSC免疫耐药之间的紧密联系。')

h('3.5  基因组不稳定性与新抗原动力学', 2)
p('虽然LUSC通常展现出相对于其他实体瘤较高的TMB——这一特征归因于慢性烟草致癌物暴露——但TMB与LUSC中ICI应答之间的关系比最初认识的要复杂 [7,26]。尽管TMB-high的LUSC肿瘤更可能携带免疫原性新抗原，但若干因素限制了单独使用TMB的预测效用。第一，克隆性新抗原（存在于所有肿瘤细胞中）比局限于少数细胞的亚克隆新抗原更能引发有效的抗肿瘤免疫，而LUSC中克隆性与亚克隆突变的比例在肿瘤间差异显著 [7]。第二，染色体不稳定性和非整倍体——两者均在LUSC中高发——可矛盾性地抑制抗肿瘤免疫，可能通过大量免疫相关基因的协同转录失调或通过损害抗原呈递的慢性ER应激诱导 [7]。第三，在ICI的治疗压力下，肿瘤可经历免疫编辑，伴随新抗原表达克隆的选择性消除和新抗原耗竭或HLA缺陷亚克隆的扩张——这是一个动态过程，对获得性耐药有贡献 [7,14]。')

p('上文讨论的肿瘤内在机制并非孤立运作，而是与周围的TME深度互联。致癌信号通路塑造细胞因子环境；EMT程序招募和激活基质细胞；表观遗传改变决定肿瘤细胞如何向免疫系统呈递自身。下一章节将审视TME的细胞和分子组分如何独立驱动LUSC的免疫治疗耐药。')

page_break()

# ============= 4. TME-MEDIATED =============
h('4  肿瘤微环境介导的耐药', 1)

p('肿瘤内在改变为免疫逃逸奠定分子基础的同时，肿瘤微环境——包括多样的免疫细胞群体、基质成分、细胞因子、代谢物和微生物群落——构成了免疫治疗成功或失败的战场。在LUSC中，TME以免疫抑制力量的汇聚为特征，即使在具有丰富T细胞浸润的肿瘤中也能集体驱动ICI耐药。')

h('4.1  T细胞耗竭与功能障碍', 2)
p('T细胞耗竭——一种由慢性抗原刺激引起的效应功能进行性丧失状态——是LUSC免疫逃逸的核心机制 [16]。耗竭T细胞的分化轨迹由一个层级转录程序主导：祖细胞耗竭T细胞（Tpex）以TCF-1表达为特征，由中等水平的TOX维持，保留增殖能力，是PD-1阻断的主要细胞靶点 [16,18]。随着耗竭进展，Tpex分化为终末耗竭T细胞（Tex-term），表达高水平TOX、多个抑制性受体，并表现出严重受损的细胞因子产生和溶细胞功能 [18]。关键的是，Tex-term细胞在很大程度上对PD-1治疗无效，因为其功能障碍状态由表观遗传印记而非持续的PD-1信号所维持 [16]。')

p('在Yang等人描述的LUSC EIC中，T细胞区室偏向于Tex-term主导状态，TCF-1+ Tpex频率降低，TOX和EOMES表达升高，以及众多抑制性受体的共表达 [18]。这种耗竭状态由多种TME衍生因子强化：高新抗原负荷导致的持续抗原暴露；包括IL-10、TGF-β和CCL18在内的免疫抑制性细胞因子；肿瘤细胞消耗葡萄糖和谷氨酰胺的代谢竞争；以及慢性I型干扰素信号——虽然最初具有免疫刺激作用，但持续存在时可矛盾性地驱动T细胞耗竭 [16,18]。')

h('4.2  多重免疫检查点共表达', 2)
p('LUSC EIC的一个定义性特征是多达九个抑制性免疫检查点的共上调——CTLA-4、PD-1、LAG-3、BTLA、TIGIT、TIM-3、IDO1、SIGLEC7和VISTA——创造了一种广泛的、多受体的免疫抑制状态 [18]。这种共表达模式具有关键的治疗意义：单独阻断PD-1/PD-L1轴留下了多个替代抑制通路完好无损，使通过其他检查点的代偿性信号得以维持T细胞抑制 [7,18]。这一概念得到Lung-MAP S1400F子研究临床经验的支持，在该研究中，PD-L1/CTLA-4双重阻断在既往抗PD-(L)1治疗获得性耐药患者中产生极小的临床获益 [15]。')

p('LAG-3和TIGIT均与PD-1在LUSC中频繁共表达，分别与不同的配体（MHC II类和CD155/PVR）结合，通过独立的胞内通路传导信号，使它们成为与PD-1阻断联合的有吸引力靶点 [7,18]。靶向PD-1和TIGIT的双特异性抗体rilvegostomig目前正在NSCLC的III期试验中进行评估；尽管TROPION-Lung10试验入组的是非鳞癌人群，但考虑到EIC中TIGIT共表达的高频率，TIGIT共阻断的强大生物学理论基础适用于包括LUSC在内的NSCLC各组织学类型 [32]。')

h('4.3  免疫抑制细胞群体', 2)

h('4.3.1  肿瘤相关巨噬细胞（TAM）', 3)
p('TAM代表LUSC TME中最丰富的免疫细胞群体，是免疫抑制的主要驱动者 [17]。在LUSC中，TAM主要向M2样表型极化，以CD163、CD206和APOE表达为特征，由CSF-1、IL-4、IL-10、乳酸和肿瘤分泌的外泌体等TME衍生信号所驱动。Ji等人证明醛缩酶A（ALDOA）——一种关键的糖酵解酶——在LUSC中显著过表达，并与巨噬细胞浸润强烈相关。空间转录组学和免疫荧光证实了ALDOA表达肿瘤细胞与CD68+巨噬细胞在LUSC组织中的共定位 [17]。功能分析揭示ALDOA驱动的糖酵解通量产生富含乳酸的微环境，促进M2巨噬细胞极化，建立一个自我强化的代谢-免疫回路：升高的糖酵解→乳酸分泌→M2 TAM极化→免疫抑制细胞因子产生→T细胞功能障碍 [17]。')

p('近期研究已在鳞状癌中鉴定出一个CAF-TAM信号轴，其中CAF衍生因子强化M2巨噬细胞极化，而TAMs相互分泌生长因子维持CAF激活 [16]。这种CAF-TAM共依赖关系在肿瘤-基质界面创造了空间上有组织的免疫抑制生态位，贡献于CD8+ T细胞排斥和功能抑制 [16]。')

h('4.3.2  癌相关成纤维细胞（CAF）', 3)
p('CAF是构成LUSC基质主要成分的活化成纤维细胞，通过多种机制促进免疫耐药 [5,16]。LUSC中CAF区室具有异质性，包含功能不同的亚群：分泌细胞因子和趋化因子的炎症性CAF（iCAF）；沉积ECM成分的肌成纤维细胞性CAF（myCAF）；以及可能调节T细胞应答的抗原呈递CAF（apCAF）[5]。CAF介导的最充分表征的免疫逃逸机制是T细胞的物理排斥。POSTN+和FAP+ CAF沉积致密ECM——特别是I型胶原、纤连蛋白和层粘连蛋白——形成阻止CTL穿透肿瘤巢的物理屏障 [5,16]。超越物理排斥，CAF分泌趋化因子——包括CXCL12和CCL2——主动排斥T细胞同时招募包括Treg和MDSC在内的免疫抑制群体 [5]。')

h('4.3.3  MDSC与Treg', 3)
p('MDSC是一类在LUSC TME中扩张的病理性活化未成熟髓系细胞的异质性群体，通过多种机制发挥强效免疫抑制功能：通过精氨酸酶-1和xCT转运体表达耗竭T细胞活化必需的氨基酸精氨酸和半胱氨酸；产生ROS和过氧亚硝酸盐亚硝基化TCR组分；分泌IL-10和TGF-β促进Treg扩张并抑制DC成熟 [5,16]。CD4+FOXP3+ Treg在LUSC TME中富集，尤其在EIC亚型中 [18]。Treg通过接触依赖性机制——包括CTLA-4介导的DC上CD80/CD86转内吞——以及通过分泌免疫抑制性细胞因子（IL-10、TGF-β、IL-35）抑制抗肿瘤免疫。LUSC TME的趋化因子环境以TAM和肿瘤细胞产生CCL22和CCL17为特征，主动招募CCR4+ Treg [18]。')

h('4.4  细胞因子网络', 2)
p('LUSC细胞因子环境由免疫抑制介质的网络主导。TGF-β是核心协调者，由肿瘤细胞、CAF、TAM和Treg产生。TGF-β信号在CD8+ T细胞中直接抑制溶细胞效应分子的表达——穿孔素、颗粒酶B和IFN-γ——同时促进Treg分化和肿瘤细胞EMT [5,16]。EIC以升高的TGF-β和CCL18表达为特征：主要由M2 TAM分泌的CCL18招募额外的Treg和未成熟DC，同时促进进一步的TAM M2极化 [18]。IL-6激活肿瘤细胞和免疫细胞中的STAT3信号，驱动促进肿瘤进展和免疫逃逸的慢性炎症前馈回路。IL-10由Treg、TAM和耗竭T细胞产生，广泛抑制DC和巨噬细胞的抗原呈递、共刺激分子表达和促炎细胞因子产生 [18]。')

h('4.5  代谢免疫调控', 2)
p('LUSC TME以Warburg效应——有氧糖酵解导致肿瘤细胞高速率消耗葡萄糖和产生乳酸——为代谢特征 [17]。乳酸积累将TME酸化至pH 6.0-6.5，直接抑制T细胞增殖、细胞因子产生和溶细胞活性，同时促进Treg稳定性和M2巨噬细胞极化 [17]。葡萄糖耗竭的TME在代谢上饥饿高度依赖糖酵解执行效应功能的CTL，而可利用脂肪酸氧化的Treg在低葡萄糖条件下维持代谢优势 [17,33]。')

p('缺氧是LUSC TME的普遍特征，源于异常肿瘤血管系统和高代谢需求 [34]。HIF-1α在低氧张力下稳定，转录激活同时促进肿瘤存活和免疫逃逸的程序：PD-L1上调；VEGF分泌促进异常血管生成和进一步缺氧；MDSC和TAM的招募。近期阐明的一个联系缺氧与免疫耐药的机制涉及STING通路——先天性免疫感知胞质DNA的关键节点 [34]。缺氧通过HIF-1α介导的机制抑制STING信号，导致I型干扰素产生减少、DC成熟受损和肿瘤特异性CD8+ T细胞交叉启动减弱。恢复STING信号——通过缓解缺氧或直接STING激动——已在临床前LUSC模型中显示增强ICI疗效 [34]。')

h('4.6  肿瘤微生物组-免疫串扰', 2)
p('LUSC TME生物学的一个新兴维度是瘤内微生物组在调控免疫应答和免疫治疗结局中的作用 [33]。多组学整合揭示LUSC中瘤内微生物群的组成与乳酸代谢、免疫细胞浸润模式和免疫检查点表达显著相关。特定细菌分类群与不同的免疫表型相关：乳杆菌属的富集与以M2巨噬细胞为主和CD8+ T细胞应答减弱的免疫抑制性TME相关 [33]。尽管LUSC微生物组与免疫耐药之间的机制联系尚在定义中，这些发现提出了微生物组导向干预——包括抗生素、益生菌或代谢物补充——可被用于调控LUSC免疫治疗应答的可能性。')

p('上文描述的TME介导的耐药机制并非在治疗时刻从头产生，而是在ICI治疗的选择压力下动态演化。下一章节将审视治疗压力如何通过克隆演化、组织学转化和表型可塑性驱动获得性耐药。')

page_break()

# ============= 5. ACQUIRED RESISTANCE =============
h('5  获得性耐药机制', 1)

p('原发性耐药反映了肿瘤-免疫系统预先存在的无法产生有效抗肿瘤应答的能力，而获得性耐药在初始有效的免疫治疗选择压力下出现，表现为一段临床获益期后的疾病进展 [14]。LUSC的获得性耐药通过多种机制——基因组、表观基因组和表型——产生，这些机制共同使肿瘤细胞能够逃避持续的免疫攻击。')

h('5.1  免疫治疗压力下的克隆演化', 2)
p('ICI的施用对遗传异质性的肿瘤细胞群体施加了强大的免疫选择压力，驱动达尔文式的克隆演化过程 [7,14]。携带赋予免疫逃逸优势突变的肿瘤亚克隆——通过降低免疫原性、增强免疫抑制因子分泌或激活抗凋亡通路——在ICI治疗期间被正向选择，最终成为导致疾病进展的主导群体 [14]。已在NSCLC中记录到若干获得性耐药的基因组机制，与LUSC相关。JAK1和JAK2功能丧失突变废除IFN-γ受体信号，使肿瘤细胞对T细胞来源的IFN-γ不敏感 [14]。同样，获得性B2M突变或杂合性缺失消除细胞表面MHC I类表达，使肿瘤细胞能够完全逃避CD8+ T细胞识别 [7,14]。')

h('5.2  组织学转化：腺鳞转化', 2)
p('与LUSC生物学相关的一种尤为有启发性的获得性耐药形式是腺癌向鳞状细胞癌转化（AST）[36]。最初在LUAD的EGFR-TKI耐药背景下被记录，AST日益被认为是ICI获得性耐药的机制。Xu等人全面回顾了腺鳞转化背后的免疫机制，证明这种组织学转换伴随TIME的深刻重塑 [36]。在AST期间，LUAD细胞经历由鳞状转录因子——包括SOX2、TP63和ZEB1——激活以及LUAD谱系决定因子NKX2-1（TTF-1）的伴随抑制所驱动的谱系可塑性程序 [36]。这种转录重编程伴随免疫微环境的广泛变化：转化的鳞状样肿瘤细胞分泌独特的细胞因子和趋化因子谱，改变免疫细胞招募模式；ECM以CAF依赖性方式重塑促进免疫排斥；免疫检查点分子表达谱发生变化 [36]。')

h('5.3  表型可塑性', 2)
p('超越完全的组织学转化，肿瘤细胞可经历更细微的表型转变来赋予ICI耐药。Wang等人表征了LUSC中表型可塑性相关基因表达模式，证明伴淋巴结转移和不伴淋巴结转移的患者展现显著不同的可塑性基因组特征 [37]。表型可塑性相关预后标志成功将LUSC患者分为具有不同生存结局的高可塑性和低可塑性组。低可塑性评分的患者被预测对PD-L1抑制剂、顺铂和紫杉醇等更为敏感，而高可塑性肿瘤与广泛治疗耐药相关 [37]。可塑性驱动的耐药尤为具有挑战性，因为它是动态和可逆的，这意味着赋予耐药的表型状态可能无法被单次时间点活检捕获 [14,37]。这强调了液体活检——包括ctDNA和CTCs——在治疗压力下纵向监测耐药机制的潜在价值 [7]。')

h('5.4  SCLC转化及其他谱系转换', 2)
p('表型可塑性的一种极端表现是NSCLC——包括LUSC——在ICI治疗下转化为小细胞肺癌（SCLC）。虽然最初在EGFR突变腺癌中被描述，SCLC转化也已在LUSC中被记录 [38]。这一现象虽然罕见（ICI治疗的NSCLC中估计频率<5%），但具有临床毁灭性，因为转化的SCLC具有侵袭性且对现有治疗应答不佳。SCLC转化以RB1和TP53功能丧失、ASCL1或NEUROD1神经内分泌转录程序的激活以及获得特征性SCLC形态为特征 [38]。LUSC生物学的特定特征——如基线TP53突变的高发率——是否赋予相较于LUAD更高的SCLC转化风险，仍是一个悬而未决的问题。')

p('获得性耐药因此代表了肿瘤在免疫压力下演化的顶点，通过基因组选择、转录重编程和表型适应性运作。获得性耐药的多因素性质——多种耐药机制在单个患者中频繁共存——提出了一个艰巨的治疗挑战，并推动了旨在先发制人或规避这些逃逸通路的联合策略的开发。')

page_break()

# ============= 6. STRATEGIES =============
h('6  克服耐药的策略', 1)

p('LUSC免疫治疗耐药的多维性质——涵盖肿瘤内在改变、TME介导的抑制和治疗诱导的适应性——需要同样多维的治疗策略。新兴的克服耐药方法可大致分为合理的免疫联合治疗、靶向肿瘤内在通路、TME重塑和生物标志物引导的精准策略四类。附图4展示了基于TIME亚型选择机制匹配干预的决策框架；表2和表3分别总结了可操作的肿瘤内在和TME介导的靶点。')

h('6.1  合理的免疫联合治疗', 2)
p('LUSC中多个抑制性免疫检查点的共表达——尤其在EIC中——为双重或三重检查点阻断提供了清晰的生物学理论基础 [7,18]。虽然Lung-MAP S1400F子研究证明在非选择的抗PD-(L)1耐药患者中增加CTLA-4阻断（tremelimumab）仅产生最小获益，但这并不否定更精准靶向组合的潜力 [15]。新型双特异性抗体——包括QL1706（PD-1/CTLA-4）和rilvegostomig（PD-1/TIGIT）——提供了同时双靶点接合的优势，通过优先靶向共表达两种受体的T细胞可能改善治疗指数 [39]。PD-1阻断联合TIGIT抑制在LUSC中尤为引人关注，鉴于EIC中TIGIT共表达的高频率 [18]。ICI联合铂类化疗仍然是LUSC的标准一线方案（KEYNOTE-407），化疗可能通过ICD诱导、MDSC耗竭和Treg减少等多种机制来克服耐药 [5]。')

h('6.2  靶向肿瘤内在通路', 2)
p('致癌信号与免疫逃逸的汇聚为通路靶向的免疫治疗联合创造了机会 [5,26]。在PIK3CA改变的LUSC中，PI3K/mTOR抑制剂可同时抑制肿瘤生长并解除mTOR驱动的PD-L1表达。在KEAP1突变的LUSC中，靶向NRF2通路的多种治疗策略正在探索中：利用NRF2激活所创造的代谢依赖性的谷氨酰胺酶抑制剂，或绕过受抑制的先天性免疫感知的STING激动剂 [5]。表观遗传启动——使用DNMT抑制剂或HDAC抑制剂恢复沉默的抗原呈递和干扰素应答基因表达，随后给予ICI——代表了由临床前证据支持的机制上有吸引力的策略 [26]。')

h('6.3  TME重塑', 2)
p('克服TME介导的耐药需要将免疫抑制环境转化为允许抗肿瘤免疫的策略 [16]。CAF靶向方法——包括FAP抑制剂、Hedgehog通路拮抗剂和TGF-β捕获剂——旨在减少ECM沉积、正常化肿瘤血管并促进T细胞穿透肿瘤巢 [16]。TAM重编程策略——如CSF1R抑制、CD47阻断或CD40激动——寻求将巨噬细胞极化平衡从M2转向M1 [16,17]。细胞因子导向治疗——包括TGF-β中和抗体（fresolimumab）和耗竭Treg的CCR4抑制剂（mogamulizumab）——旨在中断免疫抑制信号回路 [16]。在LUSC缺氧下受抑制的STING/cGAS通路代表了先天性免疫激活的有吸引力靶点，STING激动剂能够触发I型干扰素应答，即使在免疫沙漠型肿瘤中也能招募和激活DC [34]。')

h('6.4  新兴策略与生物标志物引导的精准免疫治疗', 2)
p('展望未来，若干新兴方向对LUSC具有特殊前景 [5,16]。附图4提供了基于TIME亚型选择治疗的决策框架。建立TIME分类框架——整合分子亚型、免疫浸润模式和空间架构——可指导机制匹配的联合策略选择：对具有TIGIT共表达的EIC肿瘤使用ICI+TIGIT阻断；对具有高基质含量的免疫排斥型肿瘤使用CAF靶向治疗；对缺乏T细胞浸润的免疫沙漠型肿瘤使用STING激动剂或溶瘤病毒 [16]。基于液体活检的耐药机制动态监测——通过ctDNA分析新出现的JAK1、JAK2或B2M突变，或随时间追踪TMB和新抗原克隆性——可实现适应性治疗策略 [7]。最后，瘤内微生物组在调控ICI应答中的新兴角色提出了微生物组导向干预——包括选择性抗生素、特定益生菌组合或粪菌移植——可被用于增强LUSC免疫治疗疗效的可能性 [33]。')

page_break()

# ============= 7. CONCLUSIONS =============
h('7  结论与未来展望', 1)

p('免疫治疗已改变了肺鳞状细胞癌患者的治疗格局，然而大多数患者要么初始不应答，要么最终发展出耐药。本综述从三个相互关联的维度——肿瘤内在改变、TME介导的抑制和获得性耐药——整合了对LUSC免疫治疗耐药机制的最新认识，得出若干核心原则。')

p('第一，LUSC的耐药很少归因于单一机制，而是产生于多个相互强化通路的汇聚。KEAP1/NRF2突变肿瘤细胞同时抵抗氧化杀伤、抑制先天性免疫感知并分泌促进免疫抑制TME的因子。CAF-TAM轴同时物理排斥T细胞并抑制那些得以浸润的T细胞。EIC共表达九个抑制性检查点，当其中一个被阻断时，任一其他检查点均可维持免疫抑制。这一机制冗余性原则解释了为何单药PD-1/PD-L1阻断对相当比例的LUSC患者不足，并推动了合理设计的、以机制为依据的联合策略的开发。')

p('第二，免疫浸润的功能状态比其数量更重要。EIC——具有丰富的TIL、高PD-L1表达和矛盾性不良预后——诠释了免疫浸润与免疫能力之间的关键区别。仅依赖TIL密度或PD-L1免疫组化的生物标志物策略本质上具有局限性；未来的预测模型必须纳入T细胞功能状态、检查点共表达模式和空间组织的多维评估。')

p('第三，LUSC在免疫生物学上不同于LUAD，需要组织学特异性的治疗策略。免疫排斥表型的更高发率、以抑癌基因丧失而非致癌驱动突变为特征的独特基因组景观，以及LUSC TME的特定CAF和代谢特征，都反对在没有LUSC特异性验证的情况下外推源自LUAD的治疗范式。')

p('展望未来，若干优先事项浮现。按TIME亚型、KEAP1突变状态或EIC分子标志进行患者分层的前瞻性生物标志物驱动的临床试验对于评估机制匹配的联合治疗至关重要。多组学整合——结合基因组、转录组、空间蛋白质组和微生物组数据——将提供解剖单个患者体内运作的耐药机制全复杂度所需的分辨率。通过液体活检ctDNA分析和循环免疫细胞分析的耐药动态监测将实现实时响应新兴耐药的适应性治疗策略。最后，包括双特异性抗体、表观遗传启动、STING激动剂和CAF靶向治疗在内的新型治疗模式有望扩展克服LUSC免疫治疗耐药可用的武器库。最终，克服LUSC免疫治疗耐药将需要从一刀切方法转向精准医学范式，在该范式中，每个患者肿瘤中运作的特定耐药机制被识别并通过治疗加以应对。')

page_break()

# ============= REFERENCES =============
h('参考文献', 1)

refs = [
    '[1] Sung H, Ferlay J, Siegel RL, et al. Global Cancer Statistics 2020. CA Cancer J Clin. 2021;71(3):209-249. PMID: 33538338',
    '[2] Travis WD, Brambilla E, Nicholson AG, et al. The 2015 WHO classification of lung tumors. J Thorac Oncol. 2015;10(9):1243-1260. PMID: 26291008',
    '[3] Siegel RL, Miller KD, Fuchs HE, Jemal A. Cancer statistics, 2022. CA Cancer J Clin. 2022;72(1):7-33. PMID: 35020204',
    '[4] Schiller JH, Harrington D, Belani CP, et al. Comparison of four chemotherapy regimens for advanced NSCLC. N Engl J Med. 2002;346(2):92-98. PMID: 11784886',
    '[5] Niu Z, Jin R, Zhang Y, Li H. Signaling pathways and targeted therapies in lung squamous cell carcinoma. Signal Transduct Target Ther. 2022;7:353. PMID: 36198685',
    '[6] Cancer Genome Atlas Research Network. Comprehensive genomic characterization of squamous cell lung cancers. Nature. 2012;489(7417):519-525. PMID: 22960745',
    '[7] Yuan H, Liu J, Zhang J. The current landscape of immune checkpoint blockade in metastatic LUSC. Molecules. 2021;26(5):1392. PMID: 33807509',
    '[8] Paz-Ares L, Luft A, Vicente D, et al. Pembrolizumab plus chemotherapy for squamous NSCLC. N Engl J Med. 2018;379(21):2040-2051. PMID: 30280635',
    '[9] Brahmer J, Reckamp KL, Baas P, et al. Nivolumab versus docetaxel in advanced squamous-cell NSCLC. N Engl J Med. 2015;373(2):123-135. PMID: 26028407',
    '[10] Chen W, Liu H, Li Y, et al. First-line immunotherapy efficacy in advanced squamous NSCLC with PD-L1 >=50%. Front Oncol. 2024;14:1360583. PMID: 38725635',
    '[11] Jotte R, Cappuzzo F, Vynnychenko I, et al. Atezolizumab in combination with carboplatin and nab-paclitaxel in advanced squamous NSCLC. J Thorac Oncol. 2020;15(8):1351-1360. PMID: 32302702',
    '[12] Ikeda S, Araki K, Kitagawa M, et al. Why cemiplimab? Defining a unique therapeutic niche. Cancers. 2026;18(2):272. PMID: 41595192',
    '[13] Reck M, et al. Pembrolizumab versus chemotherapy for PD-L1-positive NSCLC. N Engl J Med. 2016;375(19):1823-1833. PMID: 27718847',
    '[14] Sharma P, Hu-Lieskovan S, Wargo JA, Ribas A. Primary, adaptive, and acquired resistance to cancer immunotherapy. Cell. 2017;168(4):707-723. PMID: 28187290',
    '[15] Leighl NB, Redman MW, Rizvi N, et al. Phase II study of durvalumab plus tremelimumab in anti-PD-(L)1 resistant sqNSCLC. J Immunother Cancer. 2021;9(8):e002973. PMID: 34429332',
    '[16] Tong Y, Wang Y, Chen Y, Fan Y, Li H. Decoding the tumor immune microenvironment in LUSC. Transl Lung Cancer Res. 2025;14(4):1170-1190. PMID: 41133013',
    '[17] Ji Y, Li X, Shen X, et al. Aldolase A in pan-cancer and LUSC. Cancer Cell Int. 2025;25:184. PMID: 41239433',
    '[18] Yang M, Lin C, Wang Y, et al. Identification of a cytokine-dominated immunosuppressive class in LUSC. Genome Med. 2022;14(1):72. PMID: 35799269',
    '[19] Yin L, et al. Identification of immune subtypes of LUSC. Front Oncol. 2021;11:778324. PMID: 35186710',
    '[20] Song T, Yang Y, Wang Y, et al. Laminin gamma2-CD44 immune resistance in LUSC. Heliyon. 2024;10(11):e31299. PMID: 38803944',
    '[21] Zhang A, He J, Lin Q. EMT-based risk model for LUSC. PeerJ. 2026;14:e21117. PMID: 42089102',
    '[22] Shen Y, Chen JQ, Li XP. Differences between LUAD and LUSC. Genes Dis. 2025. PMID: 40083325',
    '[23] Yan T, et al. Immune heterogeneity between LUAD and LUSC. Front Immunol. 2021;12:703797. PMID: 34394068',
    '[24] Lin C, Li S, Yi L, et al. TGM2 regulated by NR3C1 drives p38 MAPK-mediated immune evasion in LUSC. Front Immunol. 2025;16:1547241. PMID: 41050683',
    '[25] Ju L, et al. Intrinsic resistance of LUSC to EGFR-TKI. Front Oncol. 2020;10:568878. PMID: 33133263',
    '[26] Sasa GBK, Xuan C, Chen M, Jiang Z, Ding X. lncRNAs, immunotherapy and DNA methylation in LUSC. Transl Cancer Res. 2021;10(12):5324-5341. PMID: 35116387',
    '[27] Zhang LX, Gao J, Long X, et al. circHMGB2 drives anti-PD-1 resistance via miR-181a-5p/CARM1. Mol Cancer. 2022;21(1):110. PMID: 35525959',
    '[28] Lu H, Huang W, Shen Q, Liu R. Anoikis-related genes in LUSC. Med Sci Monit. 2026;31:e951722. PMID: 41902322',
    '[29] Ou D, Wang H, Liu Y, Nie J, Liu D. Anoikis resistance gene model in LUSC. Discover Oncol. 2026. PMID: 41530460',
    '[30] Lu HP, Nong K, Pang L, et al. SLC7A11 ferroptosis-immune axis in LUSC. PeerJ. 2026;14:e20686. PMID: 41700135',
    '[31] Deng X, et al. Pyroptosis in LUSC immune microenvironment. J Cancer Res Clin Oncol. 2022. PMID: 36123889',
    '[32] Newsom-Davis T, et al. TROPION-Lung10. Front Oncol. 2025;15:1721624. PMID: 41669261',
    '[33] Qiu X, Li D. Intratumor microbiome and lactic acid metabolism in LUSC. Front Immunol. 2025. PMID: 40568577',
    '[34] Chen F, Wen X, Li S, et al. Hypoxia-STING chemo-immuno resistance in LUSC. Transl Oncol. 2025;52:102350. PMID: 40138855',
    '[35] Zhao F, et al. Hypoxia-related lncRNAs in LUSC. Front Oncol. 2021. PMID: 34250747',
    '[36] Xu H, Yang Y, Wang P, et al. Immune mechanisms in lung adenosquamous transformation. Front Immunol. 2025;16:1502584. PMID: 40568576',
    '[37] Wang F, Zhu L. Phenotypic plasticity in LUSC. Heliyon. 2023;9(4):e15083. PMID: 37025908',
    '[38] Marcoux N, Gettinger SN, O\'Kane G, et al. EGFR-mutant adenocarcinomas that transform to SCLC. J Clin Oncol. 2019;37(4):278-285. PMID: 30550363',
    '[39] Huang L, Li H. PD-1/CTLA-4 dual blockade in pulmonary sqNSCLC. Front Oncol. 2026. PMID: 42078800',
    '[40] Gettinger SN, et al. Nivolumab plus ipilimumab vs nivolumab in sqNSCLC. JAMA Oncol. 2021;7(9):1368-1377. PMID: 34264316',
    '[41] Lu T, et al. CT cavity/necrosis as efficacy predictor in LUSC. Front Immunol. 2021. PMID: 34354375',
]

for ref in refs:
    p(ref, sz=10, sp=1.2)

# ============= DECLARATIONS =============
page_break()
h('声明', 1)
p('基金资助：待填写', sz=11)
p('利益冲突：作者声明无利益冲突。', sz=11)
p('作者贡献：待填写', sz=11)
p('伦理审批：不适用（综述文章）。', sz=11)
p('数据可用性声明：本综述未产生原始数据。所有分析的文献均可通过PubMed/PubMed Central公开获取。', sz=11)
p('')
p('稿件格式：参考文献采用Vancouver编号格式', sz=10)
p('生成日期：2026-06-04', sz=10)

# Save
doc.save(OUT)
print(f"Word document saved: {OUT}")
print(f"Size: {os.path.getsize(OUT)} bytes")
