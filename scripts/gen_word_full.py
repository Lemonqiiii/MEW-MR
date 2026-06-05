#!/usr/bin/env python3
"""Generate FULL English Word document from jitc_submission.md — NO content compression."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os, re

OUT = "E:/medical-review/manuscript/NRDS_LifeCourse_Review.docx"
FIG_DIR = "E:/medical-review/manuscript/figures"
SRC = "E:/medical-review/manuscript/jitc_submission.md"

doc = Document()

# Page setup
for s in doc.sections:
    s.top_margin = Cm(2.54); s.bottom_margin = Cm(2.54)
    s.left_margin = Cm(3.18); s.right_margin = Cm(3.18)

# Styles
sty = doc.styles['Normal']; sty.font.name = 'Times New Roman'; sty.font.size = Pt(12)
sty.paragraph_format.line_spacing = 2.0
for lv in [1,2,3]:
    hs = doc.styles[f'Heading {lv}']
    hs.font.name = 'Times New Roman'; hs.font.color.rgb = RGBColor(0,0,0); hs.font.bold = True
    hs.font.size = Pt(14) if lv==1 else Pt(13) if lv==2 else Pt(12)
    # Add spacing: 18pt before, 6pt after for H1; 12pt before for H2; 6pt before for H3
    if lv == 1:
        hs.paragraph_format.space_before = Pt(18)
        hs.paragraph_format.space_after = Pt(6)
    elif lv == 2:
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(4)
    else:
        hs.paragraph_format.space_before = Pt(6)
        hs.paragraph_format.space_after = Pt(2)

def P(text, bold=False, sz=12, italic=False, align=None):
    par = doc.add_paragraph(); par.paragraph_format.line_spacing = 2.0
    if align: par.alignment = align
    r = par.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(sz)
    r.bold = bold; r.italic = italic
    return par

def H(text, lv=1):
    hd = doc.add_heading(text, level=lv)
    for r in hd.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = RGBColor(0,0,0)
    return hd

def PB(): doc.add_page_break()

def insert_fig(filename, caption, width=5.8):
    fp = os.path.join(FIG_DIR, filename)
    if not os.path.exists(fp): P(f'[Missing: {filename}]', sz=10); return
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(fp, width=Inches(width))
    cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run(caption); r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.bold = True

# =====================================================================
# Read source, parse, and convert section by section
# =====================================================================
with open(SRC, 'r', encoding='utf-8') as f:
    full_text = f.read()

# Split into body + references
parts = full_text.split('## References')
body_text = parts[0]
refs_text = parts[1] if len(parts) > 1 else ''

# ============ TITLE PAGE ============
P(''); P(''); P('')
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = tp.add_run('Mechanisms of Immunotherapy Resistance in\nSquamous Cell Carcinoma of Non-Small Cell Lung Cancer')
tr.font.name = 'Times New Roman'; tr.font.size = Pt(20); tr.bold = True
P(''); P('Review Article', sz=14, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
P('')
for lb, vl in [
    ('Running title: ','Immunotherapy Resistance in Lung Squamous Cell Carcinoma'),
    ('Word count: ','~9,000 (main text); 246 (abstract)'),
    ('Figures/Tables: ','1 Figure | 2 Tables | 41 References'),
    ('Keywords: ','lung squamous cell carcinoma; immunotherapy resistance; immune checkpoint inhibitors; tumor microenvironment; T cell exhaustion; KEAP1/NRF2'),
]:
    pp = P(''); pp.paragraph_format.line_spacing = 1.5
    rl = pp.add_run(lb); rl.font.name = 'Times New Roman'; rl.font.size = Pt(12); rl.bold = True
    rv = pp.add_run(vl); rv.font.name = 'Times New Roman'; rv.font.size = Pt(12)

PB()

# ============ PARSE AND CONVERT SECTIONS ============

# Split body by ## headers
sections = re.split(r'\n(?=## )', body_text.strip())

# Known insertion points for figures
figure_inserts = {
    '1. Introduction': [
        ('before_last_para', 'Figure2_Framework.png', 'Figure 1. Three-Dimensional Framework of Immunotherapy Resistance in LUSC.'),
    ],
    '3. Tumor-Intrinsic Resistance Mechanisms': [
        ('before_transition', 'Table2_Tumor_Intrinsic.png', 'Table 1. Tumor-Intrinsic Mechanisms of ICI Resistance in LUSC.'),
    ],
    '4. Tumor Microenvironment-Mediated Resistance': [
        ('before_transition', 'Table3_TME_Mechanisms.png', 'Table 2. TME-Mediated Resistance Mechanisms and Therapeutic Interventions.'),
    ],
}

# Process each section
para_counter = 0
for sec_idx, section in enumerate(sections):
    lines = section.strip().split('\n')
    if not lines: continue

    header_line = lines[0].strip()
    header_level = header_line.count('#')
    header_text = header_line.lstrip('#').strip()

    # Skip the top-level title
    if header_text.startswith('Mechanisms of Immunotherapy'):
        continue
    if header_text.startswith('Title Page'):
        continue

    # Render heading
    if header_level >= 2:
        lv = min(header_level - 1, 3)
        # Clean section number
        cleaned = re.sub(r'^#+\s*', '', header_line)
        H(cleaned, lv)
    elif header_text == 'Abstract':
        H('Abstract', 1)
        # Read abstract text
        body_lines = [l for l in lines[1:] if l.strip() and not l.startswith('#') and not l.startswith('---')]
        abstract_text = ' '.join(body_lines)
        if abstract_text:
            P(abstract_text)
        PB()
        continue
    elif header_text == 'Key Messages':
        H('Key Messages', 1)
        body_lines = [l for l in lines[1:] if l.strip() and not l.startswith('---')]
        for bl in body_lines:
            if bl.startswith('**'):
                P(bl.replace('**',''), bold=True, sz=12)
            elif bl.startswith('- '):
                P('• ' + bl[2:])
        PB()
        continue

    # Process body: split by blank lines to preserve paragraph structure
    body_content = '\n'.join(lines[1:])  # exclude header

    # Split into paragraph blocks on blank lines
    blocks = re.split(r'\n\s*\n', body_content)

    for block in blocks:
        block = block.strip()
        if not block: continue
        if block.startswith('---'): continue
        if block.startswith('**Figure'): continue
        if block.startswith('*Manuscript'): continue

        # Check if this block is a sub-heading
        if block.startswith('#'):
            sub_text = re.sub(r'^#+\s*', '', block.split('\n')[0])
            H(sub_text, min(header_level, 3))
            continue

        # Merge lines within this paragraph block
        para_text = ' '.join(block.split('\n'))

        # Clean markdown
        cleaned = para_text
        cleaned = re.sub(r'\*\*(.+?)\*\*', r'\1', cleaned)
        cleaned = re.sub(r'\*(.+?)\*', r'\1', cleaned)
        cleaned = re.sub(r'\[(\d+(?:,\d+)*)\]', r'[\1]', cleaned)

        if cleaned.strip():
            P(cleaned)

    # Check if we need to insert a figure after this section
    section_key = header_text
    if section_key in figure_inserts:
        for insert_type, fig_file, caption in figure_inserts[section_key]:
            insert_fig(fig_file, caption)
            P('')

# ============ DECLARATIONS ============
decl_match = re.search(r'## Declarations\n(.*?)(?=\n## |\Z)', body_text, re.DOTALL)
if decl_match:
    PB()
    H('Declarations', 1)
    decl_lines = decl_match.group(1).strip().split('\n')
    for dl in decl_lines:
        dl = dl.strip()
        if dl.startswith('**') and '**' in dl[2:]:
            text = dl.replace('**', '')
            if ':' in text:
                label, value = text.split(':', 1)
                pp = P(''); rl = pp.add_run(label.strip() + ': '); rl.font.name = 'Times New Roman'; rl.font.size = Pt(11); rl.bold = True
                rv = pp.add_run(value.strip()); rv.font.name = 'Times New Roman'; rv.font.size = Pt(11)
        else:
            P(dl, sz=11)

# ============ REFERENCES ============
PB()
H('References', 1)

# Parse individual references
ref_entries = re.findall(r'(\d+)\.\s+(.+?)(?=\n\d+\.\s|\n\n|\Z)', refs_text, re.DOTALL)
for num, entry in ref_entries:
    entry = entry.strip()
    entry = re.sub(r'\n+', ' ', entry)
    entry = re.sub(r'\s+', ' ', entry)
    P(f'[{num}] {entry}', sz=9).paragraph_format.line_spacing = 1.15

# Save
doc.save(OUT)

# === SELF-CHECK ===
import re
all_text = " ".join([p.text for p in doc.paragraphs])
words = len(all_text.split())

# Check 1: Figure/table references (skip if no figures/tables in text)
figs = re.findall(r'Figure\s+(\d+)', all_text)
tabs = re.findall(r'Table\s+(\d+)', all_text)
bad_figs = [n for n in set(figs) if n != '1'] if figs else []
bad_tabs = [n for n in set(tabs) if n not in ['1','2']] if tabs else []

# Check 2: Heading spacing
h_count = sum(1 for p in doc.paragraphs if p.style.name.startswith('Heading'))
h1_ok = all(hs.paragraph_format.space_before >= Pt(12)
            for hs in [doc.styles['Heading 1']])

# Check 3: Images embedded
rels = doc.part.rels
img_count = sum(1 for rel in rels.values() if "image" in str(rel.reltype))

print(f"Saved: {OUT}")
print(f"Size: {os.path.getsize(OUT)/1024:.0f} KB")
print(f"Words: ~{words}")
print(f"Headings: {h_count} (H1 space_before >= 12pt: {'OK' if h1_ok else 'FIX'})")
print(f"Images: {img_count}")
print(f"Figure refs: {len(figs)} (bad: {bad_figs or 'none'})")
print(f"Table refs: {len(tabs)} (bad: {bad_tabs or 'none'})")

# Verify word count
all_text = " ".join([p.text for p in doc.paragraphs])
words = len(all_text.split())
print(f"Words: ~{words}")
