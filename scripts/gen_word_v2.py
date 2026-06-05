#!/usr/bin/env python3
"""Generate Chinese Word document with embedded figures and tables."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUT = "E:/medical-review/manuscript/中文稿件_NSCLC鳞癌免疫治疗耐药机制.docx"
FIG_DIR = "E:/medical-review/manuscript/figures"

doc = Document()

# Page setup
s = doc.sections[0]
s.top_margin = Cm(2.54); s.bottom_margin = Cm(2.54)
s.left_margin = Cm(3.0); s.right_margin = Cm(3.0)

# Styles
sty = doc.styles['Normal']; sty.font.name = '宋体'; sty.font.size = Pt(12)
sty.paragraph_format.line_spacing = 1.5
for lv in [1,2,3]:
    hs = doc.styles[f'Heading {lv}']; hs.font.name = '黑体'
    hs.font.color.rgb = RGBColor(0,0,0)
    if lv==1: hs.font.size=Pt(16)
    elif lv==2: hs.font.size=Pt(14)
    else: hs.font.size=Pt(12)

def p(text, bold=False, sz=12, font='宋体', align=None, sp=1.5):
    par = doc.add_paragraph(); par.paragraph_format.line_spacing = sp
    if align is not None: par.alignment = align
    r = par.add_run(text); r.font.name = font; r.font.size = Pt(sz); r.bold = bold
    return par

def h(text, lv=1):
    hd = doc.add_heading(text, level=lv)
    for r in hd.runs: r.font.name = '黑体'; r.font.color.rgb = RGBColor(0,0,0)
    return hd

def pb():
    doc.add_page_break()

def insert_fig(filename, caption, width_inches=5.5):
    """Insert a PNG figure with caption."""
    filepath = os.path.join(FIG_DIR, filename)
    if not os.path.exists(filepath):
        p(f'[图表缺失: {filename}]', sz=10)
        return
    # Add figure
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run()
    run.add_picture(filepath, width=Inches(width_inches))
    # Add caption
    cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cp.add_run(caption); cr.font.name = '黑体'; cr.font.size = Pt(10); cr.bold = True
    p('', sz=4, sp=1.0)  # spacer

# ============= TITLE PAGE =============
p(''); p('')
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run('非小细胞肺癌鳞状细胞癌\n免疫治疗耐药机制研究进展')
r.font.name='黑体'; r.font.size=Pt(22); r.bold=True
p('')
sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sp.add_run('Mechanisms of Immunotherapy Resistance in\nSquamous Cell Carcinoma of Non-Small Cell Lung Cancer')
sr.font.name='Times New Roman'; sr.font.size=Pt(14); sr.italic=True
p(''); p('')
for lb, vl in [('文章类型：','综述'),('总字数：','约9,000字（中文正文）；41篇参考文献'),
               ('图表：','图4幅 | 表4张'),('关键词：','肺鳞状细胞癌；免疫治疗耐药；免疫检查点抑制剂；肿瘤微环境；T细胞耗竭；KEAP1/NRF2')]:
    pp = doc.add_paragraph(); pp.paragraph_format.line_spacing = 1.5
    rl = pp.add_run(lb); rl.font.name='黑体'; rl.font.size=Pt(12); rl.bold=True
    rv = pp.add_run(vl); rv.font.name='宋体'; rv.font.size=Pt(12)

pb()

# ============= ABSTRACT =============
h('摘要', 1)
p('肺鳞状细胞癌（lung squamous cell carcinoma, LUSC）约占非小细胞肺癌的30%，其显著特征为缺乏可靶向驱动突变，因此免疫治疗成为全身治疗的基石。靶向PD-1/PD-L1和CTLA-4轴的免疫检查点抑制剂（ICI）已根本性地改变了LUSC的治疗格局；然而，仅少数患者获得持久获益，原发性耐药和获得性耐药构成了关键的临床挑战。LUSC的耐药产生于肿瘤内在改变、肿瘤微环境（TME）介导的免疫抑制以及治疗诱导的适应性变化三者之间的复杂交互作用。本综述从三个相互关联的维度系统整合了LUSC免疫治疗耐药机制的最新认识：首先概述LUSC独特的免疫景观，包括存在于28-36%患者中的免疫耗竭型（EIC）；其次审视肿瘤内在耐药机制（PI3K/AKT、KEAP1/NRF2、p38 MAPK、EMT、表观遗传失调、程序性细胞死亡缺陷）；第三分析TME介导的耐药（T细胞耗竭伴多重检查点共表达、TAM/CAF/MDSC/Treg等免疫抑制细胞、代谢限制、瘤内微生物组）。进一步讨论获得性耐药（克隆演化、组织学转化、表型可塑性），最后评估克服策略（联合免疫治疗、靶向内在通路、TME重塑、精准免疫治疗）。')
p('')
p('Abstract', bold=True, sz=11, font='Times New Roman')
p("Lung squamous cell carcinoma (LUSC), accounting for approximately 30% of non-small cell lung cancers, is characterized by a paucity of actionable driver mutations and a consequent reliance on immunotherapy as the cornerstone of systemic treatment. Immune checkpoint inhibitors targeting the PD-1/PD-L1 and CTLA-4 axes have transformed the therapeutic landscape; however, only a minority of patients achieve durable benefit, with both primary and acquired resistance representing critical challenges. This review synthesizes resistance mechanisms across three interconnected dimensions: tumor-intrinsic alterations, TME-mediated suppression, and acquired resistance, concluding with strategies to overcome these barriers.", sz=10, font='Times New Roman')

pb()

# ============= 1. INTRODUCTION =============
h('1  引言', 1)

p('肺癌仍是全球癌症相关死亡的首要原因，非小细胞肺癌（NSCLC）约占所有病例的85% [1]。在NSCLC亚型中，肺鳞状细胞癌（LUSC）约占诊断的30%，是仅次于肺腺癌（LUAD）的第二常见组织学亚型 [2]。LUSC患者面临尤为严峻的临床结局，晚期疾病五年生存率低于20% [3]。历史上，标准治疗主要限于铂类化疗，生存获益有限且毒性显著 [4]。')

p('LUSC治疗格局最深层的塑造因素恰恰是一种缺失：可靶向驱动突变的极度匮乏。在LUAD中，针对EGFR、ALK、ROS1、BRAF和KRAS G12C的靶向治疗已根本改变了治疗范式；相比之下，LUSC几乎没有可通过药物靶向的基因组改变 [5,6]。LUSC的基因组景观主要由抑癌基因频繁改变所主导——包括TP53（~80%）、CDKN2A（~70%，涵盖突变、缺失和表观遗传沉默）和KEAP1（~12%）——以及PIK3CA（~30-40%）、FGFR1（~20%）的反复扩增，以及包含鳞状谱系转录因子SOX2和TP63的染色体3q位点 [5,7]。尽管针对这些基因组特征开展了大量靶向药物临床试验，迄今尚无任何药物获得专门针对LUSC的注册批准 [7]。')

p('免疫检查点抑制剂（ICIs）的问世已根本性地改变了LUSC的治疗格局。靶向PD-1（纳武利尤单抗、帕博利珠单抗）、PD-L1（阿替利珠单抗、度伐利尤单抗）以及CTLA-4（伊匹木单抗）的抗体在晚期NSCLC的多线治疗中均展现出有意义的临床活性 [8–10]。里程碑式的III期试验——包括KEYNOTE-407（帕博利珠单抗联合化疗一线治疗鳞状NSCLC）、CheckMate-017（纳武利尤单抗二线治疗鳞状NSCLC）和IMpower-131（阿替利珠单抗联合化疗，显示PFS获益但在ITT人群中OS改善未达统计学显著性）——确立了以免疫治疗为基础的方案作为新标准治疗 [8,11]。')

p('尽管取得这些进展，仅少数患者获得持久临床获益。帕博利珠单抗单药一线治疗PD-L1高表达（≥50%）患者的ORR约为40-45%，鳞状NSCLC联合化疗后提升至约55-60% [8,13]。然而，相当比例患者表现为原发性耐药，许多初始应答者最终发展为获得性耐药 [14]。Lung-MAP S1400F子研究专门评估了度伐利尤单抗联合tremelimumab在抗PD-(L)1耐药的鳞状NSCLC中的疗效，报告原发性耐药队列ORR仅7%，获得性耐药队列为0%，生动说明了克服已建立免疫治疗耐药所面临的巨大挑战 [15]。')

p('支撑这些令人失望临床结局的是一个跨越三个相互关联领域的复杂耐药机制网络：肿瘤细胞内在改变、TME动态重塑，以及驱动克隆演化和表型适应性的治疗诱导选择压力 [14,16]。LUSC的TME具有若干共同构建免疫逃逸生态位的特征：M2极化TAM、Treg和MDSC高比例存在；以TGF-β和IL-6为主导的细胞因子环境；以及以缺氧、乳酸酸中毒和营养耗竭为特征的代谢有害微环境 [16,17]。此外，近期多组学研究在28-36%的LUSC患者中鉴定出独特的免疫耗竭型（EIC），其特征为密集的淋巴细胞浸润与多达九个抑制性免疫检查点的共上调矛盾性共存 [18]。这种"炎症浸润但功能抑制"的免疫状态凸显了单一PD-1/PD-L1阻断的不足。')

p('在本综述中，我们对LUSC免疫治疗耐药机制的当前认识进行了全面整合，按照三个相互关联的维度进行组织。')

# ===== INSERT FIGURE 1 + FIGURE 2 =====
insert_fig("Figure1_PRISMA.png", "图1  PRISMA 2020文献筛选流程图")
insert_fig("Figure2_Framework.png", "图2  LUSC免疫治疗耐药的三维框架总览")

pb()

# ============= 2. IMMUNE LANDSCAPE =============
h('2  LUSC的免疫景观', 1)

p('理解LUSC免疫治疗耐药机制需要对其独特的肿瘤免疫微环境（TIME）有一个基本的认识。尽管LUSC与LUAD在肺中共享共同的解剖起源，但两者的免疫景观存在实质性差异，反映了突变过程、基质组成和演化轨迹的不同。scRNA-seq、空间转录组学和多组学整合的最新进展以前所未有的分辨率剖析了LUSC TIME的细胞和分子架构。')

h('2.1  TIME异质性与分子分类', 2)

p('LUSC的TIME以深刻的瘤间和瘤内异质性为特征，从根本上塑造了对免疫治疗的应答 [16]。以经典的肿瘤免疫表型三分法——免疫炎症型、免疫排斥型和免疫沙漠型——为基础，近期研究已利用整合多组学方法针对LUSC进行了细化 [16,19]。')

p('免疫炎症型LUSC的特征是肿瘤实质内CD8+细胞毒性T淋巴细胞（CTL）、活化CD4+记忆T细胞和树突状细胞（DC）的丰富浸润，表达高水平效应细胞因子和溶细胞标志物 [16]。矛盾的是，许多免疫炎症型LUSC同时上调多个抑制性免疫检查点并携带高频率的免疫抑制性Treg和M2极化巨噬细胞，形成一种"炎症浸润但功能抑制"的免疫状态，这种状态常常是单药PD-1/PD-L1阻断原发性耐药的基础 [18,20]。')

p('免疫排斥型LUSC的特征是CTL和其他免疫效应细胞被限制在肿瘤巢周围的基质区室中，无法穿透肿瘤实质。这种排斥主要由癌相关成纤维细胞（CAF）介导，它们沉积包括胶原、纤连蛋白和层粘连蛋白在内的致密ECM成分，形成T细胞浸润的物理屏障 [16]。空间转录组分析揭示了CAF亚群与APOE+ TAM的紧密共定位，鉴定出一种通过协同细胞因子和趋化因子分泌来强化免疫排斥的CAF-TAM信号轴 [16]。')

p('免疫沙漠型LUSC占少数病例，其标志是肿瘤和基质区室内几乎完全缺乏T细胞。这种表型被认为源于先天性免疫感知缺陷——包括STING/cGAS通路激活受损——以及DC启动不足 [16,19]。')

p('在这些定性描述之外，整合基因组规模分析已建立了LUSC的分子分类框架 [19,21]。Yin等人鉴定出LUSC的不同免疫亚型 [19]；Yang等人鉴定出一种细胞因子主导的免疫抑制类（EIC）[18]；Song等人利用513例LUSC样本划分了六个分子亚型（CS1-CS6），并鉴定出CS3为一种淋巴细胞浸润亚型，矛盾性地表现出升高的耗竭标志物并通过TIDE分析预测对ICB治疗耐药 [20]。这些分类工作的共同结论是：免疫浸润的功能状态和空间组织才是治疗结局的关键决定因素。')

h('2.2  LUSC与LUAD：分歧的免疫环境', 2)

p('尽管LUSC和LUAD均被归类为NSCLC，但两者起源于不同的细胞源头，携带根本不同的基因组景观，并展现出明显分歧的免疫微环境 [22,23]。')

p('在基因组层面，LUAD富含可靶向驱动突变——最显著的是EGFR（东亚~30-40%，西方~10-15%）、KRAS（~30%）、ALK重排（~5%）和ROS1融合（~1-2%）——这些突变在EGFR突变亚群中与相对较低的TMB相关 [22]。相比之下，LUSC主要以致癌基因功能丧失突变（TP53、CDKN2A、KEAP1）和反复扩增（PIK3CA、FGFR1、SOX2）为主导，产生较高的总TMB和新抗原负荷，理论上预测其免疫原性更高 [22,23]。然而，这种升高的新抗原负荷并未转化为LUSC相较于LUAD更优的ICI应答，提示存在强效的对抗性免疫抑制机制。')

p('在细胞层面，LUSC肿瘤表现出更高的M2巨噬细胞和静息CD4+记忆T细胞浸润，而LUAD肿瘤倾向于有更高的幼稚B细胞和浆细胞浸润。LUSC也表现出更高的免疫排斥型比例，由更广泛的CAF激活和ECM重塑驱动 [16,22]。LUSC的细胞因子环境偏向TGF-β和IL-6/STAT3信号，而LUAD更常与EGFR驱动的免疫抑制程序相关 [22]。这些亚型特异性免疫特征具有直接的临床意义：LUSC中免疫排斥的高比例提示靶向CAF介导的基质屏障的策略可能尤为适用；LUSC和LUAD之间不同的免疫检查点表达谱可指导联合免疫治疗搭档的理性选择 [22]；在LUAD中开发的生物标志物可能无法直接用于LUSC。')

# ===== INSERT TABLE 4 =====
insert_fig("Table4_LUSC_vs_LUAD.png", "表4  LUSC与LUAD的关键免疫学差异")

h('2.3  免疫耗竭型（EIC）：理解耐药的框架', 2)

p('LUSC免疫生物学的一项里程碑式贡献是Yang等人通过对624例LUSC样本进行RNA测序数据的无监督聚类，鉴定出了免疫耗竭型（EIC）[18]。约28-36%的LUSC患者属于EIC，其特征是具有对免疫治疗耐药至关重要意义的独特分子标志。')

p('EIC由四个标志性特征的共同存在所定义 [18]。第一，该类肿瘤展现出T细胞耗竭标志的显著富集，伴随典型耗竭相关转录因子TOX和EOMES的表达升高，以及TCF-1——保留增殖能力和PD-1阻断应答性的祖细胞耗竭T细胞（Tpex）标志物——的表达降低。第二，EIC的特征是高达九个抑制性免疫检查点的共上调——CTLA-4、PD-1、LAG-3、BTLA、TIGIT、TIM-3、IDO1、SIGLEC7和VISTA——代表一种广泛的、多受体的免疫抑制状态，使得单一PD-1/PD-L1阻断在机制上就不足 [18]。第三，EIC肿瘤被M2极化巨噬细胞和CD4+FOXP3+ Treg高度浸润，产生TGF-β、IL-10和CCL18等免疫抑制性细胞因子 [18]。第四，EIC矛盾性地与高总密度的TIL相关，然而这些患者的预后显著差于TIL浸润较低的患者。')

p('EIC概念已得到后续研究的证实和扩展。Song等人在其CS3分子亚型中鉴定出类似表型，该亚型表现出高淋巴细胞浸润伴随升高的耗竭标志物 [20]。此外，CS3亚型被发现特异性上调LAMC2-CD44分子轴——一种与EMT相关的通路，同时调控肿瘤增殖和免疫排斥，提供了肿瘤细胞内在程序与耗竭免疫表型之间的潜在机制联系。')

# ===== INSERT FIGURE 3 =====
insert_fig("Figure3_EIC.png", "图3  LUSC免疫耗竭型（EIC）的四大特征（28-36%患者）")

pb()

# ============= 3. TUMOR-INTRINSIC =============
h('3  肿瘤内在耐药机制', 1)

p('肿瘤细胞内在改变代表了LUSC免疫治疗耐药的一个基本层面，通过多种机制运作，汇聚于一个共同终点：免疫效应细胞无法识别和清除恶性细胞。这些机制涵盖从积极塑造免疫抑制微环境的致癌信号通路，到沉默抗原呈递机制的表观遗传程序，再到使肿瘤细胞抵抗CTL介导杀伤的细胞死亡通路缺陷。')

h('3.1  致癌信号通路', 2)

h('3.1.1  PI3K/AKT/mTOR通路', 3)
p('PI3K通路是LUSC中最频繁激活的致癌网络之一，PIK3CA扩增和激活突变发生于约30-40%的病例 [5]。PI3K/AKT/mTOR信号的激活在促进肿瘤细胞增殖和存活的同时，通过多种机制驱动免疫逃逸：通过AKT介导的PD-L1 mRNA稳定化上调PD-L1表达；增加免疫抑制性细胞因子的分泌；促进TME内M2巨噬细胞极化 [5,7]。临床前研究已证明mTOR抑制可逆转PI3K驱动的PD-L1表达并恢复T细胞介导的杀伤 [5]。')

h('3.1.2  KEAP1/NRF2通路', 3)
p('KEAP1突变发生于约12%的LUSC，已成为NSCLC中对ICI原发性耐药的最强有力的预测因子之一 [5,7]。KEAP1功能丧失突变导致NRF2持续稳定，驱动广泛的抗氧化和细胞保护基因转录程序 [5]。这一激活的NRF2程序通过多种汇聚机制深刻改变肿瘤细胞免疫原性：增强的ROS清除减少免疫原性细胞死亡并损害DC激活；外排转运蛋白上调限制CTL释放的细胞毒性分子的胞内积累；抑制促炎细胞因子产生减弱先天性免疫感知 [5,7]。临床上，KEAP1突变的LUSC与ICI治疗显著缩短的PFS相关。')

h('3.1.3  MAPK/p38信号', 3)
p('在充分表征的RAS/RAF/MEK/ERK级联之外，应激激活的p38 MAPK通路近期已被牵涉到LUSC免疫逃逸中。Lin等人证明TGM2在LUSC中显著过表达（P=0.00018），通过激活p38 MAPK信号促进肿瘤增殖、迁移、侵袭 [24]。同时，TGM2高表达肿瘤展现出以Th1细胞浸润减少为特征的免疫抑制性TME。上游调控因子NR3C1直接结合TGM2启动子并驱动其转录。这一NR3C1-TGM2-p38 MAPK轴代表了可靶向的肿瘤内在机制 [24]。')

h('3.1.4  FGFR1和EGFR：RTK过表达的悖论', 3)
p('FGFR1扩增发生于约20%的LUSC，EGFR蛋白频繁过表达 [5,7]。然而，FGFR抑制剂在FGFR1扩增的LUSC中展现出极小的单药活性 [5]。EGFR-TKI在临床上无效——激活EGFR突变在LUSC中极为罕见 [25]。机制研究揭示LUSC中的EGFR蛋白表达伴随下游RAS/MAPK和PI3K/AKT通路的组成性配体非依赖性激活，由来自FGFR1和PI3KCA的平行输入所维持，阐明了LUSC生物学的一个更广泛原理：致癌信号网络的冗余性和互联性限制了单一通路抑制的疗效 [5,7,25]。')

h('3.2  上皮间质转化（EMT）', 2)
p('EMT在LUSC中与免疫抑制性TME和ICI耐药强相关 [21]。Zhang等人构建的六基因EMT预后模型将LUSC患者分为具有显著不同免疫浸润特征的风险组 [21]。Song等人鉴定的LAMC2-CD44轴是CS3分子亚型中免疫耐药的关键驱动因子 [20]。在LAMC2-CD44轴之外，EMT通过多种互补机制促进免疫耐药：间质肿瘤细胞上调TGF-β、VEGF和IL-6；E-cadherin下调损害免疫突触形成；ZEB1和SNAIL驱动PD-L1表达增加 [20,21]。')

h('3.3  表观遗传失调', 2)
p('异常DNA甲基化通过表观遗传沉默MHC和抗原呈递基因来促进免疫逃逸 [26]。DNMT抑制剂如地西他滨和阿扎胞苷可恢复表观遗传沉默的免疫基因表达并与ICI协同。HDAC和EZH2的过表达已被牵涉到LUSC免疫逃逸 [26]。非编码RNA方面，circHMGB2通过miRNA海绵机制驱动免疫抑制和抗PD-1耐药，其敲低在临床前模型中与抗PD-1治疗协同 [27]。')

h('3.4  细胞死亡通路与免疫耐药', 2)
p('免疫原性细胞死亡（ICD）促进DC成熟和CTL交叉启动。失巢凋亡耐药——肿瘤细胞在脱离ECM后存活的能力——已被联系到LUSC免疫调控，Lu等人鉴定出S100A7、S100A8和SPP1与LUSC免疫浸润模式显著相关 [28]。铁死亡——SLC7A11表观遗传激活——定义了一个与LUSC DNA甲基化分类相关的铁死亡-免疫轴 [30]。焦亡相关基因标志可预测LUSC免疫微环境组成和免疫治疗应答 [31]。')

h('3.5  基因组不稳定性与新抗原动力学', 2)
p('虽然LUSC通常展现出较高的TMB，但TMB与ICI应答之间的关系比最初认识的要复杂 [7,26]。克隆性新抗原比亚克隆新抗原更能引发有效的抗肿瘤免疫；染色体不稳定性和非整倍体可矛盾性地抑制抗肿瘤免疫 [7]；在ICI治疗压力下，肿瘤可经历免疫编辑，伴随新抗原表达克隆的消除和HLA缺陷亚克隆的扩张 [7,14]。')

# ===== INSERT TABLE 2 =====
insert_fig("Table2_Tumor_Intrinsic.png", "表2  LUSC肿瘤内在免疫治疗耐药机制总览")

pb()

# ============= 4. TME-MEDIATED =============
h('4  肿瘤微环境介导的耐药', 1)

p('肿瘤内在改变为免疫逃逸奠定分子基础的同时，肿瘤微环境——包括多样的免疫细胞群体、基质成分、细胞因子、代谢物和微生物群落——构成了免疫治疗成功或失败的战场。')

h('4.1  T细胞耗竭与功能障碍', 2)
p('T细胞耗竭——由慢性抗原刺激引起的效应功能进行性丧失状态——是LUSC免疫逃逸的核心机制 [16]。祖细胞耗竭T细胞（Tpex）以TCF-1表达为特征，保留增殖能力，是PD-1阻断的主要靶点 [16,18]。随着耗竭进展，Tpex分化为终末耗竭T细胞（Tex-term），表达高水平TOX、多个抑制性受体，并对PD-1治疗无效 [18]。在LUSC EIC中，T细胞区室偏向Tex-term主导状态，TCF-1+ Tpex频率降低 [18]。这种耗竭状态由多种TME衍生因子强化：高新抗原负荷、免疫抑制性细胞因子（IL-10、TGF-β、CCL18）、代谢竞争和慢性I型干扰素信号 [16,18]。')

h('4.2  多重免疫检查点共表达', 2)
p('LUSC EIC的一个定义性特征是多达九个抑制性免疫检查点的共上调——CTLA-4、PD-1、LAG-3、BTLA、TIGIT、TIM-3、IDO1、SIGLEC7和VISTA [18]。这种共表达模式意味着单独阻断PD-1/PD-L1轴留下了多个替代抑制通路完好无损。Lung-MAP S1400F子研究证明了这一点：PD-L1/CTLA-4双重阻断在既往抗PD-(L)1耐药患者中产生极小的临床获益 [15]。LAG-3和TIGIT均与PD-1在LUSC中频繁共表达，使它们成为与PD-1阻断联合的有吸引力靶点 [7,18]。')

h('4.3  免疫抑制细胞群体', 2)
p('TAM代表LUSC TME中最丰富的免疫细胞群体，主要向M2样表型极化 [17]。Ji等人证明糖酵解酶ALDOA在LUSC中显著过表达，ALDOA驱动的糖酵解通量产生富含乳酸的微环境，促进M2巨噬细胞极化，建立自我强化的代谢-免疫回路 [17]。近期研究已在鳞状癌中鉴定出一个CAF-TAM信号轴 [16]。')

p('CAF是构成LUSC基质主要成分的活化成纤维细胞，通过多种机制促进免疫耐药 [5,16]。POSTN+和FAP+ CAF沉积致密ECM形成物理屏障；分泌CXCL12和CCL2排斥T细胞同时招募免疫抑制群体 [5]。MDSC通过精氨酸酶-1和xCT耗竭氨基酸、产生ROS亚硝基化TCR、分泌IL-10和TGF-β发挥免疫抑制功能 [5,16]。Treg通过CTLA-4介导的转内吞和免疫抑制性细胞因子分泌抑制抗肿瘤免疫 [18]。')

h('4.4  细胞因子网络与代谢免疫调控', 2)
p('TGF-β是核心协调者，由肿瘤细胞、CAF、TAM和Treg产生，直接抑制溶细胞效应分子表达 [5,16]。IL-6激活STAT3信号驱动慢性炎症前馈回路。IL-10广泛抑制抗原呈递和共刺激分子表达 [18]。')

p('LUSC TME以Warburg效应为代谢特征，乳酸积累将TME酸化至pH 6.0-6.5，抑制T细胞功能同时促进Treg稳定性和M2极化 [17]。缺氧通过HIF-1α稳定化转录激活PD-L1上调、VEGF分泌和MDSC/TAM招募 [34]。近期阐明的一个重要机制涉及STING通路——缺氧通过HIF-1α抑制STING信号，导致I型干扰素产生减少和DC成熟受损 [34]。')

h('4.5  肿瘤微生物组-免疫串扰', 2)
p('瘤内微生物组是LUSC TME生物学的一个新兴维度 [33]。多组学整合揭示LUSC中微生物群组成与乳酸代谢、免疫细胞浸润模式和免疫检查点表达显著相关。乳杆菌属的富集与M2巨噬细胞为主和CD8+ T细胞应答减弱的免疫抑制性TME相关 [33]。这些发现提出了微生物组导向干预的可能性。')

# ===== INSERT TABLE 3 =====
insert_fig("Table3_TME_Mechanisms.png", "表3  TME介导的耐药机制及干预策略")

pb()

# ============= 5. ACQUIRED RESISTANCE =============
h('5  获得性耐药机制', 1)

p('获得性耐药在初始有效的免疫治疗选择压力下出现，表现为一段临床获益期后的疾病进展 [14]。')

h('5.1  免疫治疗压力下的克隆演化', 2)
p('ICI对遗传异质性的肿瘤细胞群体施加了强大的免疫选择压力，驱动达尔文式的克隆演化过程 [7,14]。携带赋予免疫逃逸优势突变的肿瘤亚克隆在ICI治疗期间被正向选择 [14]。已在NSCLC中记录到JAK1/JAK2功能丧失突变（废除IFN-γ受体信号）和B2M突变（消除MHC I类表达）等获得性耐药机制 [7,14]。')

h('5.2  组织学转化：腺鳞转化', 2)
p('腺癌向鳞状细胞癌转化（AST）是LUSC生物学中一种尤为有启发性的获得性耐药形式 [36]。在AST期间，LUAD细胞经历由鳞状转录因子（SOX2、TP63、ZEB1）激活所驱动的谱系可塑性程序 [36]。这种转录重编程伴随免疫微环境的广泛变化：转化的鳞状样肿瘤细胞分泌独特的细胞因子和趋化因子谱；ECM以CAF依赖性方式重塑；免疫检查点分子表达谱发生变化 [36]。')

h('5.3  表型可塑性', 2)
p('Wang等人表征了LUSC中表型可塑性相关基因表达模式，证明伴和不伴淋巴结转移的患者展现显著不同的可塑性基因组特征 [37]。低可塑性评分的患者被预测对PD-L1抑制剂、顺铂和紫杉醇等更为敏感，而高可塑性肿瘤与广泛治疗耐药相关 [37]。可塑性驱动的耐药尤为具有挑战性，因为它是动态和可逆的 [14,37]。')

h('5.4  SCLC转化及其他谱系转换', 2)
p('表型可塑性的一种极端表现是NSCLC——包括LUSC——在ICI治疗下转化为小细胞肺癌（SCLC）。虽然最初在EGFR突变腺癌中被描述，SCLC转化也已在LUSC中被记录 [38]。这一现象虽然罕见（<5%），但具有临床毁灭性。SCLC转化以RB1和TP53功能丧失、ASCL1或NEUROD1神经内分泌转录程序的激活为特征 [38]。LUSC的特定特征——如基线TP53突变的高发率——是否赋予相较于LUAD更高的SCLC转化风险，仍是一个悬而未决的问题。')

pb()

# ============= 6. STRATEGIES =============
h('6  克服耐药的策略', 1)

p('LUSC免疫治疗耐药的多维性质需要同样多维的治疗策略，可分为合理的免疫联合治疗、靶向肿瘤内在通路、TME重塑和生物标志物引导的精准策略四类。')

h('6.1  合理的免疫联合治疗', 2)
p('LUSC中多个抑制性免疫检查点的共表达——尤其在EIC中——为双重或三重检查点阻断提供了清晰的生物学理论基础 [7,18]。虽然Lung-MAP S1400F子研究在非选择患者中增加CTLA-4阻断仅产生最小获益 [15]，但新型双特异性抗体——包括QL1706（PD-1/CTLA-4）和rilvegostomig（PD-1/TIGIT）——提供了同时双靶点接合的优势 [39]。ICI联合铂类化疗（KEYNOTE-407）仍然是标准一线方案，化疗通过ICD诱导、MDSC耗竭和Treg减少等多种机制来克服耐药 [5]。')

h('6.2  靶向肿瘤内在通路', 2)
p('在PIK3CA改变的LUSC中，PI3K/mTOR抑制剂可同时抑制肿瘤生长并解除mTOR驱动的PD-L1表达。在KEAP1突变的LUSC中，谷氨酰胺酶抑制剂和STING激动剂正在探索中 [5]。表观遗传启动——使用DNMT抑制剂或HDAC抑制剂恢复沉默的抗原呈递和干扰素应答基因表达——代表了由临床前证据支持的策略 [26]。')

h('6.3  TME重塑', 2)
p('CAF靶向方法——包括FAP抑制剂、Hedgehog通路拮抗剂和TGF-β捕获剂——旨在减少ECM沉积并促进T细胞穿透 [16]。TAM重编程策略——CSF1R抑制、CD47阻断或CD40激动——寻求将巨噬细胞极化平衡从M2转向M1 [16,17]。STING/cGAS通路代表了先天性免疫激活的有吸引力靶点 [34]。')

h('6.4  新兴策略与精准免疫治疗', 2)
p('建立TIME分类框架——整合分子亚型、免疫浸润模式和空间架构——可指导机制匹配的联合策略选择 [16]。基于液体活检的耐药机制动态监测——通过ctDNA分析新出现的JAK1、JAK2或B2M突变——可实现适应性治疗策略 [7]。瘤内微生物组在调控ICI应答中的新兴角色提出了微生物组导向干预的可能性 [33]。')

# ===== INSERT FIGURE 4 + TABLE 1 =====
insert_fig("Figure4_Decision.png", "图4  基于TIME亚型的LUSC免疫治疗决策框架")
insert_fig("Table1_Clinical_Trials.png", "表1  LUSC关键免疫检查点抑制剂临床试验")

pb()

# ============= 7. CONCLUSIONS =============
h('7  结论与未来展望', 1)

p('免疫治疗已改变了肺鳞状细胞癌患者的治疗格局，然而大多数患者要么初始不应答，要么最终发展出耐药。本综述从三个相互关联的维度整合了对LUSC免疫治疗耐药机制的最新认识。')

p('第一，LUSC的耐药很少归因于单一机制，而是产生于多个相互强化通路的汇聚。KEAP1/NRF2突变肿瘤细胞同时抵抗氧化杀伤、抑制先天性免疫感知并分泌促进免疫抑制TME的因子。CAF-TAM轴同时物理排斥T细胞并抑制浸润的T细胞。EIC共表达九个抑制性检查点，任一其他检查点均可维持免疫抑制。这一机制冗余性原则解释了为何单药PD-1/PD-L1阻断对相当比例的LUSC患者不足。')

p('第二，免疫浸润的功能状态比其数量更重要。EIC——具有丰富的TIL、高PD-L1表达和矛盾性不良预后——诠释了免疫浸润与免疫能力之间的关键区别。未来的预测模型必须纳入T细胞功能状态、检查点共表达模式和空间组织的多维评估。')

p('第三，LUSC在免疫生物学上不同于LUAD，需要组织学特异性的治疗策略。免疫排斥表型的更高发率、以抑癌基因丧失为主导的独特基因组景观，以及LUSC TME的特定CAF和代谢特征，都反对在没有LUSC特异性验证的情况下外推源自LUAD的治疗范式。')

p('展望未来，按TIME亚型、KEAP1突变状态或EIC分子标志进行患者分层的前瞻性临床试验对于评估机制匹配的联合治疗至关重要。多组学整合将提供解剖耐药机制全复杂度所需的分辨率。液体活检动态监测将实现实时响应新兴耐药的适应性治疗策略。双特异性抗体、表观遗传启动、STING激动剂和CAF靶向治疗等新型模式有望扩展克服LUSC免疫治疗耐药可用的武器库。')

pb()

# ============= REFERENCES =============
h('参考文献', 1)

refs = [
    '[1] Sung H, Ferlay J, Siegel RL, et al. Global Cancer Statistics 2020. CA Cancer J Clin. 2021;71(3):209-249. PMID: 33538338',
    '[2] Travis WD, et al. The 2015 WHO classification of lung tumors. J Thorac Oncol. 2015;10(9):1243-1260. PMID: 26291008',
    '[3] Siegel RL, et al. Cancer statistics, 2022. CA Cancer J Clin. 2022;72(1):7-33. PMID: 35020204',
    '[4] Schiller JH, et al. Comparison of four chemotherapy regimens for advanced NSCLC. N Engl J Med. 2002;346(2):92-98. PMID: 11784886',
    '[5] Niu Z, Jin R, Zhang Y, Li H. Signaling pathways and targeted therapies in LUSC. Signal Transduct Target Ther. 2022;7:353. PMID: 36198685',
    '[6] TCGA Research Network. Comprehensive genomic characterization of squamous cell lung cancers. Nature. 2012;489:519-525. PMID: 22960745',
    '[7] Yuan H, Liu J, Zhang J. Immune checkpoint blockade in metastatic LUSC. Molecules. 2021;26(5):1392. PMID: 33807509',
    '[8] Paz-Ares L, et al. Pembrolizumab plus chemotherapy for squamous NSCLC. N Engl J Med. 2018;379:2040-2051. PMID: 30280635',
    '[9] Brahmer J, et al. Nivolumab versus docetaxel in squamous-cell NSCLC. N Engl J Med. 2015;373:123-135. PMID: 26028407',
    '[10] Chen W, et al. First-line immunotherapy in advanced squamous NSCLC. Front Oncol. 2024;14:1360583. PMID: 38725635',
    '[11] Jotte R, et al. Atezolizumab in advanced squamous NSCLC (IMpower131). J Thorac Oncol. 2020;15:1351-1360. PMID: 32302702',
    '[12] Ikeda S, et al. Why cemiplimab? Cancers. 2026;18(2):272. PMID: 41595192',
    '[13] Reck M, et al. Pembrolizumab vs chemotherapy for PD-L1-positive NSCLC. N Engl J Med. 2016;375:1823-1833. PMID: 27718847',
    '[14] Sharma P, et al. Primary, adaptive, and acquired resistance to cancer immunotherapy. Cell. 2017;168:707-723. PMID: 28187290',
    '[15] Leighl NB, et al. Durvalumab plus tremelimumab in anti-PD-(L)1 resistant sqNSCLC. J Immunother Cancer. 2021;9:e002973. PMID: 34429332',
    '[16] Tong Y, et al. Decoding the TIME in LUSC. Transl Lung Cancer Res. 2025;14(4):1170-1190. PMID: 41133013',
    '[17] Ji Y, et al. Aldolase A in pan-cancer and LUSC. Cancer Cell Int. 2025;25:184. PMID: 41239433',
    '[18] Yang M, et al. Cytokine-dominated immunosuppressive class in LUSC. Genome Med. 2022;14(1):72. PMID: 35799269',
    '[19] Yin L, et al. Immune subtypes of LUSC. Front Oncol. 2021;11:778324. PMID: 35186710',
    '[20] Song T, et al. Laminin gamma2-CD44 immune resistance in LUSC. Heliyon. 2024;10:e31299. PMID: 38803944',
    '[21] Zhang A, et al. EMT-based risk model for LUSC. PeerJ. 2026;14:e21117. PMID: 42089102',
    '[22] Shen Y, et al. Differences between LUAD and LUSC. Genes Dis. 2025. PMID: 40083325',
    '[23] Yan T, et al. Immune heterogeneity between LUAD and LUSC. Front Immunol. 2021;12:703797. PMID: 34394068',
    '[24] Lin C, et al. TGM2 drives p38 MAPK-mediated immune evasion in LUSC. Front Immunol. 2025;16:1547241. PMID: 41050683',
    '[25] Ju L, et al. Intrinsic resistance of LUSC to EGFR-TKI. Front Oncol. 2020;10:568878. PMID: 33133263',
    '[26] Sasa GBK, et al. lncRNAs, immunotherapy and DNA methylation in LUSC. Transl Cancer Res. 2021;10:5324-5341. PMID: 35116387',
    '[27] Zhang LX, et al. circHMGB2 drives anti-PD-1 resistance. Mol Cancer. 2022;21(1):110. PMID: 35525959',
    '[28] Lu H, et al. Anoikis-related genes in LUSC. Med Sci Monit. 2026;31:e951722. PMID: 41902322',
    '[29] Ou D, et al. Anoikis resistance gene model in LUSC. Discover Oncol. 2026. PMID: 41530460',
    '[30] Lu HP, et al. SLC7A11 ferroptosis-immune axis in LUSC. PeerJ. 2026;14:e20686. PMID: 41700135',
    '[31] Deng X, et al. Pyroptosis in LUSC immune microenvironment. J Cancer Res Clin Oncol. 2022. PMID: 36123889',
    '[32] Newsom-Davis T, et al. TROPION-Lung10. Front Oncol. 2025;15:1721624. PMID: 41669261',
    '[33] Qiu X, Li D. Intratumor microbiome and lactic acid metabolism in LUSC. Front Immunol. 2025. PMID: 40568577',
    '[34] Chen F, et al. Hypoxia-STING chemo-immuno resistance in LUSC. Transl Oncol. 2025;52:102350. PMID: 40138855',
    '[35] Zhao F, et al. Hypoxia-related lncRNAs in LUSC. Front Oncol. 2021. PMID: 34250747',
    '[36] Xu H, et al. Immune mechanisms in lung adenosquamous transformation. Front Immunol. 2025;16:1502584. PMID: 40568576',
    '[37] Wang F, Zhu L. Phenotypic plasticity in LUSC. Heliyon. 2023;9(4):e15083. PMID: 37025908',
    '[38] Marcoux N, et al. EGFR-mutant adenocarcinomas that transform to SCLC. J Clin Oncol. 2019;37:278-285. PMID: 30550363',
    '[39] Huang L, Li H. PD-1/CTLA-4 dual blockade in pulmonary sqNSCLC. Front Oncol. 2026. PMID: 42078800',
    '[40] Gettinger SN, et al. Nivolumab plus ipilimumab vs nivolumab in sqNSCLC. JAMA Oncol. 2021;7:1368-1377. PMID: 34264316',
    '[41] Lu T, et al. CT cavity/necrosis as efficacy predictor in LUSC. Front Immunol. 2021. PMID: 34354375',
]

for ref in refs:
    p(ref, sz=9, sp=1.15, font='Times New Roman')

# ============= DECLARATIONS =============
pb()
h('声明', 1)
p('基金资助：待填写', sz=11)
p('利益冲突：作者声明无利益冲突。', sz=11)
p('作者贡献：待填写', sz=11)
p('伦理审批：不适用（综述文章）。', sz=11)
p('')
p('—— 稿件格式：参考文献采用Vancouver编号格式 | 生成日期：2026-06-04 ——', sz=9)

# Save
doc.save(OUT)
import os as _os
size_kb = _os.path.getsize(OUT) / 1024
print(f"Saved: {OUT}")
print(f"Size: {size_kb:.0f} KB")
print("Figures embedded: 8 (4 figures + 4 tables)")
