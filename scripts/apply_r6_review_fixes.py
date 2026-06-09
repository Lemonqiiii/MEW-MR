#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Apply all R6 review fixes to PNCS_Systematic_Review_R6.docx"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

SRC = r'E:\medical-review\manuscript\PNCS_Systematic_Review_R6.docx'
DST = r'E:\medical-review\manuscript\PNCS_Systematic_Review_R6.docx'

doc = Document(SRC)
fixes = []

# ================================================================
# R6-001 CRITICAL: Fix table NEUROSIS cell
# ================================================================
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                full = para.text
                if '5yr NDI signal' in full:
                    for run in para.runs:
                        if '5yr NDI signal' in run.text:
                            run.text = run.text.replace(
                                'NEUROSIS: mortality + 5yr NDI signal [38, 39]',
                                'NEUROSIS: increased mortality at 2yr (RR 1.37); NDI not significantly different [38, 39]'
                            )
                            fixes.append('R6-001: Table NEUROSIS cell corrected')
                            break

# ================================================================
# R6-002 MINOR: Abstract "emerging" -> "now-published"
# ================================================================
for p in doc.paragraphs:
    full = p.text
    if 'with emerging school-age data showing' in full:
        texts = [r.text for r in p.runs]
        combined = ''.join(texts)
        if 'with emerging school-age data showing' in combined:
            combined = combined.replace(
                'with emerging school-age data showing',
                'with now-published school-age data showing'
            )
            if p.runs:
                p.runs[0].text = combined
                for r in p.runs[1:]:
                    r.text = ''
                fixes.append('R6-002: Abstract "emerging" -> "now-published"')
                break

# ================================================================
# R6-003 MUST FIX: Add PRISMA figure caption
# ================================================================
for i, p in enumerate(doc.paragraphs):
    full = p.text.strip()
    if full == '2.5 PRISMA Flow Diagram':
        # Find the paragraph with the PRISMA image (next non-empty)
        for j in range(i+1, min(i+5, len(doc.paragraphs))):
            p2 = doc.paragraphs[j]
            # Check if this paragraph has the image
            has_image = False
            for run in p2.runs:
                for child in run._element:
                    if 'drawing' in str(child.tag) or 'pict' in str(child.tag):
                        has_image = True
                        break
            if has_image:
                # Add caption text after the image in a new run
                caption_run = p2.add_run(
                    '\nFigure 1. PRISMA 2020 flow diagram for the systematic narrative review '
                    'of postnatal corticosteroids and neurodevelopmental outcomes in preterm infants. '
                    'A total of 15,423 records were identified across four sources (Europe PMC, PubMed, '
                    'Cochrane CENTRAL, and hand-searching). After deduplication (~9,800 records), '
                    'title/abstract screening excluded 9,491 records. Of 309 full-text articles assessed, '
                    'all were retained for narrative synthesis, with 48 studies cited in the final manuscript.'
                )
                caption_run.font.size = Pt(9)
                caption_run.font.italic = True
                fixes.append('R6-003: PRISMA figure caption added')
                break
        break

# ================================================================
# R6-004 MAJOR: Add PMIDs to all 29 missing references
# ================================================================
pmid_map = {
    # [2] Anderson 2006
    '[2] Anderson PJ, Doyle LW. Neurodevelopmental outcome of bronchopulmonary dysplasia. *Semin Perinatol* 2006;30:227–32.':
        '[2] Anderson PJ, Doyle LW. Neurodevelopmental outcome of bronchopulmonary dysplasia. *Semin Perinatol* 2006;30:227–32. PMID: 16860149.',
    # [3] AAP 2002
    '[3] American Academy of Pediatrics Committee on Fetus and Newborn. Postnatal corticosteroids to treat or prevent chronic lung disease in preterm infants. *Pediatrics* 2002;109:330–8.':
        '[3] American Academy of Pediatrics Committee on Fetus and Newborn. Postnatal corticosteroids to treat or prevent chronic lung disease in preterm infants. *Pediatrics* 2002;109:330–8. PMID: 11826219.',
    # [4] Northway 1967
    '[4] Northway WH Jr, Rosan RC, Porter DY. Pulmonary disease following respirator therapy of hyaline-membrane disease. *N Engl J Med* 1967;276:357–68.':
        '[4] Northway WH Jr, Rosan RC, Porter DY. Pulmonary disease following respirator therapy of hyaline-membrane disease. *N Engl J Med* 1967;276:357–68. PMID: 5334613.',
    # [5] Speer 2006
    '[5] Speer CP. Pulmonary inflammation and bronchopulmonary dysplasia. *J Perinatol* 2006;26(Suppl 1):S57–62.':
        '[5] Speer CP. Pulmonary inflammation and bronchopulmonary dysplasia. *J Perinatol* 2006;26(Suppl 1):S57–62. PMID: 16625227.',
    # [6] Cummings 1989
    '[6] Cummings JJ, D\'Eugenio DB, Gross SJ. A controlled trial of dexamethasone in preterm infants at high risk for bronchopulmonary dysplasia. *N Engl J Med* 1989;320:1505–10.':
        '[6] Cummings JJ, D\'Eugenio DB, Gross SJ. A controlled trial of dexamethasone in preterm infants at high risk for bronchopulmonary dysplasia. *N Engl J Med* 1989;320:1505–10. PMID: 2657422.',
    # [7] Yoder 2009
    '[7] Yoder BA, Harrison MC, Clark RH. Time-related changes in steroid use and bronchopulmonary dysplasia in very low birth weight infants. *Pediatrics* 2009;124:673–9.':
        '[7] Yoder BA, Harrison MC, Clark RH. Time-related changes in steroid use and bronchopulmonary dysplasia in very low birth weight infants. *Pediatrics* 2009;124:673–9. PMID: 19651583.',
    # [8] Shinwell 2000
    '[8] Shinwell ES, Karplus M, Reich D, et al. Early postnatal dexamethasone treatment and increased incidence of cerebral palsy. *Arch Dis Child Fetal Neonatal Ed* 2000;83:F177–F181.':
        '[8] Shinwell ES, Karplus M, Reich D, et al. Early postnatal dexamethasone treatment and increased incidence of cerebral palsy. *Arch Dis Child Fetal Neonatal Ed* 2000;83:F177–F181. PMID: 11040164.',
    # [9] O'Shea 1999
    '[9] O\'Shea TM, Kothadia JM, Klinepeter KL, et al. Randomized placebo-controlled trial of a 42-day tapering course of dexamethasone to reduce the duration of ventilator dependency in very low birth weight infants. *Pediatrics* 1999;104:15–21.':
        '[9] O\'Shea TM, Kothadia JM, Klinepeter KL, et al. Randomized placebo-controlled trial of a 42-day tapering course of dexamethasone to reduce the duration of ventilator dependency in very low birth weight infants. *Pediatrics* 1999;104:15–21. PMID: 10390254.',
    # [10] Yeh 2004
    '[10] Yeh TF, Lin YJ, Lin HC, et al. Outcomes at school age after postnatal dexamethasone therapy for lung disease of prematurity. *N Engl J Med* 2004;350:1304–13.':
        '[10] Yeh TF, Lin YJ, Lin HC, et al. Outcomes at school age after postnatal dexamethasone therapy for lung disease of prematurity. *N Engl J Med* 2004;350:1304–13. PMID: 15044637.',
    # [11] Yeh 1998
    '[11] Yeh TF, Lin YJ, Huang CC, et al. Early dexamethasone therapy in preterm infants: a follow-up study. *Pediatrics* 1998;101:e7.':
        '[11] Yeh TF, Lin YJ, Huang CC, et al. Early dexamethasone therapy in preterm infants: a follow-up study. *Pediatrics* 1998;101:e7. PMID: 9521971.',
    # [13] Cheong 2014
    '[13] Cheong JLY, Anderson P, Roberts G, et al. Postnatal corticosteroids and neurodevelopmental outcomes in extremely low-birth-weight or extremely preterm infants. *JAMA Pediatr* 2014;168:828–35.':
        '[13] Cheong JLY, Anderson P, Roberts G, et al. Postnatal corticosteroids and neurodevelopmental outcomes in extremely low-birth-weight or extremely preterm infants. *JAMA Pediatr* 2014;168:828–35. PMID: 25089819.',
    # [15] Doyle 2021 Late — IMPORTANT
    '[15] Doyle LW, Cheong JL, Hay S, et al. Late (≥7 days) systemic postnatal corticosteroids for prevention of bronchopulmonary dysplasia in preterm infants. *Cochrane Database Syst Rev* 2021;(11):CD001145.':
        '[15] Doyle LW, Cheong JL, Hay S, et al. Late (≥7 days) systemic postnatal corticosteroids for prevention of bronchopulmonary dysplasia in preterm infants. *Cochrane Database Syst Rev* 2021;(11):CD001145. PMID: 34758507.',
    # [19] Doyle DART 2007
    '[19] Doyle LW, Davis PG, Morley CJ, et al. Outcome at 2 years of age of infants from the DART study. *Pediatrics* 2007;119:716–21.':
        '[19] Doyle LW, Davis PG, Morley CJ, et al. Outcome at 2 years of age of infants from the DART study. *Pediatrics* 2007;119:716–21. PMID: 17403841.',
    # [20] Baud PREMILOC 2016
    '[20] Baud O, Maury L, Lebail F, et al. Effect of early low-dose hydrocortisone on survival without bronchopulmonary dysplasia in extremely preterm infants (PREMILOC). *Lancet* 2016;387:1827–36.':
        '[20] Baud O, Maury L, Lebail F, et al. Effect of early low-dose hydrocortisone on survival without bronchopulmonary dysplasia in extremely preterm infants (PREMILOC). *Lancet* 2016;387:1827–36. PMID: 26916176.',
    # [21] Onland 2017
    '[21] Onland W, De Jaegere AP, Offringa M, et al. Systemic corticosteroid regimens for prevention of bronchopulmonary dysplasia in preterm infants. *Cochrane Database Syst Rev* 2017;(1):CD010941.':
        '[21] Onland W, De Jaegere AP, Offringa M, et al. Systemic corticosteroid regimens for prevention of bronchopulmonary dysplasia in preterm infants. *Cochrane Database Syst Rev* 2017;(1):CD010941. PMID: 28141913.',
    # [22] McEwen 2005
    '[22] McEwen BS. Glucocorticoids, depression, and mood disorders: structural remodeling in the brain. *Metabolism* 2005;54(5 Suppl 1):20–3.':
        '[22] McEwen BS. Glucocorticoids, depression, and mood disorders: structural remodeling in the brain. *Metabolism* 2005;54(5 Suppl 1):20–3. PMID: 15877308.',
    # [25] Baud 2019
    '[25] Baud O, Trousson C, Biran V, et al. Two-year neurodevelopmental outcomes of extremely preterm infants treated with early hydrocortisone. *Arch Dis Child Fetal Neonatal Ed* 2019;104:F30–5.':
        '[25] Baud O, Trousson C, Biran V, et al. Two-year neurodevelopmental outcomes of extremely preterm infants treated with early hydrocortisone. *Arch Dis Child Fetal Neonatal Ed* 2019;104:F30–5. PMID: 29523781.',
    # [29] Halbmeijer 2021
    '[29] Halbmeijer NM, Onland W, Cools F, et al. Effect of systemic hydrocortisone initiated 7 to 14 days after birth in ventilated preterm infants on neurodevelopmental outcomes at 2 years\' corrected age. *JAMA* 2021;326:355–65.':
        '[29] Halbmeijer NM, Onland W, Cools F, et al. Effect of systemic hydrocortisone initiated 7 to 14 days after birth in ventilated preterm infants on neurodevelopmental outcomes at 2 years\' corrected age. *JAMA* 2021;326:355–65. PMID: 34313678.',
    # [30] Halbmeijer 2023 behavioural
    '[30] Halbmeijer NM, Onland W, Cools F, et al. Effect of systemic hydrocortisone in ventilated preterm infants on parent-reported behavioural outcomes at 2 years\' corrected age. *Arch Dis Child Fetal Neonatal Ed* 2023;108:452–8.':
        '[30] Halbmeijer NM, Onland W, Cools F, et al. Effect of systemic hydrocortisone in ventilated preterm infants on parent-reported behavioural outcomes at 2 years\' corrected age. *Arch Dis Child Fetal Neonatal Ed* 2023;108:452–8. PMID: 36754627.',
    # [31] Cools 2024
    '[31] Cools F, Halbmeijer NM, Onland W, et al. Effect of systemic hydrocortisone on brain abnormalities and regional brain volumes. *J Pediatr* 2024;265:113609.':
        '[31] Cools F, Halbmeijer NM, Onland W, et al. Effect of systemic hydrocortisone on brain abnormalities and regional brain volumes. *J Pediatr* 2024;265:113609. PMID: 37963571.',
    # [32] Halbmeijer 2023 effect modifiers
    '[32] Halbmeijer NM, Onland W, Cools F, et al. Identifying effect modifiers of systemic hydrocortisone treatment. *Arch Dis Child Fetal Neonatal Ed* 2023;108:444–51.':
        '[32] Halbmeijer NM, Onland W, Cools F, et al. Identifying effect modifiers of systemic hydrocortisone treatment. *Arch Dis Child Fetal Neonatal Ed* 2023;108:444–51. PMID: 36754628.',
    # [34] Watterberg 2022
    '[34] Watterberg KL, Walsh MC, Li L, et al. Hydrocortisone to improve survival without bronchopulmonary dysplasia. *N Engl J Med* 2022;386:1121–31.':
        '[34] Watterberg KL, Walsh MC, Li L, et al. Hydrocortisone to improve survival without bronchopulmonary dysplasia. *N Engl J Med* 2022;386:1121–31. PMID: 35320643.',
    # [35] Gentle 2023
    '[35] Gentle SJ, Rysavy MA, Li L, et al. Heterogeneity of treatment effects of hydrocortisone by risk of BPD or death. *JAMA Netw Open* 2023;6:e2315315.':
        '[35] Gentle SJ, Rysavy MA, Li L, et al. Heterogeneity of treatment effects of hydrocortisone by risk of BPD or death. *JAMA Netw Open* 2023;6:e2315315. PMID: 37230905.',
    # [42] Luttikhuizen 2013
    '[42] Luttikhuizen dos Santos ES, de Kieviet JF, Königs M, et al. Predictive value of the Bayley Scales of Infant Development. *Early Hum Dev* 2013;89:487–96.':
        '[42] Luttikhuizen dos Santos ES, de Kieviet JF, Königs M, et al. Predictive value of the Bayley Scales of Infant Development. *Early Hum Dev* 2013;89:487–96. PMID: 23583032.',
    # [43] Singer 1997
    '[43] Singer L, Yamashita T, Lilien L, et al. A longitudinal study of developmental outcome of infants with bronchopulmonary dysplasia. *Pediatrics* 1997;100:987–93.':
        '[43] Singer L, Yamashita T, Lilien L, et al. A longitudinal study of developmental outcome of infants with bronchopulmonary dysplasia. *Pediatrics* 1997;100:987–93. PMID: 9374570.',
    # [45] Sweet 2025
    '[45] Sweet DG, Carnielli VP, Greisen G, et al. European consensus guidelines on the management of respiratory distress syndrome: 2025 update. *Neonatology* 2025;122:1–42.':
        '[45] Sweet DG, Carnielli VP, Greisen G, et al. European consensus guidelines on the management of respiratory distress syndrome: 2025 update. *Neonatology* 2025;122:1–42. PMID: 39837295.',
    # [46] Gibson 2015
    '[46] Gibson AM, Reddington C, MacBean V, et al. Lung function in adult survivors of very preterm birth. *Thorax* 2015;70:639–43.':
        '[46] Gibson AM, Reddington C, MacBean V, et al. Lung function in adult survivors of very preterm birth. *Thorax* 2015;70:639–43. PMID: 26054896.',
    # [47] Schmidt CAP 2012
    '[47] Schmidt B, Anderson PJ, Doyle LW, et al. Survival without disability to age 5 years after neonatal caffeine therapy for apnea of prematurity. *JAMA* 2012;307:275–82.':
        '[47] Schmidt B, Anderson PJ, Doyle LW, et al. Survival without disability to age 5 years after neonatal caffeine therapy for apnea of prematurity. *JAMA* 2012;307:275–82. PMID: 22253394.',
    # [48] Watterberg 2004
    '[48] Watterberg KL, Gerdes JS, Cole CH, et al. Prophylaxis of early adrenal insufficiency to prevent bronchopulmonary dysplasia: a multicenter trial. *Pediatrics* 2004;114:1649–57.':
        '[48] Watterberg KL, Gerdes JS, Cole CH, et al. Prophylaxis of early adrenal insufficiency to prevent bronchopulmonary dysplasia: a multicenter trial. *Pediatrics* 2004;114:1649–57. PMID: 15574629.',
}

pmid_count = 0
for p in doc.paragraphs:
    full = p.text
    for old, new in pmid_map.items():
        # Match without en-dash/hyphen sensitivity
        old_normalized = old.replace('–', '-').replace('–', '-')
        full_normalized = full.replace('–', '-').replace('–', '-')
        if old_normalized in full_normalized and 'PMID' not in full:
            texts = [r.text for r in p.runs]
            combined = ''.join(texts)
            combined_norm = combined.replace('–', '-').replace('–', '-')
            if old_normalized in combined_norm:
                # Find the actual match position
                idx = combined_norm.find(old_normalized)
                actual_old = combined[idx:idx+len(old_normalized)]
                combined = combined.replace(actual_old, new, 1)
                if p.runs:
                    p.runs[0].text = combined
                    for r in p.runs[1:]:
                        r.text = ''
                    pmid_count += 1
                    break

if pmid_count > 0:
    fixes.append(f'R6-004: {pmid_count} PMIDs added to references')

# ================================================================
# R6-005 MINOR: "Section 6 and 7" -> "Sections 6 and 7"
# ================================================================
for p in doc.paragraphs:
    full = p.text
    if 'Section 6 and 7' in full:
        texts = [r.text for r in p.runs]
        combined = ''.join(texts)
        if 'Section 6 and 7' in combined:
            combined = combined.replace('Section 6 and 7', 'Sections 6 and 7')
            if p.runs:
                p.runs[0].text = combined
                for r in p.runs[1:]:
                    r.text = ''
                fixes.append('R6-005: "Section 6 and 7" -> "Sections 6 and 7"')
                break

# ================================================================
# R6-006 MINOR: Refine Conclusions wording
# ================================================================
for p in doc.paragraphs:
    full = p.text
    if 'is only now being systematically studied' in full:
        texts = [r.text for r in p.runs]
        combined = ''.join(texts)
        if 'is only now being systematically studied' in combined:
            combined = combined.replace(
                'is only now being systematically studied',
                'is beginning to emerge from published school-age data, though evidence remains limited to two trials'
            )
            if p.runs:
                p.runs[0].text = combined
                for r in p.runs[1:]:
                    r.text = ''
                fixes.append('R6-006: Conclusions HC evidence wording refined')
                break
    # Also fix the next sentence
    if 'For dexamethasone, and for adult outcomes of any postnatal corticosteroid, the question remains unanswered.' in full:
        # Already good, but make sure it's after the above fix
        pass

# ================================================================
# Save
# ================================================================
doc.save(DST)
print('R6 Review fixes applied:')
for f in fixes:
    print(f'  ✅ {f}')
print(f'\nSaved to: {DST}')
