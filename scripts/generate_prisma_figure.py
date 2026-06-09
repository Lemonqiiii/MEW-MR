#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Generate PRISMA 2020 flow diagram as an image and insert into the R6 Word documents."""
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import numpy as np
from docx import Document
from docx.shared import Inches, Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUTPUT_IMG = r'E:\medical-review\manuscript\prisma_flow.png'

# ================================================================
# Create PRISMA 2020 Flow Diagram
# ================================================================
fig, ax = plt.subplots(1, 1, figsize=(10, 14))
ax.set_xlim(0, 10)
ax.set_ylim(0, 14)
ax.axis('off')

# Color scheme
COLOR_ID = '#2B579A'      # Dark blue - Identification
COLOR_SCREEN = '#5B9BD5'  # Medium blue - Screening
COLOR_ELIG = '#9DC3E6'    # Light blue - Eligibility
COLOR_INCL = '#BDD7EE'    # Very light blue - Included
COLOR_EXCL = '#FF6B6B'    # Red for excluded
BOX_BG = '#F2F7FB'        # Light background for boxes
TEXT_DARK = '#333333'

def draw_box(ax, x, y, w, h, color, text, fontsize=9, bold=False, text_color='white', border_color=None):
    """Draw a rounded box with text."""
    if border_color is None:
        border_color = color
    box = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.08",
                         facecolor=color, edgecolor=border_color, linewidth=1.2, alpha=0.95)
    ax.add_patch(box)
    weight = 'bold' if bold else 'normal'
    ax.text(x + w/2, y + h/2, text, ha='center', va='center',
            fontsize=fontsize, fontweight=weight, color=text_color,
            fontfamily='sans-serif')

def draw_text_box(ax, x, y, w, h, text, fontsize=8, color=TEXT_DARK, align='left'):
    """Draw a text box with left-aligned text."""
    box = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.05",
                         facecolor='white', edgecolor='#CCCCCC', linewidth=0.8)
    ax.add_patch(box)
    lines = text.split('\n')
    line_h = h / (len(lines) + 1)
    for i, line in enumerate(lines):
        y_pos = y + h - line_h * (i + 1)
        ax.text(x + 0.15, y_pos, line, ha='left', va='center',
                fontsize=fontsize, color=color, fontfamily='monospace')

def draw_arrow(ax, x1, y1, x2, y2, color='#555555'):
    """Draw an arrow between two points."""
    ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle='->', color=color, lw=1.5,
                               connectionstyle='arc3,rad=0'))

def draw_right_arrow(ax, x1, y1, x2, y2, color='#FF6B6B'):
    """Draw a rightward exclusion arrow."""
    ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle='->', color=color, lw=1.3,
                               connectionstyle='arc3,rad=0'))

# ---- TITLE ----
ax.text(5, 13.5, 'PRISMA 2020 Flow Diagram', ha='center', va='center',
        fontsize=16, fontweight='bold', color='#2B579A')
ax.text(5, 13.1, 'Postnatal Corticosteroids for Preterm Infants: Neurodevelopmental Outcomes',
        ha='center', va='center', fontsize=9, color='#666666', style='italic')

# ---- IDENTIFICATION SECTION ----
draw_box(ax, 0.5, 12.2, 9, 0.6, COLOR_ID, 'IDENTIFICATION', fontsize=11, bold=True)

# Database boxes
db_y = 10.5
db_text = (
    'Europe PMC (6 search angles) ·············· n = 8,406\n'
    'PubMed (MeSH + free-text) ················· n = 5,832\n'
    'Cochrane CENTRAL (via CRS Web) ············ n = 1,147\n'
    'Hand-searching of reference lists ········· n = 38\n'
    '                                    Total: n = 15,423'
)
draw_text_box(ax, 0.5, db_y, 6.5, 1.5, db_text, fontsize=7.5)

# Duplicates box
draw_text_box(ax, 0.5, 9.2, 6.5, 1.1,
              'Records after duplicates removed: n = ~9,800\n'
              '(Europe PMC includes all PubMed content;\n'
              'CENTRAL overlap ~65% with PubMed-indexed trials)',
              fontsize=7.5)

# Arrow: databases -> duplicates
draw_arrow(ax, 3.75, db_y, 3.75, 10.35)

# ---- SCREENING SECTION ----
draw_box(ax, 0.5, 8.5, 9, 0.6, COLOR_SCREEN, 'SCREENING', fontsize=11, bold=True)

# Screened box
draw_text_box(ax, 0.5, 7.2, 6.5, 1.1,
              'Records screened (title/abstract): n = 9,800\n'
              'Records excluded (not meeting PICO): n = 9,491',
              fontsize=7.5)

# Exclusion arrow
ax.annotate('Records excluded', xy=(8.8, 7.7), xytext=(7.2, 7.7),
            arrowprops=dict(arrowstyle='->', color=COLOR_EXCL, lw=1.3),
            ha='center', fontsize=7.5, color=COLOR_EXCL, fontweight='bold')

# Arrow: duplicates -> screened
draw_arrow(ax, 3.75, 9.2, 3.75, 8.35)

# ---- ELIGIBILITY SECTION ----
draw_box(ax, 0.5, 6.5, 9, 0.6, COLOR_ELIG, 'ELIGIBILITY', fontsize=11, bold=True)

draw_text_box(ax, 0.5, 5.2, 6.5, 1.1,
              'Full-text articles assessed for eligibility: n = 309\n'
              'Full-text articles excluded: n = 0\n'
              '(All 309 retained for narrative synthesis)',
              fontsize=7.5)

# Arrow: screened -> eligibility
draw_arrow(ax, 3.75, 7.2, 3.75, 6.35)

# ---- INCLUDED SECTION ----
draw_box(ax, 0.5, 4.5, 9, 0.6, COLOR_INCL, 'INCLUDED', fontsize=11, bold=True, text_color=TEXT_DARK)

draw_text_box(ax, 0.5, 3.0, 6.5, 1.3,
              'Studies included in narrative synthesis: n = 309\n'
              'Studies cited in final manuscript: n = 48\n'
              '  (Priority: Cochrane SRs, RCTs with ND follow-up,\n'
              '   and recent 2022–2026 cohort studies per protocol)',
              fontsize=7.5)

# Arrow: eligibility -> included
draw_arrow(ax, 3.75, 5.2, 3.75, 4.35)

# ---- METHODOLOGICAL NOTE ----
ax.text(5, 2.3, 'Methodological Note: Single-reviewer screening (dual independent screening not performed).',
        ha='center', va='center', fontsize=7.5, color='#888888', style='italic')
ax.text(5, 2.0, 'Protocol not pre-registered in PROSPERO. Embase and CENTRAL (via Cochrane Library) not searched.',
        ha='center', va='center', fontsize=7.5, color='#888888', style='italic')

plt.tight_layout()
plt.savefig(OUTPUT_IMG, dpi=200, bbox_inches='tight', facecolor='white', edgecolor='none')
plt.close()
print(f'PRISMA figure saved to: {OUTPUT_IMG}')

# ================================================================
# Insert into both Word documents
# ================================================================
for doc_path, label in [
    (r'E:\medical-review\manuscript\PNCS_Systematic_Review_R6.docx', 'English R6'),
    (r'E:\medical-review\manuscript\PNCS_Systematic_Review_R6_CN.docx', 'Chinese R6'),
]:
    doc = Document(doc_path)

    # Find the PRISMA Flow Diagram section and replace the text paragraph with the image
    found = False
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if txt == '2.5 PRISMA Flow Diagram' or txt == '2.5 PRISMA流程图':
            heading_idx = i
            # The next non-empty paragraph should be the old text Figure 1
            # Remove it and insert image
            for j in range(i+1, min(i+5, len(doc.paragraphs))):
                if 'Figure 1' in doc.paragraphs[j].text or 'PRISMA 2020' in doc.paragraphs[j].text:
                    # Clear the text from this paragraph
                    p_old = doc.paragraphs[j]
                    for run in p_old.runs:
                        run.text = ''
                    # Add image to this paragraph
                    run = p_old.add_run()
                    run.add_picture(OUTPUT_IMG, width=Inches(5.5))
                    # Center the paragraph
                    p_old.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    found = True
                    print(f'[{label}] Inserted PRISMA figure at paragraph {j}')
                    break

        # Also look for Figure 1 in other locations (for re-insertion after body move)
        if not found and ('Figure 1.' in txt or '图1.' in txt):
            for run in p.runs:
                run.text = ''
            run = p.add_run()
            run.add_picture(OUTPUT_IMG, width=Inches(5.5))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            print(f'[{label}] Inserted PRISMA figure at paragraph {i} (alt location)')
            found = True
            break

    if not found:
        print(f'[{label}] WARNING: Could not find PRISMA section to insert figure')

    doc.save(doc_path)

print('\nDone. PRISMA flow diagram inserted into both documents.')
