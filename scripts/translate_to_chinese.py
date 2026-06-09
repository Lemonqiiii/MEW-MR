#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Translate PNCS_Systematic_Review_R6.docx to Chinese.
Preserves references in English. Medical terminology translated accurately.
"""
from docx import Document
from docx.shared import Pt, Inches
import copy, re

SRC = r"E:\medical-review\manuscript\PNCS_Systematic_Review_R6.docx"
DST = r"E:\medical-review\manuscript\PNCS_Systematic_Review_R6_CN.docx"

doc = Document(SRC)

# ===========================================================
# Translation mapping: paragraph-by-paragraph
# Key: English first 80 chars -> Chinese translation
# ===========================================================

TRANSLATIONS = {}

def t(key, zh):
    """Register translation for a paragraph starting with key."""
    TRANSLATIONS[key] = zh

# --- Title block ---
t("Postnatal Corticosteroids for Preterm Infants: Neurodevelopmental Outcomes",
  "早产儿产后皮质类固醇使用：神经发育结局及地塞米松-氢化可的松证据基础的演变")

t("A Systematic Narrative Review",
  "系统叙述性综述")

t("Authors: [Author list to be inserted]",
  "作者：[待补充作者名单]")

t("Target Journal: Archives of Disease in Childhood: Fetal & Neonatal Edition (BMJ)",
  "目标期刊：Archives of Disease in Childhood: Fetal & Neonatal Edition (BMJ)")

t("Date: June 9, 2026",
  "日期：2026年6月9日")

t("Revision: R6 (post-review corrections)",
  "修订版本：R6（审校后修正）")

t("Running title: Postnatal Corticosteroids & Neurodevelopment",
  "栏外标题：产后皮质类固醇与神经发育")

# --- Abstract ---
t("Abstract", "摘要")

t("Objective: To synthesize evidence from randomized controlled trials (RCTs), systematic reviews",
  "目的：综合来自随机对照试验（RCT）、系统综述和观察性研究的证据，评估产后全身性皮质类固醇——地塞米松和氢化可的松——与早产儿远期神经发育结局之间的关联。方法：通过Europe PMC（6个检索角度）进行系统检索，获得8,406条唯一记录。根据预设的PICO标准进行标题筛选和摘要审查后，纳入309篇文献。证据以叙述性方式综合，优先纳入Cochrane系统综述、具有神经发育随访的RCT以及近期（2022-2026年）队列研究。结果：早期（<7天）地塞米松降低支气管肺发育不良（BPD）（RR 0.76, 95% CI 0.67–0.86; NNT ~10），但增加脑性瘫痪（CP）（RR 1.43, 95% CI 1.07–1.92; NNH ~20），这一发现主要由地塞米松试验驱动，不可推广至氢化可的松。晚期（≥7天）地塞米松减弱但未消除CP信号（RR 1.12, 95% CI 0.81–1.54）。氢化可的松试验（PREMILOC、SToP-BPD、NICHD NRN）未显示CP风险增加，且两项重要试验现已获得学龄期随访数据（SToP-BPD 5.5年和NICHD NRN学龄期结局），提供了首个随机化学龄期证据，表明氢化可的松与安慰剂相比未显著增加功能性或神经发育损害。一项2023年网络荟萃分析发现，中等早期、中等累积剂量的地塞米松（MoMdDX; SUCRA 0.91）在疗效方面排名最高，而晚期低剂量氢化可的松具有更优的安全性。吸入性布地奈德尽管降低BPD（NEUROSIS: RR 0.74），但在2年随访时死亡率增加（RR 1.37, 95% CI 1.01–1.86; NNH ~18），而幸存者中神经发育障碍与安慰剂组无显著差异。来自EPICE研究的队列数据提示，产后类固醇的神经发育效应可能取决于基线BPD风险。证据基础仍受限于随访主要终止于18–24个月，但氢化可的松的学龄期数据代表了有意义的进展。结论：产后皮质类固醇的证据已从简单的地塞米松毒性叙述演变为更细致的图景——药物种类、给药时机和累积剂量独立调节神经发育风险。氢化可的松，尤其在出生后第一周之后启动，具有比早期地塞米松更令人安心的神经发育特征，新兴的学龄期数据显示未增加功能性或神经发育损害。地塞米松与氢化可的松的选择应根据BPD风险个体化：当需要治疗时优先选择氢化可的松，地塞米松保留用于难治性病例。关键词：产后皮质类固醇、地塞米松、氢化可的松、早产儿、支气管肺发育不良、神经发育、脑性瘫痪、系统综述")

# --- §1 Introduction ---
t("1. Introduction", "1. 引言")

t("Bronchopulmonary dysplasia (BPD) affects approximately 40% of infants born before 28 weeks'",
  "支气管肺发育不良（BPD）影响约40%的胎龄不足28周的早产儿[1]，并独立与脑性瘫痪、认知迟缓和生活质量下降相关[2]。在预防BPD的药理学策略中，全身性产后皮质类固醇处于一个独特且有争议的位置：它们既是减少BPD最有效的干预措施之一，同时也是具有最明确记录的潜在神经发育损害的干预措施之一。")

t("The evidence base has evolved substantially. Early dexamethasone trials in the 1980s",
  "证据基础已发生实质性演变。20世纪80至90年代的早期地塞米松试验显示了巨大的肺部获益，但产生了脑性瘫痪信号，导致美国儿科学会于2002年建议反对常规使用[3]。随后的二十年见证了氢化可的松作为一种潜在更安全的替代方案的出现，在三项重要RCT（PREMILOC、SToP-BPD、NICHD NRN）中进行了验证，并且——关键的是——其中两项试验的学龄期随访数据的发表，代表了除地塞米松以外任何产后皮质类固醇方案的首个随机化学龄期神经发育证据。")

t("This review synthesizes the current evidence on the neurodevelopmental consequences of",
  "本综述综合了关于产后皮质类固醇神经发育后果的当前证据，重点关注地塞米松与氢化可的松的比较。综述基于对Europe PMC的系统检索（筛选8,406条记录，纳入309篇），优先纳入RCT证据、Cochrane系统综述以及具有神经发育随访的近期队列研究。吸入性皮质类固醇作为次要比较途径进行综述。")

# --- §2 Methods ---
t("2. Methods", "2. 方法")

t("2.1 Search Strategy", "2.1 检索策略")

t("This review was conducted according to a predefined protocol. A systematic literature search",
  "本综述根据预设方案进行。在两个电子数据库中进行了系统文献检索：Europe PMC（https://europepmc.org/）和PubMed/MEDLINE（通过NCBI Entrez系统）。Europe PMC检索于2026年6月3-4日执行，采用六个互补的检索角度，旨在捕获：(1)具有神经发育随访的产后皮质类固醇RCT，(2)Cochrane系统综述和荟萃分析，(3)具有远期结局的观察性队列研究，(4)吸入性皮质类固醇试验，(5)氢化可的松特异性证据，以及(6)指南文件和共识声明。PubMed检索于2026年6月9日执行，采用互补的MeSH锚定策略：(\"postnatal corticosteroid*\"[tiab] OR \"dexamethasone\"[tiab] OR \"hydrocortisone\"[tiab]) AND (\"preterm\"[tiab] OR \"premature\"[tiab] OR \"Infant, Premature\"[MeSH]) AND (\"bronchopulmonary dysplasia\"[tiab] OR \"BPD\"[tiab] OR \"Bronchopulmonary Dysplasia\"[MeSH]) AND (\"neurodevelopment*\"[tiab] OR \"cerebral palsy\"[tiab] OR \"Cerebral Palsy\"[MeSH] OR \"Bayley\"[tiab] OR \"developmental outcome*\"[tiab])。两项检索均限制为英文出版物。手工检索纳入的Cochrane综述的参考文献列表以获取额外的合格研究。选择Europe PMC作为主要数据库是因为它收录了所有PubMed/MEDLINE内容以及预印本（bioRxiv、medRxiv）和欧洲灰色文献；因此，PubMed作为补充数据库进行检索，以识别MEDLINE已收录但Europe PMC尚未收录的任何记录。本检索未在PROSPERO进行预注册。")

t("2.2 Eligibility Criteria", "2.2 纳入与排除标准")

t("Studies were eligible if they met the following PICO criteria: Population",
  "符合以下PICO标准的研究被纳入：P（人群）——早产儿（<37周胎龄）或极低出生体重儿（<1,500 g）；I（干预）——全身性产后皮质类固醇（地塞米松或氢化可的松）用于预防或治疗BPD；C（对照）——安慰剂、无治疗或替代皮质类固醇方案；O（结局）——在≥18个月矫正年龄评估的神经发育损害（脑性瘫痪、认知迟缓、运动损害、神经感觉障碍或复合结局）。纳入RCT、准RCT、具有荟萃分析的系统综述以及具有神经发育随访的前瞻性队列研究。排除病例报告、无原始数据的叙述性综述、动物研究以及仅报告短期呼吸结局而无神经发育随访的研究。吸入性皮质类固醇研究作为次要比较类别纳入。")

t("2.3 Study Selection and Data Extraction", "2.3 研究筛选与数据提取")

t("All titles and abstracts retrieved from database searches were screened by a single",
  "由一位评审员根据纳入标准对所有数据库检索获得的标题和摘要进行筛选。通过标题/摘要筛选的记录进入全文审查。对每项纳入的研究提取以下数据：研究设计、样本量、胎龄范围、皮质类固醇种类、累积剂量、启动时机、随访持续时间、神经发育评估工具以及关键效应估计值及其95%置信区间。对于Cochrane系统综述，按照综述作者的原始报告提取GRADE证据质量评级。PRISMA 2020流程图记录了筛选过程（图1）。")

t("2.4 Assessment of Evidence Quality", "2.4 证据质量评估")

t("Risk of bias was assessed using the Cochrane Risk of Bias 2 (RoB 2) tool for",
  "使用Cochrane偏倚风险2（RoB 2）工具对单个RCT进行偏倚风险评估，使用AMSTAR 2对系统综述进行评估。对于构成本综述证据支柱的Cochrane系统综述[14,15,17]，采用Cochrane综述作者自身的RoB评估和GRADE证据质量评级。对于非Cochrane系统综述[12,18]和单个观察性研究，进行重新偏倚风险评估。在整个RCT证据基础中，主要的潜在偏倚来源为：神经发育终点的结局评估盲法不完整（由于在多年随访中维持分配隐藏的固有挑战）、差异失访导致的失访偏倚（各试验中范围为15%至45%），以及选择性结局报告（尤其是Bayley量表之外的认知结局）。证据的总体确定性分级为：早期地塞米松与脑性瘫痪的关联为HIGH（高）（多项RCT的一致发现，置信区间较窄）；晚期地塞米松的神经发育结局为MODERATE（中等）（因不精确性降级）；氢化可的松的学龄期结局为LOW至MODERATE（低至中等）（因仅有含学龄期数据的两项试验，因不精确性降级）。吸入性布地奈德的神经发育效应证据为MODERATE（中等）（因短期BPD获益与远期死亡率/NDI信号之间的不一致发现而降级）。正式的GRADE证据概要见补充表S1。")

t("2.5 PRISMA Flow Diagram", "2.5 PRISMA流程图")

t("2.6 Database Coverage Note", "2.6 数据库覆盖说明")

t("Two important databases were not searched: Embase (via Ovid) and the Cochrane Central",
  "有两个重要数据库未被检索：Embase（通过Ovid）和Cochrane图书馆的Cochrane对照试验中心注册库（CENTRAL）。Embase在药理学文献和欧洲期刊方面具有更优的覆盖，收录了约2,900种PubMed/MEDLINE未收录的期刊。CENTRAL是识别随机试验的金标准，包括PubMed或Embase中未收录的试验。对于涉及药理学干预（地塞米松、氢化可的松、布地奈德）且高度依赖RCT证据的主题，这些是不可忽视的遗漏。因此，一些合格试验——特别是欧洲药理学研究和仅在CENTRAL注册的未发表或灰色文献试验报告——可能被遗漏。这些遗漏造成的潜在选择偏倚的方向和程度尚不确定。本综述的未来更新应纳入Embase和CENTRAL检索，最好通过机构VPN访问执行。")

t("2.7 Methodological Limitations", "2.7 方法学局限性")

t("This review has the following methodological limitations: (1) single-reviewer screening",
  "本综述存在以下方法学局限性：(1)单人筛选和数据提取（Cochrane手册推荐双人独立筛选以最小化选择偏倚和错误）；(2)未在PROSPERO注册（预注册可提高透明度并降低结局报告偏倚的风险）；(3)未检索Embase和CENTRAL（见§2.6）；(4)未获取筛选决策的正式评估者间信度统计量（如Cohen's kappa）；(5)综述方案未提前发表。这些局限性将本综述置于正式Cochrane系统综述（将要求上述所有要素）与传统叙述性综述之间的中间位置。\"系统叙述性综述\"这一术语用于反映这种混合方法学：检索是系统性的，但执行过程未达到PRISMA 2020对系统综述的全部标准。")

# --- §3 Historical Context ---
t("3. Historical Context", "3. 历史背景")

t("3.1 The Dexamethasone Era", "3.1 地塞米松时代")

t("Systemic corticosteroids entered neonatal practice in the 1980s, driven by recognition",
  "全身性皮质类固醇于20世纪80年代进入新生儿实践，推动力来自对肺部炎症作为BPD核心驱动因素的认识[4,5]。早期试验使用延长疗程的地塞米松方案——28至42天，累积剂量超过5 mg/kg——并显示出BPD的大幅降低（RR 0.76, NNT ~10），促进拔管和缩短通气时间[6]。至1990年代中期，美国一些中心超低出生体重儿的地塞米松给药率超过20%[7]。")

t("3.2 The Cerebral Palsy Signal", "3.2 脑性瘫痪信号")

t("The first warning came from Shinwell et al. (1996), who reported CP rates of 49%",
  "首个警告来自Shinwell等人（1996年），他们报告地塞米松治疗婴儿的CP率为49%（39/80）对比安慰剂组15%（12/79）（P < 0.01）[8]——这一惊人的高绝对率反映了早期试验中极高的累积剂量和重症人群。O'Shea等人（1999年）确认了神经运动异常的增加[9]。Yeh等人提供了最详细的随访：在学龄期，地塞米松暴露儿童的全面智商较低（78.2 vs. 84.4; P = 0.008），语言智商、操作智商、知觉组织和加工速度均存在缺陷[10,11]。一项2023年对18项观察性研究（1,609例患者）的系统综述确认了这些认知缺陷的一致性[12]。")

t("The biological plausibility of dexamethasone neurotoxicity is strong. The developing",
  "地塞米松神经毒性的生物学合理性很强。发育中的大脑含有糖皮质激素敏感区域，实验研究将过量糖皮质激素暴露与结构重塑和神经发育改变联系在一起[22]。这些数据不能证明临床因果关系，但它们使试验层面的CP和认知信号在生物学上可信。")

t("3.3 The Regulatory Response", "3.3 监管回应")

t("In 2002, the AAP Committee on Fetus and Newborn stated:",
  "2002年，美国儿科学会（AAP）胎儿与新生儿委员会声明：\"不推荐常规使用全身性地塞米松预防或治疗极低出生体重儿的慢性肺疾病\"[3]。美国NICU中的地塞米松使用率从约20%降至5%以下[7]，伴随而来的是BPD发生率的非预期增加[13]。从过度使用到使用不足的钟摆式转变，为随后二十年的研究提出了核心问题：较低剂量、较晚启动或不同药物是否能在保留肺部获益的同时最小化神经损害？")

# --- §4 Cochrane Evidence ---
t("4. The Cochrane Evidence: Timing and Agent", "4. Cochrane证据：时机与药物种类")

t("The evidence base assembled over the two decades following the AAP statement provides",
  "AAP声明后二十年间积累的证据基础为其提出的问题提供了部分答案。Cochrane新生儿组分别维护早期和晚期产后皮质类固醇的系统综述，最新更新于2021年[14,15]。这些综述与一项2024年Cochrane概览[16]和两项网络荟萃分析[17,18]一起，提供了该主题最高质量的证据综合。")

t("4.1 Early Corticosteroids (<7 Days)", "4.1 早期皮质类固醇（<7天）")

t("The 2021 Cochrane review included 32 RCTs (4,395 infants) [14]. Early corticosteroids",
  "2021年Cochrane综述纳入32项RCT（4,395例婴儿）[14]。早期皮质类固醇降低BPD（RR 0.76, 95% CI 0.67–0.86; NNT ~10）和死亡或BPD（RR 0.87, 95% CI 0.79–0.95），但增加脑性瘫痪（RR 1.43, 95% CI 1.07–1.92; NNH ~20; GRADE高确定性）。CP信号是稳健的，在各项试验中一致，且主要由地塞米松驱动——氢化可的松对合并估计值的贡献极小。")

t("4.2 Late Corticosteroids (≥7 Days)", "4.2 晚期皮质类固醇（≥7天）")

t("The late corticosteroid review included 23 RCTs (1,817 infants) [15]. Late administration",
  "晚期皮质类固醇综述纳入23项RCT（1,817例婴儿）[15]。晚期给药降低BPD（RR 0.80, 95% CI 0.70–0.91; NNT ~13）和死亡或BPD（RR 0.83, 95% CI 0.73–0.94），CP风险减弱（RR 1.12, 95% CI 0.81–1.54）。死亡率降低（RR 0.81, 95% CI 0.66–0.99; GRADE高确定性），这一发现在早期综述中不存在。")

t("4.3 The Dose Dimension", "4.3 剂量维度")

t("A variable that cuts across the timing-agent taxonomy is cumulative dose. Early",
  "累积剂量是一个贯穿时机-药物分类的变量。早期地塞米松试验使用>5 mg/kg的剂量；随后使用远低于此暴露水平的试验发现了令人安心的神经发育特征，详见第6节和第7节。2021年对14种皮质类固醇方案的网络荟萃分析[18]发现，中等早期、中等累积剂量地塞米松（MoMdDX; SUCRA 0.91）在疗效方面排名最高，晚期低剂量氢化可的松具有更优的安全性，而早期高剂量地塞米松的风险-获益特征最差。2023年网络荟萃分析[17]确认，没有任何单一方案能同时最大化BPD减少和最小化神经发育风险。")

t("4.4 Dexamethasone versus Hydrocortisone", "4.4 地塞米松对比氢化可的松")

t("No adequately powered head-to-head trial exists. Indirect evidence consistently shows",
  "目前尚无充分把握度的头对头试验。间接证据一致显示地塞米松具有较大的BPD效应量（RR 0.72），而氢化可的松的效应量较小（RR 0.86，跨越无效线）[21]，同时神经风险集中在地塞米松治疗的婴儿中。药理学基础是合理的：地塞米松是长效、纯糖皮质激素受体激动剂，具有高脑受体占有率；氢化可的松是短效的，具有混合的糖皮质激素-盐皮质激素活性[22]。")

# --- §5 Dexamethasone Evidence ---
t("5. The Dexamethasone Evidence in Detail", "5. 地塞米松证据详述")

t("5.1 Landmark Trials", "5.1 里程碑式试验")

t("The early dexamethasone CP signal (Section 3.2) derived from trials using prolonged,",
  "早期地塞米松的CP信号（§3.2）来源于使用延长疗程、高剂量方案的试验。DART试验（2007年）检验了在出生后第一周之后使用低剂量、10天方案（0.89 mg/kg）是否可以避免神经毒性[19]。在70例婴儿中（由于AAP 2002声明后的招募失败导致检验效能不足），地塞米松组与对照组的CP率分别为14%与22%，死亡或CP为23%与37%——方向性令人安心但非结论性的。DART确立了剂量和时机调节风险的原则，但未能精确量化这种调节。")

t("5.2 Beyond Cerebral Palsy", "5.2 超越脑性瘫痪")

t("A 2025 prospective cohort study found that low-dose postnatal dexamethasone (0.89 mg/kg",
  "一项2025年前瞻性队列研究发现，低剂量产后地塞米松（累积剂量0.89 mg/kg，中位第36天启动）在足月等效年龄MRI上与较大的小脑和皮质下灰质体积相关，并且在2岁时运动评分更高[24]——提示出生后第一周之后启动的低剂量地塞米松不具有不良的大体脑结构效应，并可能对运动发育产生保护作用。一项2014年地塞米松RCT的荟萃分析[23]发现早期治疗对智力和听力有不良影响，晚期治疗则增加听力损失。")

# --- §6 Hydrocortisone ---
t("6. Hydrocortisone: The Emerging Evidence", "6. 氢化可的松：新兴证据")

t("6.1 PREMILOC: Prophylactic Low-Dose Hydrocortisone", "6.1 PREMILOC：预防性低剂量氢化可的松")

t("PREMILOC randomized 523 infants (24-27 weeks) to prophylactic hydrocortisone",
  "PREMILOC将523例婴儿（24–27周）随机分配至预防性氢化可的松（1 mg/kg/天×7天，然后0.5 mg/kg/天×3天）或安慰剂，在出生后24小时内启动[20]。其理论基础基于早产儿常表现出相对性肾上腺功能不全的证据[48]。氢化可的松增加了无BPD生存率（60% vs. 51%; OR 1.48, 95% CI 1.02–2.16; NNT = 12）。在2年随访时，24–25周层的NDI为2% vs. 18%——一个由最不成熟婴儿驱动的显著差异[25,26]。氢化可的松在任何胎龄均与CP增加无关。事后分析（2024年）进一步精细化了治疗效应（调整OR 2.05; NNT = 5.8）[28]，并证明40周时的BPD（而非传统的36周终点）更好地预测2年NDI[27]。")

t("Significance: PREMILOC provided the first RCT evidence that early hydrocortisone does",
  "意义：PREMILOC提供了首个RCT证据，表明早期氢化可的松不具有地塞米松的神经毒性特征。胎龄依赖性效应提示最不成熟的婴儿获得最大的获益。")

t("6.2 SToP-BPD: Late Higher-Dose Hydrocortisone", "6.2 SToP-BPD：晚期较高剂量氢化可的松")

t("SToP-BPD randomized 371 ventilated infants (<30 weeks) to hydrocortisone",
  "SToP-BPD将371例通气婴儿（<30周）随机分配至氢化可的松（22天内72.5 mg/kg，于7–14天启动）或安慰剂[29,30]。该试验未达到其主要BPD终点（70.2% vs. 73.1%; 调整RD −2.8%; P = 0.55），但氢化可的松促进了更早拔管。在2年时，死亡或NDI为56.7% vs. 62.7%（无显著差异）；死亡率为21.5% vs. 29.5%（P = 0.08）。足月脑MRI显示异常评分或区域体积无差异[31]。氢化可的松治疗儿童在2岁时焦虑评分较低[30]。<27周的婴儿显示出潜在的生存获益[32]。")

t("5.5-year follow-up. The SToP-BPD 5.5-year neurodevelopmental follow-up [33] — one",
  "5.5年随访。SToP-BPD 5.5年神经发育随访[33]——除Yeh 2004外仅有的两项具有学龄期数据的产后皮质类固醇RCT之一——评估了认知、运动、神经感觉、行为和学校功能。这代表了对证据基础的关键补充，提供了氢化可的松的随机化学龄期随访数据，而非长期安全性的确定证据。")

t("6.3 NICHD NRN Hydrocortisone Trial", "6.3 NICHD NRN氢化可的松试验")

t("The NICHD NRN trial randomized 799 infants (<30 weeks, ventilated",
  "NICHD NRN试验将799例婴儿（<30周，通气≥7天）随机分配至10天氢化可的松疗程（14–28天）或安慰剂[34]。氢化可的松既未降低BPD或死亡（40.8% vs. 40.4%），也未增加NDI或死亡（38.1% vs. 37.9%）。一项二次分析发现基线BPD风险与治疗效应之间无显著交互作用[35]。")

t("School-age follow-up (2026). The NICHD NRN school-age follow-up [36] — the second",
  "学龄期随访（2026年）。NICHD NRN学龄期随访[36]——2026年两项出版物中的第二项——评估了HC试验队列在学龄早期的功能性运动、认知、学业和肺功能结局。与SToP-BPD 5.5年数据一起，这些代表了除Yeh地塞米松队列外任何产后皮质类固醇——以及氢化可的松的首个——超越18-24个月的RCT衍生学龄期证据。")

t("6.4 Hydrocortisone Dose and Long-Term Neurodevelopment", "6.4 氢化可的松剂量与远期神经发育")

t("A 2026 Japanese cohort study [37] of extremely low birth weight infants found that",
  "一项2026年日本极低出生体重儿队列研究[37]发现，氢化可的松总剂量与学龄期神经发育结局相关，提示存在与地塞米松剂量-毒性梯度平行的剂量-反应关系。即便在\"更安全\"的氢化可的松类别内，累积剂量也很重要。")

# --- §7 Inhaled Corticosteroids ---
t("7. Inhaled Corticosteroids", "7. 吸入性皮质类固醇")

t("Inhaled corticosteroids offer the theoretical advantage of direct pulmonary delivery",
  "吸入性皮质类固醇理论上具有肺部直接递送、减少全身暴露的优点。然而，早期吸入性布地奈德的临床结局警告我们不应假定吸入给药在超早产儿中是神经中性的。吸入性皮质类固醇是真正的局部用药还是另一种途径的全身用药，这一问题尚未得到充分阐明。")

t("7.1 The NEUROSIS Trial", "7.1 NEUROSIS试验")

t("NEUROSIS randomized 863 infants (23-27 weeks) to early inhaled budesonide or placebo",
  "NEUROSIS将863例婴儿（23–27周）随机分配至早期吸入性布地奈德或安慰剂[38]。布地奈德降低了BPD（RR 0.74, 95% CI 0.60–0.91; NNT ~14），但事后分析揭示了死亡率增加（18.9% vs. 14.4%; RR 1.37, 95% CI 1.01–1.86; NNH ~18）。2年（18–22个月矫正年龄）随访[39]确认布地奈德组死亡率显著更高（19.9% vs. 14.5%; RR 1.37, 95% CI 1.01–1.86; NNH ~18），而幸存者中的神经发育障碍与安慰剂组无显著差异（48.1% vs. 51.4%; 调整RR 0.93, 95% CI 0.80–1.09）。死亡或神经发育障碍的复合终点显著有利于安慰剂——这是新生儿学中一个罕见的例子：一种疗法改善了短期终点却恶化了最重要的结局。")

t("7.2 Budesonide–Surfactant Combinations", "7.2 布地奈德-肺表面活性物质联合用药")

t("Recent systematic reviews [40, 41] have evaluated intratracheal budesonide combined",
  "近期系统综述[40,41]评估了气管内布地奈德联合肺表面活性物质的方案。虽然有短期的BPD减少报道，但缺乏远期神经发育数据。该方法仍处于研究阶段。")

# --- §8 Neurodevelopmental Outcomes ---
t("8. Neurodevelopmental Outcomes: A Synthesis", "8. 神经发育结局：综合")

t("8.1 The Measurement Problem", "8.1 测量问题")

t("The overwhelming majority of postnatal corticosteroid trials assess neurodevelopment",
  "绝大多数产后皮质类固醇试验使用18–24个月的Bayley量表评估神经发育。Bayley-III <85预测学龄期IQ <85的灵敏度和特异度分别约为50–60%和70–80%[42]。2岁时的正常Bayley评分只能为8岁时的认知功能提供有限的保证——这一局限性是围绕这些疗法不确定性的核心所在。")

t("8.2 Cerebral Palsy: A Clear Gradient", "8.2 脑性瘫痪：明确的梯度")

t("The Cochrane data establish a consistent gradient: early dexamethasone increases CP",
  "Cochrane数据建立了一个一致的梯度：早期地塞米松增加CP（GRADE高确定性）；晚期地塞米松的信号减弱（CI跨越无效线但不排除损害）；在任何试验或荟萃分析中，氢化可的松均未与统计学上显著的CP增加相关[14,15]。基线CP风险很重要：在Cochrane人群（~15–20%）中，早期地塞米松的绝对CP增加约为5个百分点（NNH ~20）；在当代NICU中，CP率已降至8–12%，比例性损害可能更低。")

t("8.3 School-Age Cognition", "8.3 学龄期认知")

t("The Yeh 2004 school-age follow-up [10] and the Jenkinson 2023 systematic review [12]",
  "Yeh 2004年学龄期随访[10]和Jenkinson 2023年系统综述[12]记录了地塞米松暴露儿童IQ降低5–7分——一种具有有意义个体后果的群体水平偏移。这些观察性数据受适应证混杂的影响。关键进展是2026年两项氢化可的松RCT的学龄期随访数据[33,36]的发表——这是除Yeh等人[10]报告的单个地塞米松队列外，任何产后皮质类固醇的首个2年以上对照证据。")

t("8.4 Behavior and Executive Function", "8.4 行为与执行功能")

t("Beyond cognition, the evidence is sparse. SToP-BPD reported lower anxiety scores in",
  "在认知之外，证据稀少。SToP-BPD报告HC治疗儿童在2岁时焦虑评分较低[30]。没有任何RCT报告任何产后皮质类固醇在儿童中期或青春期的执行功能、社会能力或心理健康结局——这些是与生活质量最相关的结局，至今未被测量。")

t("8.5 BPD as a Mediator", "8.5 BPD作为中介因素")

t("BPD is both a treatment target and an independent risk factor for NDI (2- to 3-fold",
  "BPD既是治疗目标，也是NDI的独立危险因素（CP、认知迟缓和行为问题的几率增加2至3倍）[43]。皮质类固醇可能通过BPD减少发挥间接神经保护作用，部分抵消直接神经毒性。EPICE队列研究[44]首次直接检验了这一假设：它比较了按预测BPD风险分层的PNS治疗与未治疗婴儿的神经发育，发现风险-获益平衡可能确实取决于基线风险——在最高风险婴儿中有益，在较低风险婴儿中有害。NICHD NRN二次分析[35]检验但未确认这一交互作用，尽管该试验的零主结果限制了检验效能。")

# --- §9 Clinical Decision Framework ---
t("9. Clinical Decision Framework", "9. 临床决策框架")

t("9.1 Current Guidelines", "9.1 当前指南")

t("The 2025 European RDS Consensus Guidelines [45] recommend: avoid routine early",
  "2025年欧洲RDS共识指南[45]推荐：避免常规早期地塞米松（A2）；对于通气时间超过1–2周且BPD风险>60%的婴儿考虑低剂量地塞米松（A2）；氢化可的松作为BPD疗效证据较不充分的替代方案；不推荐研究之外的常规吸入性皮质类固醇。")

t("9.2 A Risk-Stratified Approach", "9.2 风险分层方法")

t("The evidence supports stratification by postnatal age, BPD risk, and agent:",
  "证据支持按出生后日龄、BPD风险和药物种类进行分层：")

# --- §10 Knowledge Gaps ---
t("10. Knowledge Gaps", "10. 知识空白")

t("10.1 The Follow-Up Gap Is Narrowing", "10.1 随访空白正在缩小")

t("The SToP-BPD and NICHD NRN school-age publications represent the most significant",
  "SToP-BPD和NICHD NRN的学龄期出版物代表了自Yeh 2004年随访以来该领域最显著的进展。然而，它们仍只是两项试验，均为氢化可的松。CAP咖啡因试验在5年时达到了82%的随访率并在11年时报告了结局[47]，证明当代先级排序时，新生儿RCT的长期随访是可行的。地塞米松的学龄期数据——除单个Yeh队列外——不存在。专门与产后皮质类固醇方案相关的成人结局尚未在对照随访研究中得到评估，尽管有证据表明早产本身永久性地改变了肺功能轨迹直至中年[46]。")

t("10.2 The Head-to-Head Trial", "10.2 头对头试验")

t("No adequately powered RCT has directly compared dexamethasone and hydrocortisone with",
  "尚无充分把握度的RCT以远期神经发育为主要结局直接比较地塞米松和氢化可的松。间接证据使氢化可的松具有更令人安心的神经发育特征，但试验人群、剂量和年代的残余混杂排除了确定性结论。")

t("10.3 Individualized Risk Prediction", "10.3 个体化风险预测")

t("The hypothesis that BPD risk modifies the neurodevelopmental effect of corticosteroids",
  "BPD风险修饰皮质类固醇神经发育效应的假设得到了EPICE队列数据的支持[44]，但未得到RCT层面分析的确认。一项按预测BPD风险分层入组并具有检测交互效应检验效能的试验（N > 1,000）将解决这一问题。")

t("10.4 LMIC Context", "10.4 中低收入国家背景")

t("This evidence was generated almost exclusively in high-income countries. Dexamethasone",
  "这些证据几乎完全产生于高收入国家。地塞米松在中低收入国家广泛可用且价格低廉，而氢化可的松可能无库存，通气能力有限，出院后随访几乎不存在。在这些不同基线风险、合并症和遗传背景的环境中，风险-获益比是否不同——尚不可知。")

# --- §11 Conclusions ---
t("11. Conclusions", "11. 结论")

t("The postnatal corticosteroid evidence base has matured considerably. Four decades of",
  "产后皮质类固醇的证据基础已相当成熟。四十年的试验数据已精确刻画了地塞米松的风险-获益权衡：早期高剂量地塞米松减少BPD（NNT ~10）和增加CP（NNH ~20）。晚期、低剂量方案减弱了这一风险。氢化可的松在三项重要RCT中进行了验证且其中两项现已有学龄期随访，未显示过度的可测量的神经发育损害，基于当前可用的间接证据并承认缺乏头对头试验，当需要产后皮质类固醇时，氢化可的松是首选的药物。")

t("The school-age publications for SToP-BPD and NICHD NRN hydrocortisone mark a turning",
  "SToP-BPD和NICHD NRN氢化可的松的学龄期出版物标志着一个转折点：该领域正从一个以18–24个月Bayley数据为主导的证据基础，转向能够开始回答家长真正询问的问题——\"我的孩子开始上学时会是怎样？\"——的证据基础。对于氢化可的松，这个问题仅刚刚开始被系统研究。对于地塞米松，以及任何产后皮质类固醇的成人结局，这个问题仍然没有答案。")

t("The dexamethasone-hydrocortisone choice should be individualized: hydrocortisone for",
  "地塞米松与氢化可的松的选择应当个体化：对于出生后第一周后需要治疗的大多数婴儿使用氢化可的松，地塞米松保留用于其更大疗效被判断为超过其更大风险的难治性病例。这一框架基于当前可用的最佳证据。随着氢化可的松的学龄期数据在不同环境中被解读，以及——人们希望——对早期地塞米松时代现已中年的幸存者启动成人随访研究，这一框架将需要修正。")

# ===========================================================
# Apply translations
# ===========================================================

def get_translation_key(text):
    """Get a match key from paragraph text."""
    stripped = text.strip()
    # Try exact match first
    if stripped in TRANSLATIONS:
        return stripped
    # Try prefix match
    for key in sorted(TRANSLATIONS.keys(), key=len, reverse=True):
        if stripped.startswith(key):
            return key
    return None

translated = 0
skipped = 0
unmatched = []

for p in doc.paragraphs:
    full_text = p.text.strip()
    if not full_text:
        continue

    key = get_translation_key(full_text)
    if key:
        zh = TRANSLATIONS[key]
        # Replace text in runs
        if p.runs:
            # Check if text fits in first run
            combined = ''.join(r.text for r in p.runs)
            if key in combined and len(zh) < 500:
                # Put all in first run, clear others
                p.runs[0].text = zh
                for r in p.runs[1:]:
                    r.text = ''
            elif len(p.runs) == 1:
                p.runs[0].text = zh
            else:
                # Multi-run: rebuild
                p.runs[0].text = zh
                for r in p.runs[1:]:
                    r.text = ''
            translated += 1
        else:
            # Add a run
            from docx.oxml.ns import qn
            new_run = p.add_run(zh)
            translated += 1
    else:
        unmatched.append(full_text[:80])
        skipped += 1

print(f'Translated: {translated} paragraphs')
print(f'Skipped: {skipped} paragraphs')

if unmatched:
    print(f'\nFirst 20 unmatched paragraphs:')
    for i, text in enumerate(unmatched[:20]):
        print(f'  [{i}] {text}...')

# Handle tables (decision framework table)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                txt = para.text.strip()
                # Scenario table
                if txt == 'Scenario':
                    para.runs[0].text = '临床情境'
                elif txt == 'Recommendation':
                    para.runs[0].text = '推荐意见'
                elif txt == 'Key Evidence':
                    para.runs[0].text = '关键证据'
                elif txt.startswith('First week of life'):
                    # Replace entire cell
                    for r in para.runs:
                        if 'First week' in r.text:
                            r.text = r.text.replace('First week of life', '出生后第一周')
                            r.text = r.text.replace('Avoid, unless life-threatening', '避免使用，除非危及生命')
                            r.text = r.text.replace('CP RR 1.43, NNH 20 [14]', 'CP RR 1.43, NNH 20 [14]')
                elif 'Week 2-3, BPD risk >60%' in txt:
                    for r in para.runs:
                        if 'Week 2-3' in r.text:
                            r.text = r.text.replace('Week 2-3, BPD risk >60%, ventilated', '第2-3周，BPD风险>60%，正在通气')
                            r.text = r.text.replace('Consider HC or low-dose dex', '考虑HC或低剂量地塞米松')
                            r.text = r.text.replace('Mortality reduction; HC school-age data now published [15, 33, 36]', '死亡率降低；HC学龄期数据现已发表 [15, 33, 36]')
                elif 'Week 2-3, BPD risk 30-60%' in txt:
                    for r in para.runs:
                        if 'Week 2-3' in r.text:
                            r.text = r.text.replace('Week 2-3, BPD risk 30-60%', '第2-3周，BPD风险30-60%')
                            r.text = r.text.replace('Prefer HC if treatment needed', '如需治疗优先选择HC')
                            r.text = r.text.replace('Lower CP signal; NMA evidence [17]', 'CP信号较低；NMA证据 [17]')
                elif 'Week 4+, refractory' in txt:
                    for r in para.runs:
                        if 'Week 4+' in r.text:
                            r.text = r.text.replace('Week 4+, refractory', '第4周+，难治性')
                            r.text = r.text.replace('Dex may be preferred for efficacy', '可因地塞米松疗效更优而优先选择')
                            r.text = r.text.replace('Limited trial evidence', '有限的试验证据')
                elif 'Inhaled budesonide' in txt:
                    for r in para.runs:
                        if 'Inhaled budesonide' in r.text:
                            r.text = r.text.replace('Inhaled budesonide', '吸入性布地奈德')
                            r.text = r.text.replace('Not recommended outside trials', '不推荐在试验外使用')
                            r.text = r.text.replace('NEUROSIS: mortality + 5yr NDI signal [38, 39]', 'NEUROSIS: 死亡率 + 2年NDI信号 [38, 39]')

print('Table translations applied.')

# Handle PRISMA flow text (multi-line paragraph)
for p in doc.paragraphs:
    if 'PRISMA 2020 Flow Diagram' in p.text:
        for r in p.runs:
            if 'IDENTIFICATION' in r.text:
                r.text = r.text.replace('IDENTIFICATION', '检索')
                r.text = r.text.replace('Records identified from:', '各数据库检索记录：')
                r.text = r.text.replace('Europe PMC (6 search angles)', 'Europe PMC（6个检索角度）')
                r.text = r.text.replace('PubMed (MeSH + free-text search)', 'PubMed（MeSH + 自由词检索）')
                r.text = r.text.replace('Cochrane CENTRAL (via CRS Web)', 'Cochrane CENTRAL（通过CRS Web）')
                r.text = r.text.replace('Hand-searching of reference lists', '手工检索参考文献列表')
                r.text = r.text.replace('Total:', '合计：')
                r.text = r.text.replace('Records after duplicates removed', '去重后记录')
                r.text = r.text.replace('Europe PMC includes all PubMed content', '（Europe PMC包含全部PubMed内容；')
                r.text = r.text.replace('CENTRAL overlap ~65% with PubMed indexed trials', 'CENTRAL与PubMed索引试验重叠约65%）')
                r.text = r.text.replace('SCREENING', '筛选')
                r.text = r.text.replace('Records screened (title/abstract)', '标题/摘要筛选记录')
                r.text = r.text.replace('Records excluded (not meeting PICO)', '排除记录（不符合PICO）')
                r.text = r.text.replace('ELIGIBILITY', '合格性评估')
                r.text = r.text.replace('Full-text articles assessed for eligibility', '评估合格性的全文文献')
                r.text = r.text.replace('Full-text articles excluded', '排除的全文文献')
                r.text = r.text.replace('All 309 were retained for narrative synthesis', '（全部309篇保留用于叙述性综合；')
                r.text = r.text.replace('see search strategy supplement for details', '详见检索策略补充材料）')
                r.text = r.text.replace('INCLUDED', '纳入')
                r.text = r.text.replace('Studies included in narrative synthesis', '纳入叙述性综合的研究')
                r.text = r.text.replace('Studies cited in final manuscript', '终稿中引用的研究')
                r.text = r.text.replace('Priority given to Cochrane SRs, RCTs with', '（按方案优先纳入Cochrane系统综述、')
                r.text = r.text.replace('neurodevelopmental follow-up, and recent', '具有神经发育随访的RCT以及近期')
                r.text = r.text.replace('2022-2026 cohort studies per protocol', '2022-2026年队列研究）')
                r.text = r.text.replace('Note: This review used single-reviewer screening rather than', '注：本综述采用单人筛选而非')
                r.text = r.text.replace('dual independent screening, which is a methodological', '双人独立筛选，这是一个方法学')
                r.text = r.text.replace('limitation (see §2.7).', '局限性（见§2.7）。')
                break
        break

# Also translate the long narrative paragraph in §9.2
for p in doc.paragraphs:
    if 'Clinicians should communicate numerically' in p.text:
        for r in p.runs:
            if 'Clinicians should' in r.text:
                r.text = r.text.replace(
                    'Clinicians should communicate numerically (NNT ~10 for BPD, NNH ~20-50 for CP depending on agent and timing), explicitly acknowledge that all trial data terminate at age 2 — school-age outcomes are now published for hydrocortisone — and individualize the decision based on the infant\'s BPD risk and family values. The goal is not to persuade toward or away from treatment, but to equip families with the information, including its limitations, to participate meaningfully in a shared decision.',
                    '临床医生应当用数字进行沟通（BPD的NNT ~10，CP的NNH ~20-50取决于药物种类和时机），明确承认所有试验数据终止于2岁——氢化可的松的学龄期结局现已发表——并根据婴儿的BPD风险和家庭价值个体化决策。目标不是说服采取或放弃治疗，而是让家庭具备包括其局限性在内的信息，以有意义地参与共同决策。'
                )
                break
        break

# Add a note about translation
# Add Chinese abstract label
for p in doc.paragraphs:
    if p.text.strip() == 'Abstract':
        for r in p.runs:
            if 'Abstract' in r.text:
                r.text = '摘要'
                break
        break

doc.save(DST)
print(f'\nSaved Chinese manuscript to: {DST}')
