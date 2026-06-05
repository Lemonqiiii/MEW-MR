#!/usr/bin/env python3
"""
Generate a comprehensive project overview Word document for the
LUSC Immunotherapy Resistance Review Project at E:/medical-review/
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os, datetime

OUT = "E:/medical-review/NSCLC_ICI_Resistance_Project_Overview.docx"

doc = Document()

# ── Page Setup ──
for s in doc.sections:
    s.top_margin = Cm(2.54)
    s.bottom_margin = Cm(2.54)
    s.left_margin = Cm(2.54)
    s.right_margin = Cm(2.54)

# ── Styles ──
sty = doc.styles['Normal']
sty.font.name = 'Microsoft YaHei'
sty.font.size = Pt(11)
sty.paragraph_format.line_spacing = 1.5
sty.paragraph_format.space_after = Pt(6)

for lv in [1, 2, 3]:
    hs = doc.styles[f'Heading {lv}']
    hs.font.name = 'Microsoft YaHei'
    hs.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    hs.font.bold = True
    sz_map = {1: Pt(18), 2: Pt(14), 3: Pt(12)}
    hs.font.size = sz_map[lv]
    hs.paragraph_format.space_before = Pt(18 if lv == 1 else 12)
    hs.paragraph_format.space_after = Pt(8 if lv == 1 else 6)

def heading(text, level=1):
    return doc.add_heading(text, level=level)

def para(text, bold=False, size=11, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text)
    r.font.name = 'Microsoft YaHei'
    r.font.size = Pt(size)
    r.bold = bold
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    r = p.add_run(text)
    r.font.name = 'Microsoft YaHei'
    r.font.size = Pt(10.5)
    return p

def code_block(text, size=9):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_table(headers, rows, col_widths=None):
    """Add a formatted table"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for par in cell.paragraphs:
            for run in par.runs:
                run.font.size = Pt(9.5)
                run.font.bold = True
    # Data
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for par in cell.paragraphs:
                for run in par.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()  # spacer
    return table

def page_break():
    doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run('NSCLC鳞癌免疫治疗耐药机制\n综述写作项目')
r.font.name = 'Microsoft YaHei'
r.font.size = Pt(26)
r.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
r.bold = True

doc.add_paragraph()

tp2 = doc.add_paragraph()
tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = tp2.add_run('项目完整档案与部署文档')
r2.font.name = 'Microsoft YaHei'
r2.font.size = Pt(16)
r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
doc.add_paragraph()

info_items = [
    ('项目路径', 'E:\\medical-review'),
    ('生成日期', datetime.date.today().strftime('%Y-%m-%d')),
    ('项目状态', '初稿完成，待最终审阅 (Phase 6 ~90%)'),
    ('目标期刊', 'Journal for ImmunoTherapy of Cancer (JITC)'),
    ('综述语言', '英文 (内部文档使用中文)'),
    ('文档版本', 'v1.0'),
]
for label, value in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.3
    rl = p.add_run(f'{label}：')
    rl.font.size = Pt(11)
    rl.font.name = 'Microsoft YaHei'
    rl.bold = True
    rv = p.add_run(value)
    rv.font.size = Pt(11)
    rv.font.name = 'Microsoft YaHei'

page_break()

# ════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ════════════════════════════════════════════════════════════════════
heading('目录', 1)

toc_items = [
    '一、项目概述',
    '二、综述主题与PICO框架',
    '三、项目目录结构',
    '四、Agent体系架构',
    '五、完整工作流程',
    '六、文献检索详细流程',
    '七、文献筛选与PRISMA流程',
    '八、深度阅读与论文笔记',
    '九、综述稿件结构',
    '十、图表体系',
    '十一、质量管理体系 (Harness Engineering)',
    '十二、质量关卡系统 (Quality Gates)',
    '十三、写作纪律与规则',
    '十四、数据文件清单',
    '十五、脚本工具链',
    '十六、项目统计',
    '十七、关键决策记录',
    '十八、经验教训',
    '十九、技术栈',
    '二十、当前状态与后续计划',
]
for item in toc_items:
    p = doc.add_paragraph()
    r = p.add_run(item)
    r.font.size = Pt(11)
    r.font.name = 'Microsoft YaHei'

page_break()

# ════════════════════════════════════════════════════════════════════
# 一、项目概述
# ════════════════════════════════════════════════════════════════════
heading('一、项目概述', 1)

para('本项目利用Claude Code（DeepSeek V4 Pro后端）构建了一套完整的AI辅助医学综述写作系统，旨在撰写一篇高质量的英文综述论文，主题为"非小细胞肺癌鳞状细胞癌（LUSC）的免疫治疗耐药机制"。项目从零开始，在一日内完成了从项目初始化、文献检索、筛选、深度阅读、写作到审校的全流程。')

heading('项目核心特征', 2)
bullet('三层上下文体系（Tier 1常驻核心指引 + Tier 2按需加载 + Tier 3持久化知识库）')
bullet('7个专业化子Agent协作系统，覆盖搜索、筛选、分析、写作、审校、评估、编码全流程')
bullet('五维度Harness Engineering质量保证框架（成功率、效率、鲁棒性、安全性、一致性）')
bullet('六级质量关卡系统（Quality Gates），确保每阶段产出物质量')
bullet('单源真理原则：所有稿件内容唯一源文件为 jitc_submission.md，Word由脚本自动生成')
bullet('规范引用管理：每条声明必须由至少一篇引用文献摘要直接支撑')

heading('关键成果', 2)
bullet('完成系统文献检索：PubMed/Europe PMC + Semantic Scholar共检索4,106篇')
bullet('严格两轮筛选（432→62→37篇最终纳入）')
bullet('62篇论文结构化笔记（按年份/第一作者/PMID命名）')
bullet('18个交叉主题提取，覆盖三大耐药维度')
bullet('完成8,969词英文综述初稿（7章节+摘要+图表）')
bullet('41篇参考文献，4张图+4张表')
bullet('建立完整的质量保证和错误模式库')

page_break()

# ════════════════════════════════════════════════════════════════════
# 二、综述主题与PICO框架
# ════════════════════════════════════════════════════════════════════
heading('二、综述主题与PICO框架', 1)

heading('综述标题', 2)
para('英文：Mechanisms of Immunotherapy Resistance in Squamous Cell Carcinoma of Non-Small Cell Lung Cancer')
para('中文：非小细胞肺癌鳞状细胞癌免疫治疗耐药机制研究进展')

heading('PICO框架', 2)
add_table(
    ['维度', '内容'],
    [
        ['Population', 'NSCLC患者，限鳞状细胞癌亚型（含早晚期）'],
        ['Intervention/Exposure', '免疫检查点抑制剂：anti-PD-1 (nivolumab, pembrolizumab), anti-PD-L1 (atezolizumab, durvalumab), anti-CTLA-4 (ipilimumab), 单药或联合'],
        ['Comparison', 'ICI应答者 vs 无应答者；原发性耐药 vs 获得性耐药；鳞癌 vs 非鳞NSCLC'],
        ['Outcome', '耐药机制的识别与表征——肿瘤内在因素、TME改变、基因组/转录组决定因素、免疫逃逸通路'],
    ]
)

heading('排除标准', 2)
bullet('非鳞NSCLC亚型：肺腺癌(LUAD)、大细胞癌等')
bullet('化疗耐药、靶向治疗耐药（非ICI耐药）')
bullet('喉鳞癌(laryngeal)、头颈鳞癌、口腔鳞癌、食管鳞癌、皮肤鳞癌、宫颈鳞癌等非肺鳞癌')
bullet('病例报告、勘误、编辑信件等非原始研究')

heading('检索关键词', 2)
para('PubMed检索式（初稿）：')
code_block('("non-small cell lung cancer"[MeSH] OR "NSCLC" OR "lung cancer")\nAND ("squamous cell carcinoma" OR "squamous" OR "lung squamous cell carcinoma")\nAND ("immunotherapy" OR "immune checkpoint inhibitor" OR "anti-PD-1" OR "anti-PD-L1"\n     OR "anti-CTLA-4" OR "pembrolizumab" OR "nivolumab" OR "atezolizumab"\n     OR "durvalumab" OR "ipilimumab")\nAND ("resistance" OR "resistant" OR "refractory" OR "immune evasion"\n     OR "non-response" OR "tumor microenvironment" OR "TME")')

page_break()

# ════════════════════════════════════════════════════════════════════
# 三、项目目录结构
# ════════════════════════════════════════════════════════════════════
heading('三、项目目录结构', 1)

para('项目采用精心设计的三层架构组织信息，确保上下文利用率最大化和token消耗最小化：')

code_block('''E:/medical-review/
├── CLAUDE.md                    # Tier 1: 会话常驻核心指引 (5.6 KB)
├── README.md                    # 项目介绍与使用说明
├── .gitignore
├── .claude/
│   └── settings.local.json      # 项目级Claude Code配置
│
├── memory/                      # Tier 2: 结构化Memory指针 (按需加载)
│   ├── MEMORY.md                #   Memory索引
│   ├── project-status.md        #   项目状态 (当前阶段/进度/统计)
│   ├── active-focus.md          #   当前研究方向与PICO框架
│   ├── agent-specializations.md #   7个子Agent完整定义
│   ├── key-findings.md          #   18个交叉主题与核心发现
│   ├── decisions.md             #   关键决策记录
│   └── lessons-learned.md       #   8条经验教训+错误模式库
│
├── docs/                        # Tier 3: 持久化知识库
│   ├── index.md                 #   文献知识库索引 (含PRISMA流程)
│   ├── glossary.md              #   医学术语表 (~80个术语)
│   ├── methods/                 #   方法论指南
│   │   ├── systematic-review.md #     系统综述方法 (PRISMA, Cochrane)
│   │   ├── meta-analysis.md     #     荟萃分析方法指南
│   │   ├── statistical-methods.md #   生物医学常用统计方法
│   │   └── database-coverage.md #     数据库覆盖目录 (7个数据库)
│   ├── papers/                  #   论文笔记
│   │   ├── template.md          #     笔记模板 (含PICO框架)
│   │   └── lusc_ici_resistance/ #     论文笔记目录 (62篇)
│   │       └── README.md        #     论文笔记索引
│   └── search-results/
│       └── manual-search-checklist.md # VPN手动检索清单模板
│
├── data/                        # 检索与筛选数据 (总计 ~18 MB)
│   ├── pubmed_merged_all.json   #   全量检索结果 (4,106篇)
│   ├── pubmed_search_results.json # PubMed初步检索结果
│   ├── pubmed_search_supplement.json # 补充检索结果
│   ├── pubmed_relevant_for_screening.json # 去重后初筛列表
│   ├── screening_final_40.json  #   最终纳入文献 (37篇)
│   ├── screening_final_80.json  #   全文候选列表
│   ├── screening_excluded.json  #   排除文献记录
│   ├── screening_included.json  #   纳入文献记录
│   ├── screening_uncertain.json #   不确定文献
│   ├── screening_log.md         #   筛选日志
│   └── batch1-3_papers.json     #   分批深度阅读数据
│
├── manuscript/                  # 综述稿件
│   ├── jitc_submission.md       #   ★ 唯一源文件 (8,969词, 41引用)
│   ├── complete_draft.md        #   合并草稿
│   ├── outline.md               #   综述大纲 (7章节估8,300词)
│   ├── figures_tables.md        #   图表详细定义 (4 Fig + 4 Table)
│   ├── English_Manuscript_NSCLC_ICI_Resistance.docx # 生成的Word文档
│   ├── md2docx.py               #   Markdown→Word转换脚本
│   ├── sections/                #   分章节草稿
│   │   ├── 01_introduction.md
│   │   ├── 02_immune_landscape.md
│   │   ├── 03_tumor_intrinsic.md
│   │   ├── 04_tme_mediated.md
│   │   ├── 05_acquired_resistance.md
│   │   ├── 06_overcoming_strategies.md
│   │   └── 07_conclusions.md
│   └── figures/                 #   图表PNG文件 (8张, 总计 ~1.2 MB)
│       ├── Figure1_PRISMA.png
│       ├── Figure2_Framework.png
│       ├── Figure3_EIC.png
│       ├── Figure4_Decision.png
│       ├── Table1_Clinical_Trials.png
│       ├── Table2_Tumor_Intrinsic.png
│       ├── Table3_TME_Mechanisms.png
│       └── Table4_LUSC_vs_LUAD.png
│
├── scripts/                     # 工具脚本 (11个Python文件)
│   ├── gen_word_full.py         #   核心：Markdown→Word生成器+自检
│   ├── gen_english_word.py      #   旧版生成器
│   ├── gen_word.py / gen_word_v2.py # 历史版本
│   ├── generate_word.py         #   另一个生成器变体
│   ├── generate_figures.py      #   图表生成脚本
│   ├── redraw_figures.py        #   图表重绘脚本
│   ├── pubmed_search.py         #   PubMed检索脚本
│   ├── pubmed_search_supplement.py # 补充检索脚本
│   ├── fulltext_screening.py    #   全文筛选脚本
│   └── screen_abstracts.py      #   摘要筛选脚本
│
├── features/                    # 任务跟踪
│   └── FEATURE_LIST.md          #   6个Phase共30+子任务
│
├── progress/                    # 进度记录
│   ├── SESSION_LOG.md           #   会话日志
│   ├── MILESTONES.md            #   里程碑
│   └── metrics-raw.json         #   效率指标原始数据
│
└── harness/                     # 质量保证框架
    ├── README.md                #   Harness概述与架构
    ├── metrics.md               #   五维度度量定义
    ├── quality-gate.md          #   六级质量关卡+可执行脚本
    ├── safety-policy.md         #   安全策略与越权检测
    ├── test-scenarios.md        #   鲁棒性测试场景库 (L1-L5, 13场景)
    ├── consistency-benchmarks.md #  一致性基准测试 (5个基准)
    ├── remediation-plan.md      #   稿件质量修复计划
    └── eval-log.md              #   评估日志''')

page_break()

# ════════════════════════════════════════════════════════════════════
# 四、Agent体系架构
# ════════════════════════════════════════════════════════════════════
heading('四、Agent体系架构', 1)

para('项目定义了7个专业化的子Agent，分为5个横向执行Agent和2个纵向基础设施Agent，每个Agent有极简命令触发和完善的工作流定义：')

heading('Agent全景图', 2)
add_table(
    ['编号', 'Agent名称', '定位', '触发命令', '核心职责'],
    [
        ['0', '编码Agent', '纵向基础设施', '"编码"/"记"/"6"', '进度记录+效率数据收集+安全审计+Git提交'],
        ['1', '文献搜索Agent', '横向执行', '"搜索"/"搜"/"1"', '多数据库系统检索、五层全面性验证、全文获取分级'],
        ['2', '论文分析Agent', '横向执行', '"分析"/"读"/"3"', '结构化提取论文信息，按模板生成文献笔记'],
        ['3', '综述写作Agent', '横向执行', '"写作"/"写"/"4"', '基于论文笔记撰写综述段落，声明-引用配对验证'],
        ['4', '审校Agent', '横向执行', '"审校"/"审"/"5"', '事实核查+逻辑审查+语言润色+引用溯源+图表一致性'],
        ['5', '评估Agent', '纵向基础设施', '"评估"/"评"/"7"', 'L2业务质量评分+鲁棒性测试+一致性测试+安全复核'],
        ['6', '筛选Agent', '横向执行', '"筛选"/"筛"/"2"', '两轮筛选（标题/摘要→全文），宁滥勿缺策略'],
    ]
)

heading('Agent间协作流程', 2)
code_block('''用户:"主题" → 确定综述方向 (active-focus.md)
    ↓
用户:"搜索" → 文献搜索Agent [Agent 1]
    ├─ Tier 1 自动检索 (PubMed + S2 + EPMC + CT.gov)
    ├─ Tier 2 手动清单 (Embase/Cochrane/CNKI等)
    └─ 五层全面性验证 → Handoff
    ↓
用户:"筛选" → 筛选Agent [Agent 6]
    ├─ Round 1: 标题/摘要 (宁滥勿缺 → 432→62)
    └─ Round 2: 全文筛选 (≤20%仅摘要 → 62→37)
    ↓
用户:"分析" → 论文分析Agent [Agent 2]
    └─ 逐篇结构化笔记 → docs/papers/lusc_ici_resistance/
    ↓
用户:"写作" → 综述写作Agent [Agent 3]
    └─ 笔记→草稿段落 → jitc_submission.md
    ↓
用户:"审校" → 审校Agent [Agent 4]
    └─ 事实核查+逻辑+引用溯源+Gate 4/5验证
    ↓
用户:"编码"→编码Agent / "评估"→评估Agent
    └─ 进度+效率+安全 / L2+鲁棒+一致性''')

heading('Agent 3v2 & 4v2 改进版 (2026-06-05编码)', 2)
para('基于项目实践中的教训，对写作Agent和审校Agent进行了重大升级：')
bullet('Agent 3v2新增：单源真理原则、引用铁律、扩展前验证、图表纪律、喉鳞癌过滤规则')
bullet('Agent 4v2新增：引用-声明配对验证(Gate 4)、格式完整性(Gate 5)、图表编号一致性检查、喉鳞癌/非鳞文献标记')
bullet('Agent 5v2新增：Gate 4/5可执行脚本、错误模式库匹配')

page_break()

# ════════════════════════════════════════════════════════════════════
# 五、完整工作流程
# ════════════════════════════════════════════════════════════════════
heading('五、完整工作流程', 1)

para('项目按照以下6个Phase顺序推进，每个Phase结束后编码Agent自动更新项目状态：')

add_table(
    ['Phase', '名称', '状态', '核心任务'],
    [
        ['1', '项目初始化', '✅ 完成', '搭建目录结构、配置Claude Code hooks、初始化Git仓库、搭建Harness Engineering框架（五维度度量+测试场景+安全策略）'],
        ['2', '文献搜索', '✅ 完成', '确定检索策略和关键词、PubMed/Europe PMC/Semantic Scholar三路检索、去重生成初筛列表、导出JSON格式'],
        ['3', '文献筛选', '✅ 完成', '标题/摘要筛选(432→62)、全文获取(60/62 PMC OA)、全文筛选(62→37)、PRISMA流程图'],
        ['4', '深度阅读', '✅ 完成', '分批笔记Batch 1(10篇)+Batch 2(10篇)+Batch 3(11篇)、提取18个交叉主题、构建证据表(Tables 1-4)'],
        ['5', '写作', '✅ 完成', '确定大纲、撰写7章节+Abstract+Key Messages、初稿8,969词/41引用'],
        ['6', '修改与定稿', '~90%', '内部审校✅、语言润色✅、引用核查✅、图表生成✅、格式化参考文献(待定期刊)、最终定稿(待用户审阅)'],
    ]
)

heading('阶段流转规则', 2)
code_block('''init → literature-search → screening → deep-reading → writing → revision → 投稿归档
   ↑ 用户确认主题        ↑ PRISMA完成        ↑ 核心论点提取   ↑ 初稿完成
                          ↑ 432→40            30篇精读+18主题 7章节8,969词''')

page_break()

# ════════════════════════════════════════════════════════════════════
# 六、文献检索详细流程
# ════════════════════════════════════════════════════════════════════
heading('六、文献检索详细流程', 1)

heading('数据库分层', 2)
add_table(
    ['层级', '数据库', '访问方式', '说明'],
    [
        ['Tier 1 (自动)', 'PubMed/MEDLINE', 'E-utilities API / WebFetch', '5,600+期刊，36M+记录，MeSH词表'],
        ['Tier 1 (自动)', 'Semantic Scholar', 'REST API', '214M+论文，语义搜索，TLDR摘要'],
        ['Tier 1 (自动)', 'Europe PMC', 'REST API', '含预印本(bioRxiv/medRxiv)，OA全文链接'],
        ['Tier 1 (自动)', 'ClinicalTrials.gov', 'API v2', '450K+试验记录（本综述未激活）'],
        ['Tier 2 (VPN手动)', 'Embase (Ovid)', 'VPN+预编译检索式', '药理学金标准，8,500+期刊'],
        ['Tier 2 (VPN手动)', 'Cochrane CENTRAL', 'VPN+预编译检索式', '临床试验注册库'],
        ['Tier 2 (VPN手动)', 'CNKI/万方/SinoMed', 'VPN+预编译检索式', '中文文献数据库'],
    ]
)

heading('五层全面性验证', 2)
add_table(
    ['层级', '验证方法', '说明'],
    [
        ['L1 多策略检索覆盖', '四路独立检索', 'PubMed + S2 + EPMC + CT.gov 确认响应状态'],
        ['L2 多点引文扩散', '四锚点扩散', '共识端(高引)+时间端(最新)+方法学端(RCT/队列/SR/基础)+地域端 → 反向查参考文献+前向查引文'],
        ['L3 外部金标准验证', '交叉验证', '从近2-3篇高质量系统综述提取纳入文献列表，计算命中率（≥90%=可接受）'],
        ['L4 灰色文献补充', '预印本+会议+未发表试验', 'Europe PMC自动覆盖预印本；CT.gov检查未发表试验'],
        ['L5 饱和+对抗检验', '检索饱和+Discussion验证', '连续扩展检索式直至PMID增量<5%；随机5篇Discussion中引文是否被检索覆盖'],
    ]
)

heading('实际检索统计', 2)
bullet('Europe PMC (PubMed/MEDLINE)：7个关键词查询(2025-2026)共1,614篇 + 年份分层检索(2020-2024)共2,492篇 = 4,106篇')
bullet('Semantic Scholar：429限流，进行了有限补充检索')
bullet('Tier 2数据库：未激活（此综述为叙述性综述，无需中文/Embase库；VPN不可用情况下标记为方法学局限性）')

page_break()

# ════════════════════════════════════════════════════════════════════
# 七、文献筛选与PRISMA流程
# ════════════════════════════════════════════════════════════════════
heading('七、文献筛选与PRISMA流程', 1)

heading('筛选流程', 2)

heading('Round 1: 标题/摘要筛选 (432→62)', 3)
para('宁滥勿缺策略：任何不确定(UNCERTAIN)→纳入。只有明确违反关键纳入标准才排除。')
add_table(
    ['排除原因代码', '说明', '数量'],
    [
        ['WRONG_POPULATION', '人群不匹配（非NSCLC/非鳞癌/喉鳞癌等）', '主要排除原因'],
        ['WRONG_INTERVENTION', '干预不匹配（非ICI治疗）', '次要排除原因'],
        ['WRONG_DESIGN', '研究设计不符合（病例报告等）', '部分排除'],
        ['NOT_ORIGINAL', '非原始研究/综述/评论/信件', '部分排除'],
        ['UNCERTAIN/LOW_INFO', '信息不足→宁滥勿缺→INCLUDE', '纳入Round 2'],
    ]
)

heading('Round 2: 全文筛选 (62→37)', 3)
para('逐篇全文评估，复核PICO匹配度，检查方法学质量。关键发现：')
bullet('PMC开放获取：60/62篇（96.8%），仅2篇为预印本')
bullet('排除25篇：喉鳞癌1篇 + 非鳞NSCLC 2篇 + 机制内容不足13篇 + 冗余预后模型9篇')
bullet('最终纳入37篇（综述8篇21.6% + 原始研究26篇70.3% + 系统综述/荟萃分析3篇8.1%）')
bullet('仅摘要比例：0%（全部37篇均获取全文）——远优于≤20%标准')

heading('PRISMA 2020 流程图', 2)
code_block('''检索获得 (n=4,106)    Europe PMC: 1,614 (2025-2026) + 2,492 (2020-2024)
     │
     ├─ 去重 + 排除勘误/病例报告 (n=267)
     ▼
初筛候选 (n=432)       关键词筛选: squamous + resistance
     │
     ├─ 标题/摘要筛选 → 排除370篇
     ▼
全文候选 (n=62)         PMC OA: 60/62 (96.8%)
     │
     ├─ 全文评估 → 排除25篇
     ├─   喉鳞癌误分类: 1
     ├─   非鳞NSCLC: 2
     ├─   机制内容不足: 13
     ├─   冗余预后模型: 9
     ▼
最终纳入 (n=37)         综述8 + 原始研究26 + SR/MA 3
                        年份分布: 2020(3) 2021(6) 2022(8) 2023(2) 2024(5) 2025(8) 2026(5)''')

page_break()

# ════════════════════════════════════════════════════════════════════
# 八、深度阅读与论文笔记
# ════════════════════════════════════════════════════════════════════
heading('八、深度阅读与论文笔记', 1)

heading('笔记模板', 2)
para('所有论文笔记遵循统一的模板 (docs/papers/template.md)，包含以下结构化字段：')
bullet('元数据：PMID/DOI、期刊、年份、第一作者、引用次数、检索来源')
bullet('研究问题：一句话描述核心问题')
bullet('PICO框架：Population / Intervention / Comparison / Outcome')
bullet('方法：设计类型(RCT/队列/基础实验等)、样本量、关键方法、偏倚风险评估')
bullet('核心发现：分条列出关键结果')
bullet('关键数据：效应量、95% CI、p值等定量数据')
bullet('局限性：作者自述+审稿人视角')
bullet('与本综述的关系：支撑论点、引用位置(outline对应节)、重要性评级(★★★/★★/★)')
bullet('交叉引用：相关论文笔记PMID、与其他论点的关联')

heading('阅读进度', 2)
add_table(
    ['批次', '篇数', '内容', '状态'],
    [
        ['Batch 1', '10篇', '核心综述+关键机制论文', '✅ 精读完成'],
        ['Batch 2', '10篇', '信号通路+免疫微环境论文', '✅ 精读完成'],
        ['Batch 3', '11篇', '耐药机制+联合策略论文', '✅ 精读完成'],
        ['剩余', '7篇', '纯预后模型，仅做摘要笔记', '✅ 摘要笔记完成'],
    ]
)

heading('18个交叉主题', 2)
para('从30篇精读文献中提取了18个跨文献的交叉主题（详见 memory/key-findings.md）：')
bullet('LUSC TIME异质性与免疫分型 (PMID: 33807509, 41133013, 38803944, 35799269)')
bullet('T细胞功能障碍/耗竭 (PMID: 33807509, 38803944, 41133013)')
bullet('EMT驱动的免疫逃逸 (PMID: 42089102, 38803944, 41530460)')
bullet('表观遗传调控与免疫耐药 (PMID: 35116387, 41133013)')
bullet('CAF-TAM免疫抑制轴 (PMID: 41133013, 42111396)')
bullet('LUSC独特基因组特征 (PMID: 33807509, 41133013, 41595192)')
bullet('ICI联合策略以克服耐药 (PMID: 33807509, 42078800, 34429332)')
bullet('免疫耗竭型(EIC)与多重检查点共表达 (PMID: 35799269, 38803944)')
bullet('KEAP1/NRF2通路与免疫耐药 (PMID: 36198685, 33807509)')
bullet('表型可塑性与组织学转化 (PMID: 40568576, 37025908)')
bullet('代谢重编程驱动的免疫抑制 (PMID: 41239433, 36198685)')
bullet('p38 MAPK信号驱动的免疫逃逸 (PMID: 41050683)')
bullet('原发vs获得性耐药的临床证据 (PMID: 34429332, 33807509)')
bullet('肿瘤微生物组-免疫代谢crosstalk (PMID: 40568577)')
bullet('缺氧-STING通路驱动免疫逃逸 (PMID: 40138855)')
bullet('circRNA介导的抗PD-1耐药 (PMID: 35525959)')
bullet('LUAD vs LUSC免疫异质性 (PMID: 34394068, 40083325)')
bullet('EGFR-TKI在LUSC中的固有耐药 (PMID: 33133263)')

page_break()

# ════════════════════════════════════════════════════════════════════
# 九、综述稿件结构
# ════════════════════════════════════════════════════════════════════
heading('九、综述稿件结构', 1)

para('稿件遵循标准学术综述结构，包含标题页、摘要、关键信息、7个正文章节、声明和参考文献：')

add_table(
    ['章节', '标题', '字数', '核心内容'],
    [
        ['—', 'Title Page', '—', '标题、Running title、作者信息、字数统计、关键词'],
        ['—', 'Abstract', '246词', '结构化摘要：背景→问题→范围→结论'],
        ['—', 'Key Messages', '~200词', '已知背景/新增贡献/研究影响 (JITC格式)'],
        ['1', 'Introduction', '~1,300词', 'LUSC流行病学、基因组特征、ICI标准治疗地位、耐药问题定义、综述范围'],
        ['2', 'Immune Landscape of LUSC', '~1,600词', 'TIME异质性与分子分类、LUSC vs LUAD差异、Exhausted Immune Class (EIC)'],
        ['3', 'Tumor-Intrinsic Resistance', '~2,200词', 'PI3K/AKT/mTOR、KEAP1/NRF2、p38 MAPK/TGM2、FGFR1/EGFR悖论、EMT、表观遗传、细胞死亡通路、基因组不稳定性'],
        ['4', 'TME-Mediated Resistance', '~2,100词', 'T细胞耗竭、多检查点共表达、TAMs/CAFs/MDSCs/Tregs、细胞因子网络、代谢调控(糖酵解/缺氧/色氨酸)、微生物组'],
        ['5', 'Acquired Resistance', '~950词', '克隆进化、组织学转化(AST)、表型可塑性、SCLC转化'],
        ['6', 'Overcoming Strategies', '~800词', '联合免疫治疗、靶向肿瘤内在通路、TME重塑、新兴策略(TIME精准/Biomarker)'],
        ['7', 'Conclusions', '~470词', '三大原则总结、临床意义、未来研究方向'],
        ['—', 'Declarations', '~150词', 'Funding/Competing interests/Author contributions/Ethics/Data availability'],
        ['—', 'References', '41篇', 'Vancouver格式，每篇包含PMID/DOI'],
    ]
)

heading('稿件元数据', 2)
bullet('总字数：8,969词（正文）+ 246词（摘要）')
bullet('图表数：1 Figure + 2 Tables（简化为代表性图表）')
bullet('参考文献：41篇')
bullet('目标期刊：Journal for ImmunoTherapy of Cancer (JITC)')
bullet('引用格式：Vancouver (BMJ house style)')

page_break()

# ════════════════════════════════════════════════════════════════════
# 十、图表体系
# ════════════════════════════════════════════════════════════════════
heading('十、图表体系', 1)

para('项目中定义了4张图和4张表（实际Word文档中嵌入2张代表性的Figure和2张Table）：')

heading('Figures', 2)
add_table(
    ['编号', '文件名', '内容', '大小'],
    [
        ['Figure 1', 'Figure1_PRISMA.png', 'PRISMA 2020流程图（检索→筛选→纳入）', '60 KB'],
        ['Figure 2', 'Figure2_Framework.png', '三维耐药框架图（Tumor-Intrinsic / TME-Mediated / Acquired → Convergence）', '355 KB'],
        ['Figure 3', 'Figure3_EIC.png', 'Exhausted Immune Class结构图（T细胞耗竭+9检查点+免疫抑制+临床特征）', '51 KB'],
        ['Figure 4', 'Figure4_Decision.png', '基于TIME亚型的治疗决策框架', '60 KB'],
    ]
)

heading('Tables', 2)
add_table(
    ['编号', '文件名', '内容', '大小'],
    [
        ['Table 1', 'Table1_Clinical_Trials.png', 'LUSC ICI关键临床试验汇总（KEYNOTE-407等7项）', '36 KB'],
        ['Table 2', 'Table2_Tumor_Intrinsic.png', '肿瘤内在耐药机制与治疗策略（8种机制）', '264 KB'],
        ['Table 3', 'Table3_TME_Mechanisms.png', 'TME介导的耐药机制与干预策略（9种机制）', '388 KB'],
        ['Table 4', 'Table4_LUSC_vs_LUAD.png', 'LUSC vs LUAD关键免疫学差异（8个维度）', '43 KB'],
    ]
)

page_break()

# ════════════════════════════════════════════════════════════════════
# 十一、质量管理体系
# ════════════════════════════════════════════════════════════════════
heading('十一、质量管理体系 (Harness Engineering)', 1)

para('项目构建了完整的Agent质量保证与可观测性框架，称为"Harness Engineering"，从五个维度对Agent执行质量进行系统性度量：')

add_table(
    ['维度', '核心问题', '判定方式', '执行频率'],
    [
        ['成功率', 'Agent是否完成了任务？', 'L1技术成功(自动) + L2业务成功(审校Agent评定1-5分)', '每次任务'],
        ['效率', '完成任务的代价多大？', '编码Agent自动收集：墙钟时间、工具调用次数、token消耗、数据库覆盖度', '每次任务'],
        ['鲁棒性', '输入扰动时能否正常运行？', 'L1-L5对抗性测试场景（输入漂移/中途变更/矛盾指令/撤回重建/模糊歧义）', '每Phase'],
        ['安全性', '是否发生了越权操作？', '编码Agent扫描：文件越界/网络越界/命令越界/配置篡改/信息泄露 → 评估Agent复核', '每次任务'],
        ['一致性', '同一任务多次运行结果是否稳定？', '重跑基准任务 + 工具调用序列编辑距离 + 语义diff（论点Jaccard/引用重叠/数据误差）', '每Phase'],
    ]
)

heading('综合评分公式', 2)
code_block('Phase 综合分 = 成功率×0.35 + 效率×0.10 + 鲁棒性×0.20 + 安全性×0.15 + 一致性×0.20')

heading('安全策略', 2)
para('定义了完整的越权检测规则，包括：')
bullet('文件越界检测：Read/Write/Edit限制在 E:\\medical-review\\ 和白名单系统路径')
bullet('网络越界检测：WebFetch域名白名单（pubmed, pmc, europepmc, semanticscholar, springer, ncbi, ebi）')
bullet('命令越界检测：允许 git/python/pip/mkdir/find 等受限模式')
bullet('配置篡改检测：CLAUDE.md、settings.json等敏感文件的非用户明确请求修改')
bullet('信息泄露检测：API key模式、Bearer token、Authorization header')
bullet('严重度响应：CRITICAL→立即终止+通知 | HIGH→终止+记录 | MEDIUM→评估复核 | LOW→趋势监控')

heading('鲁棒性测试场景库', 2)
para('定义了5个层级共13个测试场景：')
bullet('L1 输入漂移：中英文术语等价性、术语粒度变化、输入格式变化（DOI/PMID/URL）')
bullet('L2 中途变更：方向切换（机制→临床）、范围收窄（500→2024-2025 RCT）、格式要求中途变更')
bullet('L3 矛盾指令：篇幅vs覆盖矛盾、证据强度vs前沿热点矛盾、方法学矛盾（系统综述vs无限制）')
bullet('L4 撤回重建：数据源替换(Pubmed→S2)、筛选标准重置、笔记模板替换')
bullet('L5 模糊歧义：数量模糊（"几篇"）、质量模糊（"更好"）、概念模糊（"最近Nature那个"）')

heading('一致性基准测试', 2)
para('5个基准任务，每Phase结束时每个跑2次，对比行为路径和语义输出：')
bullet('Bench-001：PubMed精确检索（Jaccard≥80%）')
bullet('Bench-002：论文元数据提取（精确匹配PMID 38200123）')
bullet('Bench-003：术语定义查找（ITT→glossary.md）')
bullet('Bench-004：统计解释（I²=75%→"高异质性"）')
bullet('Bench-005：功能清单读取（下一个未完成任务）')

page_break()

# ════════════════════════════════════════════════════════════════════
# 十二、质量关卡系统
# ════════════════════════════════════════════════════════════════════
heading('十二、质量关卡系统 (Quality Gates)', 1)

para('每个Phase的输出在进入下一Phase之前必须通过质量关卡检查。项目定义了6个关卡：')

add_table(
    ['关卡', '阶段切换', '通过标准', '执行情况'],
    [
        ['Gate 1', '检索→筛选', '查全率5/5命中 + 去重0重复 + 数据完整性<5%缺失', '未执行（轻量流程）'],
        ['Gate 2', '筛选→深度阅读', 'Cohen\'s Kappa>0.7 + 排除理由10/10合理 + 假阳性0', '未执行（轻量流程）'],
        ['Gate 3', '深度阅读→写作', '笔记质量≥8/10 + 机制类别0空白 + 交叉主题引用100%', '未执行（轻量流程）'],
        ['Gate 4 ★', '写作（正文）', '逐条打开引用摘要验证≥95%通过 + 逻辑连贯7/7 + 数据准确性0事实错误', '✅ 通过（回退后19/20验证）'],
        ['Gate 5 ★', '修改/扩展', '新增声明100%可溯源 + 修改不越界 + 100%可安全回退', '✅ 通过（41/41引用已用）'],
        ['Gate 6', '终稿→投稿', '引用格式100%一致 + 图表嵌入100%对应 + 语言终审≤5标记', '待执行'],
    ]
)

heading('Gate 4 关键发现', 2)
para('首次执行Gate 4时发现严重问题：扩展的14条声明中13条（93%）未被所引文献摘要支撑。根因为Agent从训练数据中提取知识嫁接到不相关引用上（"引用嫁接"）。')
para('修复策略：回退所有无支撑扩展内容 → 逐引用验证 → 仅保留有摘要直接支撑的声明。回退后19/20声明直接验证通过，41篇引用去重后全部有效。')

heading('Gate 4/5 可执行脚本', 2)
para('Gate 4和Gate 5已编码为可执行的Python脚本，集成在 harness/quality-gate.md 和 scripts/gen_word_full.py 中：')
bullet('Gate 4：20项声明 vs PMID摘要交叉验证 → 关键词匹配 → 通过率检查')
bullet('Gate 5：范围引用展开[N-M]、编号一致性、未使用引用检测 → 正文引用集合==列表引用集合')
bullet('Word格式自检（gen_word_full.py内置8项）：Figure refs, Table refs, Bad refs, Body citations, Images embedded, Headings, Word count, References used')

page_break()

# ════════════════════════════════════════════════════════════════════
# 十三、写作纪律与规则
# ════════════════════════════════════════════════════════════════════
heading('十三、写作纪律与规则', 1)

para('基于项目实践中发现的错误模式，在CLAUDE.md中编码了严格的写作纪律：')

heading('单源真理原则', 2)
bullet('稿件内容的唯一源文件是 manuscript/jitc_submission.md')
bullet('Word文档由 scripts/gen_word_full.py 自动生成，该脚本只负责格式化，不包含内容')
bullet('任何修改：先改源文件 → 运行生成器 → 运行自检')

heading('引用铁律', 2)
bullet('每条声明必须有至少一篇引用文献的摘要直接支撑（全文支撑更好）')
bullet('禁止从训练数据中提取知识贴到不相关的引用上（"引用嫁接"）')
bullet('扩展段落前必须先验证引用-声明配对')
bullet('引用非鳞NSCLC文献支撑LUSC论点时，必须加限定语')

heading('图表纪律', 2)
bullet('图表编号必须在源文件中全局唯一')
bullet('删除图表 = 删除所有正文引用 + 更新Title Page声明 + 重新编号')
bullet('自检必须同时验证：正文引用编号、图片标题文字、Word Caption')

heading('喉鳞癌过滤规则', 2)
bullet('筛选时必须显式排除：laryngeal, head and neck, oral, esophageal, cutaneous, cervical, thymic squamous')
bullet('"LSCC"缩写在纳入前必须验证为 lung squamous cell carcinoma')

heading('错误模式库', 2)
para('每遇到一种新错误即追加到CLAUDE.md中，当前包含：')
add_table(
    ['错误模式', '描述', '修复'],
    [
        ['引用嫁接', '训练数据知识 + 不匹配的引用', '回退 + 逐引用验证 Gate 4'],
        ['编号漂移', '删除图表后手工重编号 → 不一致', '自动化自检 + 删除引用同步'],
        ['压缩丢失', '节省token压缩段落 → 内容丢失', '用源文件解析，不做智能压缩'],
        ['范围遗漏', '[8-10]不被正则匹配 → 误报', '展开范围引用再验证'],
    ]
)

page_break()

# ════════════════════════════════════════════════════════════════════
# 十四、数据文件清单
# ════════════════════════════════════════════════════════════════════
heading('十四、数据文件清单', 1)

add_table(
    ['文件名', '大小', '内容', '用途'],
    [
        ['pubmed_merged_all.json', '8.6 MB', '4,106篇全量检索结果', '原始检索存档'],
        ['pubmed_search_results.json', '3.4 MB', 'PubMed初步检索结果', '7个关键词查询(2025-2026)'],
        ['pubmed_search_supplement.json', '5.2 MB', '年份分层补充检索', '2020-2024年份补充'],
        ['pubmed_relevant_for_screening.json', '1.1 MB', '去重后432篇初筛列表', '筛选输入'],
        ['screening_final_40.json', '118 KB', '37篇最终纳入文献', '核心引用库（含摘要）'],
        ['screening_final_80.json', '210 KB', '62篇全文候选列表', 'Round 1输出'],
        ['screening_excluded.json', '415 KB', '排除文献+原因', '排除记录审计'],
        ['screening_included.json', '341 KB', '纳入文献记录', '纳入记录'],
        ['screening_uncertain.json', '373 KB', '不确定文献', '筛选决策记录'],
        ['screening_strict_included.json', '531 KB', '严格纳入（高置信度）', '质量子集'],
        ['screening_final_curated.json', '158 KB', '人工策展纳入列表', '最终人工审核版'],
        ['screening_final_included.json', '558 KB', '纳入文献完整版', '全量纳入数据'],
        ['screening_final_inclusion.json', '118 KB', '最终纳入完整数据', '同screening_final_40'],
        ['batch1_papers.json', '29 KB', 'Batch 1论文元数据', '分批深度阅读'],
        ['batch2_papers.json', '30 KB', 'Batch 2论文元数据', '分批深度阅读'],
        ['batch3_papers.json', '30 KB', 'Batch 3论文元数据', '分批深度阅读'],
        ['screening_log.md', '3 KB', '筛选过程日志', '筛选过程记录'],
        ['fulltext_screening_log.md', '9 KB', '全文筛选日志', '全文筛选记录'],
    ]
)

bullet('数据总量：约18 MB（19个JSON+2个MD文件）')
bullet('所有文献数据包含PMID、标题、摘要、作者、期刊、年份等标准化字段')

page_break()

# ════════════════════════════════════════════════════════════════════
# 十五、脚本工具链
# ════════════════════════════════════════════════════════════════════
heading('十五、脚本工具链', 1)

add_table(
    ['脚本', '大小', '功能', '状态'],
    [
        ['gen_word_full.py', '9.5 KB', '★ 核心：Markdown→Word生成器 + 8项内置自检', '当前使用'],
        ['gen_english_word.py', '43.7 KB', '旧版Word生成器（含硬编码文本——已废弃）', '历史遗留'],
        ['gen_word.py', '52.0 KB', '旧版Word生成器变体', '历史遗留'],
        ['gen_word_v2.py', '35.6 KB', 'v2版Word生成器', '历史遗留'],
        ['generate_word.py', '23.1 KB', '另一个生成器变体', '历史遗留'],
        ['generate_figures.py', '15.3 KB', '图表PNG生成脚本', '辅助工具'],
        ['redraw_figures.py', '12.1 KB', '图表重绘脚本（调整尺寸/颜色）', '辅助工具'],
        ['pubmed_search.py', '4.3 KB', 'PubMed E-utilities API检索', '检索工具'],
        ['pubmed_search_supplement.py', '4.1 KB', '年份分层补充检索', '检索工具'],
        ['fulltext_screening.py', '13.1 KB', '全文筛选自动化', '筛选工具'],
        ['screen_abstracts.py', '14.1 KB', '摘要筛选自动化', '筛选工具'],
    ]
)

heading('gen_word_full.py 核心流程', 2)
code_block('''1. 读取 jitc_submission.md（唯一源文件）
2. 解析Section分离 → 标题页/摘要/Key Messages/正文章节/声明/参考文献
3. 按JITC格式渲染：Times New Roman 12pt, 双倍行距, 1英寸页边距
4. 在指定位置插入Figure/Table图片
5. 自动运行8项自检：
   • Figure引用编号验证
   • Table引用编号验证
   • 异常引用检测
   • 正文引用完整性
   • 图片嵌入数量验证
   • 标题层级格式验证
   • 总字数统计
   • 已使用引用列表''')

page_break()

# ════════════════════════════════════════════════════════════════════
# 十六、项目统计
# ════════════════════════════════════════════════════════════════════
heading('十六、项目统计', 1)

heading('文献统计', 2)
add_table(
    ['指标', '数值'],
    [
        ['全量检索命中', '4,106篇'],
        ['去重后初筛', '432篇'],
        ['全文候选', '62篇'],
        ['PMC开放获取', '60/62 (96.8%)'],
        ['最终纳入', '37篇 (综述8 + 原始研究26 + SR/MA 3)'],
        ['精读笔记', '31篇 (Batch 1+2+3)'],
        ['浅读笔记', '31篇 (62篇全覆盖笔记)'],
        ['交叉主题', '18个'],
    ]
)

heading('稿件统计', 2)
add_table(
    ['指标', '数值'],
    [
        ['总字数', '8,969词 (正文) + 246词 (摘要)'],
        ['章节数', '7个正文章节 + Abstract + Key Messages + Declarations'],
        ['参考文献', '41篇 (含PMID/DOI)'],
        ['图表', '1 Figure + 2 Tables (发布版) / 4 Figures + 4 Tables (全部定义)'],
        ['目标期刊', 'Journal for ImmunoTherapy of Cancer (JITC)'],
    ]
)

heading('文件统计', 2)
add_table(
    ['指标', '数值'],
    [
        ['项目总文件数', '约150个（含62篇论文笔记）'],
        ['数据文件', '19个JSON + 2个Markdown日志'],
        ['脚本文件', '11个Python脚本'],
        ['图片文件', '8个PNG图表 (总计 ~1.2 MB)'],
        ['文档文件', '约30个Markdown（含论文笔记）'],
        ['Word输出', '1个.docx文件'],
    ]
)

heading('Agent体系统计', 2)
add_table(
    ['指标', '数值'],
    [
        ['子Agent总数', '7个 (5横向执行 + 2纵向基础设施)'],
        ['极简触发命令', '15+个'],
        ['Harness维度', '5个 (成功率/效率/鲁棒性/安全性/一致性)'],
        ['质量关卡', '6个 (Gate 1-6)'],
        ['鲁棒性测试场景', '13个 (L1-L5)'],
        ['一致性基准', '5个'],
        ['安全检测规则', '5条 (文件/网络/命令/配置/信息泄露)'],
        ['错误模式', '4个已编码'],
    ]
)

page_break()

# ════════════════════════════════════════════════════════════════════
# 十七、关键决策记录
# ════════════════════════════════════════════════════════════════════
heading('十七、关键决策记录', 1)

add_table(
    ['决策编号', '决策', '选项', '选择', '理由'],
    [
        ['#1', '项目基础设施架构', 'A:单文件CLAUDE.md / B:三层Context体系', 'B 三层体系', '文献量大(50-150篇)，单文件无法承载；分层与Claude Code memory/hooks机制匹配'],
        ['#2', '编码Agent触发方式', 'A:纯手动 / B:Stop hook自动 / C:两者结合', 'C 两者结合', '手动确保用户确认；Stop hook作为兜底防遗忘'],
        ['#3', 'Tier 3知识库存储', 'A:纯MD文件 / B:纯Memory / C:混合方案', 'C 混合方案', 'MD可读可git友好；Memory存元信息提供快速导航'],
        ['待定', '目标期刊选择', '—', '→ JITC', '影响格式要求和写作深度'],
        ['待定', '引用管理工具', 'Zotero vs EndNote vs BibTeX', '待定', '—'],
        ['待定', 'PROSPERO注册', '注册 vs 不注册', '不注册', '此综述为叙述性综述，非系统综述'],
        ['待定', '荟萃分析 vs 叙述性综述', '荟萃分析 vs 叙述性', '叙述性综述', '主题广泛，机制多样性不适合定量合并'],
    ]
)

heading('目标期刊选定：JITC', 2)
para('Journal for ImmunoTherapy of Cancer (JITC) 是BMJ出版的免疫肿瘤学领域高影响力期刊。选定理由：')
bullet('主题高度匹配：JITC专注肿瘤免疫治疗，本综述论ICI耐药机制完全契合')
bullet('接受叙述性综述：JITC发表Review类型文章')
bullet('Vancouver引用格式：与项目使用的引用格式一致')

page_break()

# ════════════════════════════════════════════════════════════════════
# 十八、经验教训
# ════════════════════════════════════════════════════════════════════
heading('十八、经验教训', 1)

para('项目全流程中发现并编码了8条系统化教训，已反馈到Agent系统设计和CLAUDE.md中：')

add_table(
    ['#', '教训', '问题描述', '修复措施'],
    [
        ['1', '单源真理', '两套Word生成脚本（含硬编码内容vs解析markdown），修改不同步导致引用丢失', '所有内容→jitc_submission.md；Word由gen_word_full.py解析生成'],
        ['2', '引用验证前置', '从训练数据提取知识（如alpelisib/CB-839）贴在无关引用上，13/14扩展声明无引文支撑', '声明-引用配对验证为Gate 4硬性要求；禁止从训练数据嫁接到不相关引用'],
        ['3', '图表编号系统级约束', '删减图表后手动重编号，三处(正文/图像/Word Caption)不一致', '删除=删除引用+重编号+自检；自检同时检查三处一致性'],
        ['4', '自检覆盖格式渲染', '只检查文本内容(引用编号)，忽略了Word标题格式(小标题挤在一起)', '自检加：段落分隔、标题段前段后间距、图片清晰度'],
        ['5', '范围引用解析', '评测脚本只解析[N]和[N,M]，不解析[N-M]，refs 9-10被误报', '所有引用解析必须处理[N]/[N,M]/[N-M]三种格式'],
        ['6', '喉鳞癌过滤', '"LSCC"缩写混淆喉鳞癌和肺鳞癌，PMID 42111396错误纳入', '筛选时显式检查laryngeal/head neck/oral/esophageal等标志'],
        ['7', '非鳞引用限定', 'TROPION-Lung10(PMID 41669261)针对non-squamous，用于LUSC论点', '引用非鳞试验必须加限定语；LUSC论点优先LUSC文献'],
        ['8', '扩展前验证', '用户修改建议后扩展段落，新增声明无引文支撑，被迫回退', '扩展前先将新增声明与引用配对验证；未通过=删除'],
    ]
)

heading('错误模式库', 2)
bullet('引用嫁接：声明关键词 vs 引用摘要关键词匹配度 < 30% → Gate 4回退')
bullet('编号漂移：图/表引用集合 != 已声明保留集合 → Gate 5自检')
bullet('范围遗漏：正文有[N-M]但解析器只提取了N → 展开范围重新验证')
bullet('压缩丢失：Word词数 < 源文件词数×0.85 → 检查生成过程')

page_break()

# ════════════════════════════════════════════════════════════════════
# 十九、技术栈
# ════════════════════════════════════════════════════════════════════
heading('十九、技术栈', 1)

add_table(
    ['组件', '技术/工具', '说明'],
    [
        ['AI协作平台', 'Claude Code (CLI)', 'DeepSeek V4 Pro后端，支持自定义Agent和三层上下文体系'],
        ['AI记忆系统', 'Claude Code Memory', 'Tier 2结构化记忆指针，自动加载+按需检索'],
        ['AI Hooks', 'Claude Code Hooks', 'Stop hook提醒编码Agent执行；Settings hook管理权限'],
        ['文献检索', 'PubMed E-utilities API', 'MeSH词表+自由词检索，3 req/sec无API key'],
        ['文献检索', 'Europe PMC REST API', '含预印本bioRxiv/medRxiv，宽松限流'],
        ['文献检索', 'Semantic Scholar API', '语义搜索，TLDR AI摘要，100req/5min限流'],
        ['版本控制', 'Git', '项目初始化+增量进度提交'],
        ['文档生成', 'python-docx', 'Python→Word(.docx)生成，段落格式+图片嵌入+Heading样式'],
        ['图表生成', 'matplotlib / PIL', 'PRISMA流程图+框架图+表格PNG渲染'],
        ['数据格式', 'JSON', '所有检索/筛选/论文数据标准化为JSON'],
        ['文档格式', 'Markdown', '所有项目文档、论文笔记、稿件源文件使用Markdown'],
        ['术语管理', 'docs/glossary.md', '~80个医学术语缩写（研究设计/分子生物学/免疫学/统计学/临床试验）'],
    ]
)

page_break()

# ════════════════════════════════════════════════════════════════════
# 二十、当前状态与后续计划
# ════════════════════════════════════════════════════════════════════
heading('二十、当前状态与后续计划', 1)

heading('当前状态', 2)
add_table(
    ['维度', '状态'],
    [
        ['项目阶段', 'Phase 6: 修改与定稿 (~90%)'],
        ['稿件状态', '初稿完成，41篇引用验证通过 (Gate 4/5 ✅)'],
        ['图表状态', '8张PNG图表已生成，2张嵌入Word'],
        ['待完成', '格式化参考文献（JITC格式）、最终用户审阅、投稿前检查清单'],
        ['阻碍', '无'],
    ]
)

heading('待完成任务', 2)
bullet('6.4 格式化参考文献（目标期刊格式）—— 待选定最终期刊后执行')
bullet('6.6 最终定稿，提交前检查清单 —— 待用户最终审阅')
bullet('Gate 6 终稿检查：引用格式100%一致 + 图表嵌入100%对应 + 语言终审≤5标记')

heading('后续可选任务', 2)
bullet('确定最终投稿期刊并调整格式')
bullet('补充作者信息和利益冲突声明')
bullet('同行评审前的预审（内部/外部专家）')
bullet('Cover Letter撰写')
bullet('补充Tier 2数据库检索（如需提升全面性）')
bullet('注册PROSPERO（如转为系统综述）')

heading('项目成果总结', 2)
para('本项目在一天内（2026-06-04至2026-06-05）完成了从一个空目录到一篇完整英文综述初稿的全流程，包括：')
bullet('✅ 完整的项目基础设施（三层上下文体系 + Git版本控制）')
bullet('✅ 4,106篇文献系统检索与去重')
bullet('✅ 严格两轮筛选（432→62→37篇纳入，0%仅摘要比例）')
bullet('✅ 62篇论文结构化笔记（可追溯、可审计）')
bullet('✅ 18个跨文献交叉主题提取')
bullet('✅ 8,969词英文综述初稿（7章节 + 41引用）')
bullet('✅ 8张专业PNG图表（4 Figure + 4 Table）')
bullet('✅ 完整Harness Engineering质量保证系统（5维度+6关卡+13测试场景+5基准）')
bullet('✅ 8条经验教训编码到Agent系统中')
bullet('✅ 可复用的Agent定义和工作流模板')

page_break()

# ════════════════════════════════════════════════════════════════════
# 二十一、编码Agent问题诊断与改进建议
# ════════════════════════════════════════════════════════════════════
heading('二十一、编码Agent问题诊断与改进建议', 1)

para('经过对编码Agent（Agent 0）的全流程审计，发现其实际执行情况与设计期望之间存在严重差距。核心问题：编码Agent虽然在设计文档中定义完善，但在实际操作中几乎从未被触发执行，导致项目在"可恢复性"方面存在重大风险——如果关闭窗口再新建会话，Agent无法从正确状态接续。')

heading('问题1：CLAUDE.md 启动状态严重过时（致命）', 2)
para('当前 CLAUDE.md 第6行：')
code_block('**当前阶段**: 初始化')
para('实际问题：项目已进行到 Phase 6 revision（初稿完成），但 CLAUDE.md 从未被编码Agent更新。当用户新建会话时，CLAUDE.md 会自动加载到上下文中，Agent 会读取到"当前阶段: 初始化"并认为项目刚刚开始。虽然 CLAUDE.md 中定义了启动行为规则（读取 project-status.md），但：')
bullet('启动行为规则写在 Tier 1 常驻文件中，但"当前阶段: 初始化"也是写在同文件中——存在自相矛盾的信号')
bullet('编码Agent从未更新过 CLAUDE.md 的"项目身份"字段中的以下过时信息：')
bullet('  - 当前阶段: 初始化 → 应为 revision')
bullet('  - 目标期刊: 待定 → 应为 JITC (Journal for ImmunoTherapy of Cancer)')
bullet('  - 综述主题/领域未在 CLAUDE.md 中记录（仅在 memory/active-focus.md 中）')

heading('问题2：SESSION_LOG.md 几乎空白（致命）', 2)
para('编码Agent的设计工作流（Part A）要求"追加 progress/SESSION_LOG.md"。实际情况：')
bullet('只有1条记录：Session 0 — 项目初始化 (2026-06-04)')
bullet('缺失的会话记录：Phase 2 文献搜索、Phase 3 文献筛选、Phase 4 深度阅读（3个Batch）、Phase 5 写作（7个章节+摘要）、Phase 6 修改与定稿（3轮审校+引用核查+图表生成）')
bullet('合计缺失约10+个会话日志条目')
para('后果：新会话无法了解"上次做到了哪里"、"中间经历了什么"、"有哪些陷阱需要注意"。')

heading('问题3：MILESTONES.md 从未更新（严重）', 2)
para('当前 milestones.md 状态：')
add_table(
    ['里程碑', '当前标记', '实际应标记'],
    [
        ['项目基础设施搭建完成', '✅', '✅（正确）'],
        ['检索策略确定', '⬜ TBD', '✅ 2026-06-04 完成'],
        ['文献初筛完成 (Round 1)', '⬜ TBD', '✅ 2026-06-04 完成（432→62）'],
        ['全文筛选完成 (Round 2)', '⬜ TBD', '✅ 2026-06-04 完成（62→37）'],
        ['深度阅读完成', '⬜ TBD', '✅ 2026-06-04 完成（30篇精读）'],
        ['综述大纲确定', '⬜ TBD', '✅ 2026-06-04 完成'],
        ['初稿完成', '⬜ TBD', '✅ 2026-06-04 完成（8,969词）'],
        ['内部审校完成', '⬜ TBD', '✅ 2026-06-05 完成（Gate 4/5通过）'],
        ['提交投稿', '⬜ TBD', '⬜ 待完成'],
    ]
)
para('所有里程碑都被标记为"未开始"，给人感觉项目仍处于早期阶段。')

heading('问题4：metrics-raw.json 仅为模板占位（严重）', 2)
para('编码Agent的 Part B 要求"从会话上下文提取效率指标，写入 progress/metrics-raw.json"。实际文件内容：')
bullet('仅1条记录，标注为"模板占位"')
bullet('所有效率字段为 null：wall_time_sec=null, tool_calls=null, tokens_in=null, tokens_out=null')
bullet('所有安全字段为零：violations=[], critical=0, high=0, medium=0, low=0')
bullet('整个 Phase 2-6 没有任何一条真实的效率数据记录')
para('后果：评估Agent（Agent 5）无法读取任何真实数据来做效率分析和L2判定。Harness Engineering系统的"数据收集层"完全空转。')

heading('问题5：Git提交极为稀少（严重）', 2)
para('编码Agent的 Part D 要求"git commit 所有变更，使用结构化 commit message [phase] 简短描述"。实际情况：')
code_block('37fed49 feat: simplify interaction model — single-character commands\n+ df258d9 feat: enhance literature search with comprehensive coverage\n+ 7a59efa feat: add Harness Engineering — agent quality assurance framework\n+ 6ace12e chore: initialize medical review project infrastructure\n= 仅4次提交')
bullet('Phase 2-6 的所有产出物（数十个文件的新增和修改）从未被编码Agent提交')
bullet('当前 git status 显示大量未跟踪/未提交文件（如果仓库状态干净则是另一回事）')
bullet('一旦文件被误删或修改，没有版本历史可以回溯')

heading('问题6：project-status.md 数据不一致（中等）', 2)
para('project-status.md 的内容实际上有所更新（phase=revision, progress_pct=100），但存在内部不一致：')
bullet('YAML frontmatter 中 last_session_id 字段出现了两次——顶部的"2026-06-05-agent-encoding"和底部的"null"')
bullet('papers_read_in_depth: 30 — 但实际62篇都有笔记（虽然7篇是摘要笔记）')
bullet('papers_included: 37 — 但实际有效纳入是37篇，排除3篇（喉鳞癌1+非鳞NSCLC 2）后用于引用的也是37篇')
bullet('last_update: 2026-06-04 — 但2026-06-05进行了Agent系统编码和Gate 5验证')

heading('问题7：无自动触发机制（设计缺陷）', 2)
para('CLAUDE.md 决策#2记录选择"方案C：Stop hook提醒 + 手动执行"，但 settings.local.json 中只有 permissions，完全没有配置 hooks。这意味着：')
bullet('用户关闭窗口时没有任何提醒')
bullet('用户找不到编码Agent的入口（虽然有"编码"命令，但用户可能不知情）')
bullet('编码Agent完全依赖用户主动记得去触发')

heading('问题8：编码Agent工作流过重（设计缺陷）', 2)
para('编码Agent定义了4个Parts（进度记录+效率数据收集+安全审计+Git提交），每次执行的预计成本很高。在实际的快节奏写作流程中（用户说"写第3节"→Agent写→用户说"审校"→...），编码Agent的完整工作流过于繁重。缺少"轻量级快速编码"模式——只记录核心状态，不做完整审计。')

heading('根因分析', 2)
add_table(
    ['根因', '影响', '严重度'],
    [
        ['编码Agent从未被用户手动触发（用户不知道/忘记了"编码"命令）', '所有Part A-D职责落空', '🔴 致命'],
        ['CLAUDE.md 中有启动行为规则，但"当前阶段"字段写在同文件中且从未更新', '新会话可能错误判断阶段', '🔴 致命'],
        ['settings.local.json 中未配置 Stop hook', '无自动提醒机制', '🟠 严重'],
        ['编码Agent工作流过于复杂（4 Parts）', '执行心理成本高', '🟡 中等'],
        ['项目节奏过快（一日完成全流程），编码被忽略', '边做边忘', '🟡 中等'],
    ]
)

heading('改进建议', 2)

heading('短期修复（立即可行）', 3)

para('1. 更新 CLAUDE.md "项目身份" 字段', bold=True)
code_block('- **当前阶段**: revision\n- **目标期刊**: JITC (Journal for ImmunoTherapy of Cancer)\n- **综述主题**: Mechanisms of ICI Resistance in LUSC (NSCLC鳞癌免疫治疗耐药)\n- **领域**: 肿瘤学 / 免疫学 (肺癌免疫治疗)')
para('')

para('2. 补充 SESSION_LOG.md 批量记录', bold=True)
para('至少追加以下条目：Phase 2 文献搜索、Phase 3 文献筛选、Phase 4 深度阅读(Batch 1-3)、Phase 5 写作、Phase 6 修改与定稿。记录关键完成事项和下一步。')

para('3. 更新 MILESTONES.md', bold=True)
para('将已完成的7个里程碑标记为 ✅ 并填写实际日期（均为2026-06-04或2026-06-05）。')

para('4. 执行一次完整 Git 提交', bold=True)
para('提交所有当前变更，使用结构化 message。')

heading('中期改进（架构层面）', 3)

para('5. 配置 Stop hook 实现自动提醒', bold=True)
code_block('{\n  "hooks": {\n    "Stop": [\n      {\n        "matcher": "",\n        "hooks": [{\n          "type": "command",\n          "command": "cat << EOF\\n🔔 会话即将结束。\\n建议执行: \\"编码\\" 或 \\"记\\" 来保存当前进度。\\nEOF"\n        }]\n      }\n    ]\n  }\n}')

para('6. 拆分编码Agent为"轻量编码"和"完整编码"', bold=True)
bullet('轻量编码（触发命令 "记" / "快记"）：仅执行 Part A（进度记录）+ Git add+commit，每次任务后快速执行')
bullet('完整编码（触发命令 "编码"）：执行全部4 Parts，每个 Phase 结束时执行')
bullet('在关键任务里程碑处自动建议执行轻量编码（写入 CLAUDE.md 中的 Agent 提示）')

para('7. 在 CLAUDE.md 中增加"自愈"启动逻辑', bold=True)
bullet('启动时不仅读取 project-status.md，还要交叉验证 CLAUDE.md 中的"当前阶段"字段')
bullet('如果发现 CLAUDE.md 的"当前阶段"与 project-status.md 的 phase 不一致→以 project-status.md 为准，并提示用户"注意：CLAUDE.md 中的阶段信息过时"')
bullet('如果 MILESTONES.md 与 FEATURE_LIST.md 的完成状态矛盾→以 FEATURE_LIST.md 为准')

heading('长期改进（流程层面）', 3)

para('8. 建立"Phase结束检查清单"', bold=True)
para('每个Phase结束时（用户说"审校"后），Agent主动提醒：')
code_block('📍 Phase X 即将完成。建议检查清单：\n  □ 编码Agent已记录本Phase进度？\n  □ Git已提交本Phase变更？\n  □ MILESTONES已更新？\n  □ SESSION_LOG已追加？\n  □ Gate已通过？')

para('9. 增加会话恢复测试', bold=True)
para('定期执行"冷启动测试"：关闭会话→打开新会话→只读CLAUDE.md→检查Agent是否正确识别项目阶段和下一步任务。将此测试加入Harness Engineering的鲁棒性测试场景库。')

page_break()

# ════════════════════════════════════════════════════════════════════
# APPENDIX: 命令行参考
# ════════════════════════════════════════════════════════════════════
heading('附录A：极简命令表', 1)

add_table(
    ['用户说', '触发Agent', '说明'],
    [
        ['1 / 搜索 / 搜', 'Agent 1 文献搜索', '多数据库系统检索'],
        ['2 / 筛选 / 筛', 'Agent 6 筛选Agent', '两轮筛选(标题/摘要→全文)'],
        ['3 / 分析 / 读', 'Agent 2 论文分析', '结构化提取论文信息'],
        ['4 / 写作 / 写', 'Agent 3 综述写作', '基于笔记撰写综述段落'],
        ['5 / 审校 / 审', 'Agent 4 审校Agent', '事实核查+逻辑+引用溯源'],
        ['6 / 编码 / 记', 'Agent 0 编码Agent', '记录进度+效率+安全+Git提交'],
        ['7 / 评估 / 评', 'Agent 5 评估Agent', 'L2评分+鲁棒+一致性+安全复核'],
        ['状态 / 进度', '—', '展示当前项目状态'],
        ['下一步 / next', '—', '自动判断建议下一步'],
        ['帮助 / ? / help', '—', '显示可用命令'],
        ['就绪 / 好了 / done', '—', '标记手动操作完成'],
    ]
)

heading('附录B：生成Word文档命令', 1)
code_block('cd E:/medical-review\npython scripts/gen_word_full.py')
para('输出文件：E:/medical-review/manuscript/English_Manuscript_NSCLC_ICI_Resistance.docx')
para('运行后自动执行8项自检并在控制台输出结果。')

# ── Save ──
doc.save(OUT)
file_size = os.path.getsize(OUT)
print(f'✅ 项目概览文档已生成：{OUT}')
print(f'   文件大小：{file_size/1024:.0f} KB')
